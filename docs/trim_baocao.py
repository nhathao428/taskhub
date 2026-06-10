"""
Trim BAO_CAO_DO_AN_CO_SO.docx về 30-40 trang.
Strategy:
- Cắt toàn bộ code/config/tree-ASCII trong Chương 4 (giữ ở Phụ lục)
- Trim Chương 2 sections 2.5-2.14 còn 1 đoạn tóm tắt mỗi mục
- Giữ nguyên Chương 1, 3, 5 (đã hợp lý), Ch6
"""
import re
import docx
from docx.oxml.ns import qn

PATH = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'
d = docx.Document(PATH)

# ============================================================
# Helpers
# ============================================================
def heading_level(p):
    name = p.style.name
    if name.startswith('Heading '):
        try: return int(name.split(' ')[1])
        except: return 0
    return 0

def chapter_of_heading(text):
    text = text.upper()
    m = re.match(r'CHƯƠNG (\d)', text)
    if m: return int(m.group(1))
    if 'KẾT LUẬN' in text: return 6
    if 'TÀI LIỆU THAM KHẢO' in text or 'PHỤ LỤC' in text: return 7
    return None

def looks_code(text):
    t = text.strip()
    if not t: return False
    # ASCII tree / box / pipe
    if re.search(r'[│├└─┌┐┘┤┬┴┼]', t): return True
    # Code patterns (Java/JS/Dart/YAML/XML/Dockerfile/properties)
    starters = ('import ', 'package ', '@', 'public ', 'private ', 'const ', 'function ',
                'def ', 'class ', 'interface ', 'export ', 'from ', 'spring:', 'app:',
                'gemini:', 'server:', 'services:', 'networks:', 'FROM ', 'COPY ', 'RUN ',
                'WORKDIR ', 'CMD ', 'ENV ', 'EXPOSE ', '<?xml', '<dependency>', '<plugin>',
                '{', '[', '}', ']', 'jdbc:', 'jwt.', '//', '/*', '#!/', '#include',
                'http.', 'auth.', 'sb.', 'return ', 'if (', 'else ', 'final ',
                'protected ', 'static ', 'void ', 'this.', 'var ', 'let ', 'await ',
                'async ', 'try {', 'catch ', 'throw ', '} catch', '} else', 'mvn ',
                'npm ', 'docker', 'curl ', 'ssh ', '$', 'apiVersion:', 'kind:',
                'metadata:', 'image:', '- name', '- "', '-D', '--', 'DROP TABLE',
                'CREATE TABLE', 'INSERT INTO', 'SELECT ', 'UPDATE ', 'DELETE FROM',
                'ALTER ', 'ON CONFLICT', 'CONSTRAINT ', 'FOREIGN KEY', 'PRIMARY KEY')
    if t.startswith(starters): return True
    # YAML key-value lines that look like config
    if re.match(r'^[a-z][\w\.-]*\s*[:=]\s', t.lower()): return True
    # Property style
    if re.match(r'^[\w.-]+=[\S]', t): return True
    # Properties or HTTP requests
    if re.match(r'^(GET|POST|PUT|DELETE|PATCH)\s+/', t): return True
    # Very long single line ≥250 chars (likely code/data dump)
    if len(t) >= 250 and not re.search(r'[.!?]\s', t): return True
    return False

# ============================================================
# 1. Cắt code blocks ở Chương 4 (giữ Heading)
# ============================================================
in_ch4 = False
removed_ch4 = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    h = heading_level(p)

    if h == 1:
        c = chapter_of_heading(t)
        in_ch4 = (c == 4)
        continue
    if not in_ch4: continue
    # Skip heading lines
    if h in (1, 2, 3, 4): continue
    if not t: continue
    # Skip bullet/list items
    if p.style.name.startswith('List'): continue
    # Cut code-like content
    if looks_code(t):
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el)
            removed_ch4 += 1
print(f'Chương 4: removed {removed_ch4} code/config paragraphs')

# ============================================================
# 2. Trim Chương 2 — mỗi tiểu mục (Heading 3) giữ tối đa 2 đoạn
# ============================================================
in_ch2 = False
para_count_in_h3 = 0
removed_ch2 = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    h = heading_level(p)
    if h == 1:
        c = chapter_of_heading(t)
        in_ch2 = (c == 2)
        para_count_in_h3 = 0
        continue
    if not in_ch2: continue
    if h in (2, 3, 4):
        para_count_in_h3 = 0
        continue
    if not t: continue
    if p.style.name.startswith('List'): continue
    para_count_in_h3 += 1
    if para_count_in_h3 > 2:
        # Remove paragraph
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el)
            removed_ch2 += 1
print(f'Chương 2: removed {removed_ch2} extra paragraphs (keep 2 per Heading 3)')

# ============================================================
# 3. Trim Chương 5 sections 5.2.x test detail (giữ 1 đoạn)
# ============================================================
in_ch5 = False
in_52 = False
para_count = 0
removed_ch5 = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    h = heading_level(p)
    if h == 1:
        c = chapter_of_heading(t)
        in_ch5 = (c == 5)
        in_52 = False
        continue
    if not in_ch5: continue
    if h == 2:
        in_52 = t.startswith('5.2.')
        para_count = 0
        continue
    if h == 3:
        para_count = 0
        continue
    if not in_52: continue
    if not t: continue
    if p.style.name.startswith('List'): continue
    para_count += 1
    if para_count > 1:
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el)
            removed_ch5 += 1
print(f'Chương 5.2.x: removed {removed_ch5} extra test-detail paragraphs')

d.save(PATH)
print(f'\nSaved → {PATH}')
