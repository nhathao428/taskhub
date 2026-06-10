# -*- coding: utf-8 -*-
"""
Sinh BÁO CÁO ĐỒ ÁN CƠ SỞ — BẢN CƠ BẢN (Times New Roman 13, line 1.5, lề T2 B2 L3 R2).
Phạm vi: phần LÕI (xác thực/phân quyền JWT, CRUD nhân viên–dự án–công việc, chấm công,
dashboard) + module AI GỢI Ý NHÂN VIÊN (đã đăng ký trong đề cương đồ án cơ sở).
Các tính năng nâng cao khác (GPS/geofence, Redis, rate-limit, mobile, cloud) chỉ nêu ở
mục "Hướng phát triển" để dành cho Đồ án chuyên ngành.
Tái dùng bộ helper định dạng từ build_baocao_docx.py.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import os as _os
OUTPUT = r"C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx"
SHOTS_DIR = r"C:\Users\Admin\taskhub\docs\screenshots"
UML_DIR = r"C:\Users\Admin\taskhub\docs\uml\png"

FONT = "Times New Roman"
SIZE_BODY = 13
SIZE_H1 = 14

# ------------------------------------------------------------------ HELPERS
def set_run(run, size=SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_p(doc, text="", size=SIZE_BODY, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=0.0,
          space_before=0, space_after=0, line=1.5):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    set_run(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def _style_set_font(style, size, bold=False, italic=False):
    f = style.font
    f.name = FONT; f.size = Pt(size); f.bold = bold; f.italic = italic
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(a), FONT)


def configure_heading_styles(doc):
    h1 = doc.styles["Heading 1"]; _style_set_font(h1, SIZE_H1, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(0); h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]; _style_set_font(h2, SIZE_BODY, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(12); h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]; _style_set_font(h3, SIZE_BODY, bold=True, italic=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(10); h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True


def add_h1(doc, text):
    p = doc.add_paragraph(style="Heading 1"); p.add_run(text.upper()); return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2"); p.add_run(text); return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3"); p.add_run(text); return p


def add_toc_field(doc, levels="1-3"):
    p = doc.add_paragraph(); run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); b.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t"); ph.text = "Nhấn F9 (chuột phải → Update Field) trong Word để cập nhật mục lục."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (b, instr, sep, ph, end):
        run._element.append(el)
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        u = OxmlElement("w:updateFields"); u.set(qn("w:val"), "true"); settings.append(u)
    return p


def add_para(doc, text, indent=0.75):
    return add_p(doc, text, indent_first=indent)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(text), size=SIZE_BODY)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), size=SIZE_BODY, italic=True, bold=True)
    return p


def _add_pic(doc, base_dir, filename, width_cm, missing_label):
    path = _os.path.join(base_dir, filename)
    if not _os.path.exists(path):
        add_p(doc, f"[{missing_label}: {filename}]", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Cm(width_cm))
    return p


def add_image(doc, filename, width_cm=14.5):
    return _add_pic(doc, SHOTS_DIR, filename, width_cm, "Không tìm thấy ảnh")


def add_uml_image(doc, filename, width_cm=15.0):
    return _add_pic(doc, UML_DIR, filename, width_cm, "Không tìm thấy sơ đồ")


def add_figure(doc, base, filename, number, caption, width_cm=14.5):
    (add_uml_image if base == "uml" else add_image)(doc, filename, width_cm)
    add_caption(doc, f"Hình {number}. {caption}")


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = ""
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(c.paragraphs[0].add_run(h), size=SIZE_BODY, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            c = tbl.rows[r_idx].cells[c_idx]; c.text = ""
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run(c.paragraphs[0].add_run(str(val)), size=SIZE_BODY)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    add_p(doc, "", space_after=4)
    return tbl


def add_table_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(text), size=SIZE_BODY, italic=True, bold=True)
    return p


def set_page_number(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def add_footer_page_number(section, show=True):
    footer = section.footer; footer.is_linked_to_previous = False
    for p in list(footer.paragraphs):
        p.clear()
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if show:
        run = p.add_run(); set_run(run, size=SIZE_BODY)
        f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
        f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
        for el in (f1, it, f2):
            run._element.append(el)


def new_section(doc, fmt="decimal", start=None, show_pn=True):
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    s.top_margin = Cm(2); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)
    s.header.is_linked_to_previous = False; s.footer.is_linked_to_previous = False
    set_page_number(s, fmt=fmt, start=start); add_footer_page_number(s, show=show_pn)
    return s


# ------------------------------------------------------------------ DOC INIT
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = FONT; normal.font.size = Pt(SIZE_BODY)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.paragraph_format.line_spacing = 1.5; normal.paragraph_format.space_after = Pt(0)
configure_heading_styles(doc)

sec1 = doc.sections[0]
sec1.top_margin = Cm(2); sec1.bottom_margin = Cm(2.5); sec1.left_margin = Cm(3); sec1.right_margin = Cm(2)
add_footer_page_number(sec1, show=False)


# ------------------------------------------------------------------ TRANG BÌA
def add_cover_page():
    lg = doc.add_paragraph(); lg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lg.paragraph_format.space_after = Pt(6)
    if _os.path.exists("hutech_logo.png"):
        lg.add_run().add_picture("hutech_logo.png", width=Cm(2.8))

    def c(text, size, bold=True, italic=False, after=0):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        set_run(p.add_run(text), size=size, bold=bold, italic=italic)

    c("BỘ GIÁO DỤC VÀ ĐÀO TẠO", 13)
    c("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH", 14)
    c("KHOA CÔNG NGHỆ THÔNG TIN", 13, after=0)
    add_p(doc, "─" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    c("ĐỒ ÁN MÔN HỌC:", 16)
    c("ĐỒ ÁN CƠ SỞ", 20, after=24)
    c("TÊN ĐỀ TÀI:", 14)
    c("HỆ THỐNG QUẢN LÝ CÔNG VIỆC CHO", 18)
    c("DOANH NGHIỆP NHỎ ĐA NGÀNH TÍCH HỢP AI", 18, after=20)

    info = [
        ("Ngành:", "CÔNG NGHỆ THÔNG TIN"),
        ("Lớp:", "23DTHC1"),
        ("Giảng viên hướng dẫn:", "ThS. DƯƠNG THÀNH PHẾT"),
        ("Sinh viên thực hiện:", "NGUYỄN NHẬT HÀO"),
        ("Mã số sinh viên:", "2380612688"),
    ]
    for label, value in info:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        set_run(p.add_run(label.ljust(28)), size=14, bold=True)
        set_run(p.add_run(value), size=14, bold=True)

    add_p(doc, "", space_after=18)
    c("TP. Hồ Chí Minh, tháng 5 năm 2026", 14, italic=True)


add_cover_page()
doc.add_page_break()
add_cover_page()


# ============================ FRONT MATTER (La Mã i, ii, …) ============================
new_section(doc, fmt="lowerRoman", start=1, show_pn=True)

# --------- LỜI CAM ĐOAN ---------
add_h1(doc, "LỜI CAM ĐOAN")
for t in [
    "Em xin cam đoan đồ án “Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ Đa ngành Tích hợp AI” là công trình nghiên cứu của riêng em, được thực hiện dưới sự hướng dẫn của ThS. Dương Thành Phết.",
    "Các nội dung, số liệu, kết quả trình bày trong đồ án là trung thực và do chính em thực hiện. Những phần tham khảo từ tài liệu, công trình của các tác giả khác đều được trích dẫn đầy đủ và ghi rõ nguồn trong mục Tài liệu tham khảo.",
    "Em xin hoàn toàn chịu trách nhiệm trước Khoa và Nhà trường nếu có bất kỳ sự gian lận hay sao chép nào trong đồ án này.",
]:
    add_para(doc, t)
add_p(doc, "", space_after=24)
add_p(doc, "TP. Hồ Chí Minh, tháng 5 năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=2)
add_p(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=36)
add_p(doc, "Nguyễn Nhật Hào", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

# --------- LỜI CẢM ƠN ---------
add_h1(doc, "LỜI CẢM ƠN")
for t in [
    "Lời đầu tiên, em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến ThS. Dương Thành Phết – giảng viên hướng dẫn môn Đồ án cơ sở. Thầy đã tận tình hướng dẫn, định hướng giải pháp và đưa ra những góp ý quý báu giúp em hoàn thành đồ án này.",
    "Em xin cảm ơn Ban Giám hiệu Trường Đại học Công nghệ TP.HCM cùng quý thầy cô Khoa Công nghệ Thông tin đã tạo điều kiện học tập thuận lợi và trang bị cho em nền tảng kiến thức về Lập trình hướng đối tượng, Cơ sở dữ liệu, Kỹ thuật phần mềm, Lập trình Web và Trí tuệ nhân tạo để thực hiện đề tài.",
    "Cuối cùng, em xin cảm ơn gia đình và bạn bè đã luôn động viên, hỗ trợ em trong suốt quá trình học tập và thực hiện đồ án. Do thời gian và kinh nghiệm còn hạn chế, đồ án không tránh khỏi thiếu sót, em rất mong nhận được sự góp ý của quý thầy cô.",
    "Em xin chân thành cảm ơn!",
]:
    add_para(doc, t)
add_p(doc, "", space_after=24)
add_p(doc, "TP. Hồ Chí Minh, tháng 5 năm 2026", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=2)
add_p(doc, "Sinh viên thực hiện", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=36)
add_p(doc, "Nguyễn Nhật Hào", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)

# --------- LỜI MỞ ĐẦU ---------
add_h1(doc, "LỜI MỞ ĐẦU")
for t in [
    "Trong xu thế chuyển đổi số hiện nay, việc ứng dụng phần mềm vào quản lý nhân sự và công việc ngày càng trở nên cần thiết đối với các doanh nghiệp, đặc biệt là doanh nghiệp nhỏ. Phần lớn doanh nghiệp nhỏ vẫn quản lý nhân viên, dự án và chấm công bằng sổ sách hoặc bảng tính Excel rời rạc, dẫn đến khó theo dõi tiến độ, dễ sai sót và thiếu minh bạch. Bên cạnh đó, việc phân công nhân viên cho từng công việc thường dựa vào cảm tính của người quản lý.",
    "Xuất phát từ nhu cầu thực tế đó, đề tài “Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ Đa ngành Tích hợp AI” được lựa chọn nhằm xây dựng một ứng dụng web giúp doanh nghiệp quản lý tập trung nhân viên, dự án, công việc và chấm công, phân quyền rõ ràng cho ba vai trò; đồng thời tích hợp module trí tuệ nhân tạo gợi ý nhân viên phù hợp nhất cho mỗi công việc dựa trên dữ liệu thực tế.",
    "Trong phạm vi Đồ án cơ sở, báo cáo trình bày phần lõi của hệ thống gồm xác thực – phân quyền, quản lý nhân viên – dự án – công việc, chấm công, thống kê và module AI gợi ý nhân viên. Một số tính năng mở rộng khác sẽ được tiếp tục nghiên cứu trong Đồ án chuyên ngành. Báo cáo được tổ chức thành sáu chương, trình bày từ tổng quan, cơ sở lý thuyết, phân tích – thiết kế, xây dựng ứng dụng, kiểm thử đến kết luận.",
    "Em xin trân trọng giới thiệu báo cáo và mong nhận được những ý kiến đóng góp quý báu của quý thầy cô.",
]:
    add_para(doc, t)

# --------- MỤC LỤC ---------
add_h1(doc, "MỤC LỤC")
add_toc_field(doc, levels="1-3")

# --------- DANH MỤC TỪ VIẾT TẮT ---------
add_h1(doc, "DANH MỤC CÁC TỪ VIẾT TẮT")
add_table(doc, ["Từ viết tắt", "Diễn giải"], [
    ("AI", "Artificial Intelligence – Trí tuệ nhân tạo"),
    ("API", "Application Programming Interface – Giao diện lập trình ứng dụng"),
    ("BCrypt", "Thuật toán băm mật khẩu dựa trên Blowfish"),
    ("CRUD", "Create, Read, Update, Delete – Bốn thao tác cơ bản với dữ liệu"),
    ("CSDL", "Cơ sở dữ liệu"),
    ("DTO", "Data Transfer Object – Đối tượng truyền dữ liệu"),
    ("ERD", "Entity Relationship Diagram – Sơ đồ quan hệ thực thể"),
    ("HTTP", "HyperText Transfer Protocol – Giao thức truyền siêu văn bản"),
    ("JPA", "Java Persistence API – Chuẩn ánh xạ đối tượng – quan hệ của Java"),
    ("JSON", "JavaScript Object Notation – Định dạng trao đổi dữ liệu"),
    ("JWT", "JSON Web Token – Chuẩn token dùng cho xác thực"),
    ("LLM", "Large Language Model – Mô hình ngôn ngữ lớn"),
    ("ORM", "Object–Relational Mapping – Ánh xạ đối tượng – quan hệ"),
    ("REST", "Representational State Transfer – Phong cách kiến trúc API"),
    ("SPA", "Single Page Application – Ứng dụng web một trang"),
    ("SQL", "Structured Query Language – Ngôn ngữ truy vấn cấu trúc"),
    ("UML", "Unified Modeling Language – Ngôn ngữ mô hình hóa thống nhất"),
    ("UI", "User Interface – Giao diện người dùng"),
], col_widths=[3.0, 12.5])

# --------- DANH MỤC BẢNG ---------
add_h1(doc, "DANH MỤC CÁC BẢNG")
add_table(doc, ["Số hiệu", "Tên bảng", "Trang"], [
    ("Bảng 2.1", "Các công nghệ sử dụng trong hệ thống", "…"),
    ("Bảng 3.1", "Danh sách yêu cầu chức năng", "…"),
    ("Bảng 3.2", "Danh sách yêu cầu phi chức năng", "…"),
    ("Bảng 3.3", "Đặc tả use case Đăng nhập", "…"),
    ("Bảng 3.4", "Đặc tả use case Tạo công việc và phân công", "…"),
    ("Bảng 3.5", "Đặc tả use case AI gợi ý nhân viên", "…"),
    ("Bảng 3.6", "Mô tả bảng users", "…"),
    ("Bảng 3.7", "Mô tả bảng employees", "…"),
    ("Bảng 3.8", "Mô tả bảng projects", "…"),
    ("Bảng 3.9", "Mô tả bảng tasks", "…"),
    ("Bảng 3.10", "Mô tả bảng attendance", "…"),
    ("Bảng 3.11", "Mô tả bảng skills", "…"),
    ("Bảng 3.12", "Mô tả bảng suggestions", "…"),
    ("Bảng 3.13", "Các tiêu chí dữ liệu đưa vào AI gợi ý nhân viên", "…"),
    ("Bảng 4.1", "Môi trường và công cụ triển khai", "…"),
    ("Bảng 5.1", "Một số kịch bản kiểm thử chức năng", "…"),
    ("Bảng 5.2", "Tổng hợp kết quả kiểm thử", "…"),
], col_widths=[2.5, 11.0, 2.0])

# --------- DANH MỤC HÌNH ---------
add_h1(doc, "DANH MỤC CÁC HÌNH ẢNH, SƠ ĐỒ")
add_table(doc, ["Số hiệu", "Tên hình", "Trang"], [
    ("Hình 2.1", "Mô hình kiến trúc tổng thể của hệ thống", "…"),
    ("Hình 3.1", "Sơ đồ use case tổng thể", "…"),
    ("Hình 3.2", "Sơ đồ use case chức năng xác thực", "…"),
    ("Hình 3.3", "Sơ đồ use case quản lý dự án – công việc", "…"),
    ("Hình 3.4", "Sơ đồ use case AI gợi ý nhân viên", "…"),
    ("Hình 3.5", "Sơ đồ hoạt động chức năng đăng nhập", "…"),
    ("Hình 3.6", "Sơ đồ hoạt động quản lý công việc", "…"),
    ("Hình 3.7", "Sơ đồ hoạt động AI gợi ý nhân viên", "…"),
    ("Hình 3.8", "Sơ đồ tuần tự chức năng đăng nhập", "…"),
    ("Hình 3.9", "Sơ đồ tuần tự AI gợi ý nhân viên", "…"),
    ("Hình 3.10", "Sơ đồ quan hệ thực thể (ERD)", "…"),
    ("Hình 4.1", "Giao diện đăng nhập", "…"),
    ("Hình 4.2", "Giao diện đăng ký", "…"),
    ("Hình 4.3", "Giao diện Dashboard quản lý", "…"),
    ("Hình 4.4", "Giao diện quản lý nhân viên", "…"),
    ("Hình 4.5", "Giao diện quản lý dự án", "…"),
    ("Hình 4.6", "Giao diện quản lý công việc", "…"),
    ("Hình 4.7", "Giao diện chấm công", "…"),
    ("Hình 4.8", "Trang cá nhân của nhân viên", "…"),
    ("Hình 4.9", "Danh sách công việc của nhân viên", "…"),
    ("Hình 4.10", "Lịch sử chấm công của nhân viên", "…"),
    ("Hình 4.11", "Giao diện yêu cầu AI gợi ý nhân viên", "…"),
    ("Hình 4.12", "Kết quả AI gợi ý nhân viên cho công việc", "…"),
], col_widths=[2.5, 11.0, 2.0])


# ============================ NỘI DUNG (Ả Rập 1, 2, …) ============================
new_section(doc, fmt="decimal", start=1, show_pn=True)

# ====================== CHƯƠNG 1 ======================
add_h1(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI")

add_h2(doc, "1.1. Lý do chọn đề tài")
add_para(doc, "Doanh nghiệp nhỏ và vừa (DNNVV) giữ vai trò quan trọng trong nền kinh tế Việt Nam. Theo Sách trắng Doanh nghiệp Việt Nam năm 2022 của Tổng cục Thống kê, DNNVV chiếm khoảng 98% tổng số doanh nghiệp đang hoạt động, đóng góp khoảng 40–45% GDP và thu hút gần 60% lực lượng lao động xã hội [1]. Tuy giữ vai trò lớn như vậy, phần lớn DNNVV, đặc biệt là các doanh nghiệp nhỏ đa ngành, vẫn quản lý nhân sự và công việc theo cách thủ công bằng sổ sách, bảng tính Excel hoặc các nhóm chat. Cách làm này khiến dữ liệu phân tán, khó theo dõi tiến độ công việc, khó tổng hợp báo cáo và dễ phát sinh sai sót khi tính công. Đặc biệt, việc phân công nhân viên cho từng công việc thường dựa vào cảm tính, thiếu cơ sở dữ liệu để ra quyết định.")
add_para(doc, "Trong những năm gần đây, trí tuệ nhân tạo (AI) ngày càng được ứng dụng rộng rãi trong quản trị nguồn nhân lực. Nghiên cứu tổng quan hệ thống của Ekuma (2024) chỉ ra rằng AI và tự động hóa giúp giảm bớt các tác vụ lặp lại, hỗ trợ nhà quản lý ra quyết định dựa trên dữ liệu và nâng cao hiệu quả phát triển nguồn nhân lực [2]. Tương tự, nghiên cứu đăng trên tạp chí Frontiers in Psychology (2024) cho thấy AI mang lại nhiều cơ hội cho các hoạt động nhân sự như tuyển dụng, đánh giá năng lực và phân bổ công việc, giúp giảm sự phụ thuộc vào cảm tính của con người [3]. Đây là cơ sở khoa học cho thấy việc đưa AI vào hỗ trợ phân công nhân viên là một hướng đi có giá trị thực tiễn.")
add_para(doc, "Xuất phát từ thực tiễn và các nghiên cứu nêu trên, việc xây dựng một phần mềm quản lý công việc tập trung, có thêm khả năng gợi ý nhân viên phù hợp bằng AI, sẽ giúp doanh nghiệp số hóa quy trình, lưu trữ dữ liệu nhất quán, phân quyền rõ ràng và hỗ trợ người quản lý ra quyết định phân công hợp lý hơn. Đây cũng là cơ hội để vận dụng các kiến thức đã học về lập trình hướng đối tượng, cơ sở dữ liệu, kỹ thuật phần mềm, lập trình web và trí tuệ nhân tạo vào một sản phẩm thực tế. Vì những lý do trên, em chọn đề tài “Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ Đa ngành Tích hợp AI”.")

add_h2(doc, "1.2. Mục tiêu của đề tài")
add_para(doc, "Mục tiêu của đồ án là xây dựng một ứng dụng web quản lý công việc hoàn chỉnh ở mức cơ bản, đáp ứng các yêu cầu nghiệp vụ cốt lõi của một doanh nghiệp nhỏ và có tích hợp tính năng gợi ý nhân viên bằng AI. Cụ thể, hệ thống cần đạt được các mục tiêu sau:")
for t in [
    "Cho phép đăng ký, đăng nhập và phân quyền cho ba vai trò: Quản trị viên (Admin), Quản lý (Manager) và Nhân viên (Employee).",
    "Quản lý (thêm, sửa, xóa, tìm kiếm) thông tin nhân viên, dự án và công việc.",
    "Phân công công việc cho nhân viên và theo dõi trạng thái hoàn thành.",
    "Ghi nhận và tra cứu dữ liệu chấm công vào/ra theo ngày.",
    "Tích hợp module AI gợi ý danh sách nhân viên phù hợp nhất cho mỗi công việc kèm lý do.",
    "Cung cấp trang thống kê tổng quan giúp người quản lý nắm bắt tình hình công việc.",
]:
    add_bullet(doc, t)

add_h2(doc, "1.3. Đối tượng và phạm vi nghiên cứu")
add_h3(doc, "1.3.1. Đối tượng nghiên cứu")
add_para(doc, "Đối tượng nghiên cứu của đề tài là quy trình quản lý nhân viên, dự án, công việc và chấm công trong doanh nghiệp nhỏ; cùng các công nghệ phát triển ứng dụng web gồm Spring Boot (backend), React (frontend), hệ quản trị cơ sở dữ liệu PostgreSQL và mô hình ngôn ngữ lớn Google Gemini phục vụ gợi ý nhân viên.")
add_h3(doc, "1.3.2. Phạm vi nghiên cứu")
add_para(doc, "Trong khuôn khổ Đồ án cơ sở, đề tài tập trung xây dựng phần lõi của hệ thống gồm: xác thực và phân quyền, quản lý nhân viên – dự án – công việc, chấm công, thống kê và module AI gợi ý nhân viên. Một số tính năng mở rộng khác như xác thực chấm công bằng định vị GPS, bộ nhớ đệm, ứng dụng di động và triển khai trên môi trường đám mây sẽ được tiếp tục nghiên cứu, phát triển trong giai đoạn Đồ án chuyên ngành.")

add_h2(doc, "1.4. Phương pháp thực hiện")
for t in [
    "Phương pháp nghiên cứu tài liệu: tìm hiểu lý thuyết về kiến trúc ứng dụng web, REST API, xác thực JWT, ORM, hệ quản trị cơ sở dữ liệu quan hệ và mô hình ngôn ngữ lớn.",
    "Phương pháp phân tích – thiết kế: khảo sát yêu cầu, mô hình hóa hệ thống bằng UML (use case, activity, sequence) và thiết kế cơ sở dữ liệu.",
    "Phương pháp thực nghiệm: lập trình, tích hợp module AI và kiểm thử hệ thống theo từng chức năng.",
]:
    add_bullet(doc, t)

add_h2(doc, "1.5. Cấu trúc đồ án")
add_para(doc, "Báo cáo được tổ chức thành sáu chương:")
for t in [
    "Chương 1. Tổng quan đề tài: trình bày lý do chọn đề tài, mục tiêu, đối tượng, phạm vi và phương pháp thực hiện.",
    "Chương 2. Cơ sở lý thuyết: giới thiệu các công nghệ và khái niệm nền tảng, bao gồm cả trí tuệ nhân tạo.",
    "Chương 3. Phân tích và thiết kế hệ thống: trình bày yêu cầu, các sơ đồ UML, thiết kế cơ sở dữ liệu và thiết kế module AI.",
    "Chương 4. Xây dựng ứng dụng: giới thiệu môi trường triển khai và các giao diện chính.",
    "Chương 5. Kiểm thử và đánh giá kết quả: trình bày phương pháp, kịch bản và kết quả kiểm thử.",
    "Chương 6. Kết luận và hướng phát triển: tổng kết kết quả đạt được và đề xuất hướng mở rộng.",
]:
    add_bullet(doc, t)


# ====================== CHƯƠNG 2 ======================
add_h1(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT")

add_h2(doc, "2.1. Kiến trúc ứng dụng web")
add_para(doc, "Hệ thống được xây dựng theo mô hình kiến trúc ba lớp (3-tier), tách biệt rõ ràng giữa giao diện người dùng, xử lý nghiệp vụ và lưu trữ dữ liệu. Tầng giao diện (Client) là ứng dụng web viết bằng React, giao tiếp với máy chủ thông qua REST API. Tầng ứng dụng (Application) là máy chủ Spring Boot chịu trách nhiệm xử lý toàn bộ logic nghiệp vụ và bảo mật, đồng thời gọi tới dịch vụ AI bên ngoài. Tầng dữ liệu (Data) là hệ quản trị cơ sở dữ liệu PostgreSQL lưu trữ dữ liệu của hệ thống.")
add_figure(doc, "uml", "architecture.png", "2.1", "Mô hình kiến trúc tổng thể của hệ thống", width_cm=15.0)
add_para(doc, "Việc phân tách thành ba lớp giúp hệ thống dễ bảo trì, dễ mở rộng và cho phép phát triển độc lập giữa giao diện và máy chủ.")

add_h2(doc, "2.2. Spring Boot và REST API")
add_h3(doc, "2.2.1. Spring Boot")
add_para(doc, "Spring Boot là một framework mã nguồn mở thuộc hệ sinh thái Spring của ngôn ngữ Java, giúp xây dựng ứng dụng web và dịch vụ một cách nhanh chóng. Spring Boot cung cấp cơ chế tự động cấu hình (auto-configuration), máy chủ web nhúng và quản lý phụ thuộc thông qua Maven, giúp lập trình viên tập trung vào nghiệp vụ thay vì cấu hình thủ công.")
add_h3(doc, "2.2.2. RESTful API")
add_para(doc, "REST (Representational State Transfer) là phong cách kiến trúc cho các dịch vụ web, trong đó mỗi tài nguyên được định danh bằng một đường dẫn URL và được thao tác thông qua các phương thức HTTP như GET (lấy dữ liệu), POST (tạo mới), PUT (cập nhật) và DELETE (xóa). Dữ liệu trao đổi giữa máy khách và máy chủ thường ở định dạng JSON. Trong hệ thống, backend Spring Boot cung cấp các REST API để frontend gọi đến và hiển thị dữ liệu.")

add_h2(doc, "2.3. Xác thực và phân quyền với JWT")
add_para(doc, "JWT (JSON Web Token) là một chuẩn mở dùng để truyền thông tin xác thực dưới dạng một chuỗi token được ký số. Khi người dùng đăng nhập thành công, máy chủ tạo ra một token chứa thông tin định danh và vai trò của người dùng, sau đó gửi về cho máy khách. Ở những yêu cầu tiếp theo, máy khách đính kèm token này trong phần header để máy chủ xác thực mà không cần lưu phiên đăng nhập trên máy chủ.")
add_para(doc, "Hệ thống sử dụng Spring Security kết hợp với JWT để bảo vệ các REST API. Mỗi yêu cầu sẽ đi qua một bộ lọc kiểm tra token; nếu hợp lệ, thông tin người dùng được nạp vào ngữ cảnh bảo mật và hệ thống sẽ kiểm tra vai trò để quyết định cho phép hay từ chối truy cập. Mật khẩu người dùng được mã hóa bằng thuật toán băm BCrypt trước khi lưu vào cơ sở dữ liệu nhằm đảm bảo an toàn.")

add_h2(doc, "2.4. JPA/Hibernate và ánh xạ ORM")
add_para(doc, "JPA (Java Persistence API) là chuẩn ánh xạ đối tượng – quan hệ (ORM) của Java, cho phép thao tác với cơ sở dữ liệu thông qua các đối tượng Java thay vì viết câu lệnh SQL thủ công. Hibernate là một bộ hiện thực phổ biến của JPA. Mỗi lớp thực thể (Entity) trong mã nguồn được ánh xạ tới một bảng trong cơ sở dữ liệu, mỗi thuộc tính ánh xạ tới một cột. Nhờ đó, các thao tác thêm, sửa, xóa, truy vấn dữ liệu trở nên ngắn gọn và an toàn hơn.")

add_h2(doc, "2.5. Hệ quản trị cơ sở dữ liệu PostgreSQL")
add_para(doc, "PostgreSQL là một hệ quản trị cơ sở dữ liệu quan hệ mã nguồn mở mạnh mẽ, hỗ trợ đầy đủ chuẩn SQL, ràng buộc toàn vẹn dữ liệu, khóa chính – khóa ngoại và giao dịch. Hệ thống sử dụng PostgreSQL để lưu trữ dữ liệu người dùng, nhân viên, dự án, công việc, chấm công và gợi ý. Quy ước đặt tên bảng và cột theo kiểu snake_case.")

add_h2(doc, "2.6. Công nghệ phía giao diện (Frontend)")
add_para(doc, "Giao diện người dùng được xây dựng bằng thư viện React kết hợp công cụ đóng gói Vite và framework CSS Tailwind. React cho phép xây dựng giao diện theo hướng thành phần (component) và là ứng dụng một trang (SPA), giúp trải nghiệm người dùng mượt mà. Thư viện Axios được dùng để gọi REST API kèm token xác thực, còn Chart.js được dùng để vẽ biểu đồ thống kê.")
add_table_caption(doc, "Bảng 2.1. Các công nghệ sử dụng trong hệ thống")
add_table(doc, ["Thành phần", "Công nghệ", "Vai trò"], [
    ("Backend", "Java, Spring Boot, Maven", "Cung cấp REST API và xử lý nghiệp vụ"),
    ("Bảo mật", "Spring Security, JWT, BCrypt", "Xác thực, phân quyền, mã hóa mật khẩu"),
    ("ORM", "Spring Data JPA / Hibernate", "Ánh xạ đối tượng – cơ sở dữ liệu"),
    ("Frontend", "React, Vite, Tailwind CSS", "Xây dựng giao diện người dùng"),
    ("Biểu đồ", "Chart.js", "Vẽ biểu đồ thống kê trên Dashboard"),
    ("Cơ sở dữ liệu", "PostgreSQL", "Lưu trữ dữ liệu hệ thống"),
    ("Trí tuệ nhân tạo", "Google Gemini (gemini-2.5-flash)", "Gợi ý nhân viên phù hợp cho công việc"),
], col_widths=[3.5, 5.5, 6.5])

add_h2(doc, "2.7. Trí tuệ nhân tạo và mô hình ngôn ngữ lớn")
add_para(doc, "Trí tuệ nhân tạo (AI) là lĩnh vực nghiên cứu giúp máy tính thực hiện những công việc đòi hỏi trí thông minh của con người như suy luận, phân tích và ra quyết định. Mô hình ngôn ngữ lớn (Large Language Model – LLM) là một dạng mô hình AI được huấn luyện trên khối lượng dữ liệu văn bản khổng lồ, có khả năng hiểu ngữ cảnh và sinh ra văn bản tự nhiên.")
add_para(doc, "Trong đề tài, hệ thống sử dụng mô hình Google Gemini (phiên bản gemini-2.5-flash) thông qua API để thực hiện chức năng gợi ý nhân viên. Backend thu thập dữ liệu làm việc thực tế của các nhân viên, xây dựng một đoạn mô tả (prompt) bằng tiếng Việt kèm các tiêu chí ưu tiên, sau đó gửi cho Gemini. Mô hình sẽ phân tích và xếp hạng những nhân viên phù hợp nhất cho công việc, kèm theo lý do bằng ngôn ngữ tự nhiên. Việc đưa AI vào hệ thống giúp người quản lý có thêm cơ sở tham khảo khách quan khi phân công, thay vì chỉ dựa vào cảm tính.")


# ====================== CHƯƠNG 3 ======================
add_h1(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

add_h2(doc, "3.1. Phân tích yêu cầu")
add_h3(doc, "3.1.1. Yêu cầu chức năng")
add_para(doc, "Trên cơ sở khảo sát nghiệp vụ quản lý công việc của doanh nghiệp nhỏ, hệ thống cần đáp ứng các yêu cầu chức năng được liệt kê trong Bảng 3.1.")
add_table_caption(doc, "Bảng 3.1. Danh sách yêu cầu chức năng")
add_table(doc, ["Mã", "Chức năng", "Vai trò sử dụng"], [
    ("CN-01", "Đăng ký, đăng nhập, đăng xuất", "Tất cả"),
    ("CN-02", "Phân quyền theo ba vai trò Admin/Manager/Employee", "Hệ thống"),
    ("CN-03", "Quản lý nhân viên (thêm/sửa/xóa/tìm kiếm)", "Admin, Manager"),
    ("CN-04", "Quản lý dự án", "Admin, Manager"),
    ("CN-05", "Quản lý công việc và phân công nhân viên", "Admin, Manager"),
    ("CN-06", "Cập nhật trạng thái công việc được giao", "Employee"),
    ("CN-07", "Chấm công vào/ra và xem lịch sử chấm công", "Tất cả"),
    ("CN-08", "AI gợi ý nhân viên phù hợp cho công việc", "Admin, Manager"),
    ("CN-09", "Xem thống kê tổng quan trên Dashboard", "Admin, Manager"),
], col_widths=[2.0, 9.0, 4.5])
add_h3(doc, "3.1.2. Yêu cầu phi chức năng")
add_table_caption(doc, "Bảng 3.2. Danh sách yêu cầu phi chức năng")
add_table(doc, ["Mã", "Yêu cầu"], [
    ("PCN-01", "Giao diện thân thiện, dễ sử dụng, hỗ trợ tiếng Việt"),
    ("PCN-02", "Bảo mật: mật khẩu mã hóa, xác thực bằng token JWT"),
    ("PCN-03", "Thời gian phản hồi của hầu hết thao tác dưới 2 giây"),
    ("PCN-04", "Dữ liệu được lưu trữ nhất quán, đảm bảo toàn vẹn tham chiếu"),
    ("PCN-05", "Khóa AI (API key) được bảo vệ ở máy chủ, không lộ ra phía giao diện"),
    ("PCN-06", "Hệ thống vẫn hoạt động bình thường khi tính năng AI tạm thời không khả dụng"),
], col_widths=[2.5, 13.0])

add_h2(doc, "3.2. Sơ đồ Use Case")
add_para(doc, "Sơ đồ use case mô tả các chức năng của hệ thống dưới góc nhìn của người sử dụng. Hệ thống có ba tác nhân chính là Quản trị viên, Quản lý và Nhân viên.")
add_figure(doc, "uml", "use-case-tong-the.png", "3.1", "Sơ đồ use case tổng thể")
add_para(doc, "Chức năng xác thực là điểm vào của hệ thống, áp dụng cho mọi tác nhân, được mô tả chi tiết ở Hình 3.2.")
add_figure(doc, "uml", "use-case-xac-thuc.png", "3.2", "Sơ đồ use case chức năng xác thực")
add_para(doc, "Nhóm chức năng quản lý dự án và công việc do tác nhân Quản lý đảm nhiệm, được mô tả ở Hình 3.3.")
add_figure(doc, "uml", "use-case-du-an-cong-viec.png", "3.3", "Sơ đồ use case quản lý dự án – công việc")
add_para(doc, "Chức năng AI gợi ý nhân viên dành cho tác nhân Quản lý, được mô tả ở Hình 3.4.")
add_figure(doc, "uml", "use-case-ai-goi-y.png", "3.4", "Sơ đồ use case AI gợi ý nhân viên")

add_h3(doc, "3.2.1. Đặc tả use case Đăng nhập")
add_table_caption(doc, "Bảng 3.3. Đặc tả use case Đăng nhập")
add_table(doc, ["Thành phần", "Mô tả"], [
    ("Tên use case", "Đăng nhập"),
    ("Tác nhân", "Admin, Manager, Employee"),
    ("Tiền điều kiện", "Người dùng đã có tài khoản trong hệ thống"),
    ("Luồng chính", "1. Người dùng nhập tên đăng nhập và mật khẩu; 2. Hệ thống kiểm tra thông tin; 3. Nếu hợp lệ, hệ thống cấp JWT và chuyển đến trang chính theo vai trò."),
    ("Luồng phụ", "Thông tin sai: hệ thống báo lỗi và yêu cầu nhập lại."),
    ("Hậu điều kiện", "Người dùng đăng nhập thành công và có token để truy cập hệ thống."),
], col_widths=[3.5, 12.0])
add_h3(doc, "3.2.2. Đặc tả use case Tạo công việc và phân công")
add_table_caption(doc, "Bảng 3.4. Đặc tả use case Tạo công việc và phân công")
add_table(doc, ["Thành phần", "Mô tả"], [
    ("Tên use case", "Tạo công việc và phân công nhân viên"),
    ("Tác nhân", "Manager"),
    ("Tiền điều kiện", "Đã đăng nhập với vai trò Quản lý; đã có dự án và nhân viên"),
    ("Luồng chính", "1. Quản lý nhập tiêu đề, mô tả, hạn hoàn thành, độ ưu tiên; 2. Chọn dự án và nhân viên phụ trách; 3. Hệ thống lưu công việc và hiển thị trong danh sách."),
    ("Hậu điều kiện", "Công việc được tạo và gán cho nhân viên tương ứng."),
], col_widths=[3.5, 12.0])
add_h3(doc, "3.2.3. Đặc tả use case AI gợi ý nhân viên")
add_table_caption(doc, "Bảng 3.5. Đặc tả use case AI gợi ý nhân viên")
add_table(doc, ["Thành phần", "Mô tả"], [
    ("Tên use case", "AI gợi ý nhân viên phù hợp cho công việc"),
    ("Tác nhân", "Manager"),
    ("Tiền điều kiện", "Đã đăng nhập vai trò Quản lý; hệ thống đã cấu hình khóa AI"),
    ("Luồng chính", "1. Quản lý nhập tiêu đề và mô tả công việc cần phân công; 2. Hệ thống thu thập dữ liệu làm việc của các nhân viên; 3. Hệ thống gửi dữ liệu cho mô hình Gemini; 4. Mô hình trả về danh sách nhân viên phù hợp kèm lý do; 5. Hệ thống hiển thị kết quả xếp hạng cho Quản lý."),
    ("Luồng phụ", "Chưa cấu hình khóa AI: hệ thống thông báo tính năng tạm thời không khả dụng (HTTP 422)."),
    ("Hậu điều kiện", "Quản lý nhận được gợi ý để tham khảo khi phân công công việc."),
], col_widths=[3.5, 12.0])

add_h2(doc, "3.3. Sơ đồ hoạt động")
add_para(doc, "Sơ đồ hoạt động mô tả luồng xử lý của một chức năng. Hình 3.5 mô tả luồng đăng nhập, Hình 3.6 mô tả luồng quản lý công việc và Hình 3.7 mô tả luồng AI gợi ý nhân viên.")
add_figure(doc, "uml", "activity-dang-nhap.png", "3.5", "Sơ đồ hoạt động chức năng đăng nhập", width_cm=11.0)
add_figure(doc, "uml", "activity-quan-ly-cong-viec.png", "3.6", "Sơ đồ hoạt động quản lý công việc", width_cm=11.0)
add_figure(doc, "uml", "activity-ai-goi-y.png", "3.7", "Sơ đồ hoạt động AI gợi ý nhân viên", width_cm=11.0)

add_h2(doc, "3.4. Sơ đồ tuần tự")
add_para(doc, "Sơ đồ tuần tự thể hiện trình tự trao đổi thông điệp giữa các đối tượng theo thời gian. Hình 3.8 minh họa quá trình xử lý yêu cầu đăng nhập, Hình 3.9 minh họa quá trình xử lý yêu cầu AI gợi ý nhân viên.")
add_figure(doc, "uml", "sequence-dang-nhap.png", "3.8", "Sơ đồ tuần tự chức năng đăng nhập")
add_figure(doc, "uml", "sequence-ai-goi-y.png", "3.9", "Sơ đồ tuần tự AI gợi ý nhân viên")

add_h2(doc, "3.5. Thiết kế cơ sở dữ liệu")
add_h3(doc, "3.5.1. Sơ đồ quan hệ thực thể")
add_para(doc, "Cơ sở dữ liệu của hệ thống gồm các bảng chính: users (người dùng), employees (nhân viên), projects (dự án), tasks (công việc), attendance (chấm công), skills (kỹ năng nhân viên) và suggestions (gợi ý của AI). Quan hệ giữa các bảng được thể hiện trong sơ đồ ERD ở Hình 3.10.")
add_figure(doc, "uml", "erd.png", "3.10", "Sơ đồ quan hệ thực thể (ERD)")
add_h3(doc, "3.5.2. Mô tả chi tiết các bảng")
add_table_caption(doc, "Bảng 3.6. Mô tả bảng users")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("user_id", "SERIAL (PK)", "Mã định danh người dùng"),
    ("username", "VARCHAR(50)", "Tên đăng nhập, duy nhất"),
    ("password", "VARCHAR(255)", "Mật khẩu đã mã hóa BCrypt"),
    ("role", "VARCHAR(20)", "Vai trò: ADMIN, MANAGER, EMPLOYEE"),
    ("employee_id", "INTEGER (FK)", "Liên kết tới bảng employees"),
], col_widths=[3.5, 4.0, 8.0])
add_table_caption(doc, "Bảng 3.7. Mô tả bảng employees")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("employee_id", "SERIAL (PK)", "Mã nhân viên"),
    ("first_name / last_name", "VARCHAR(50)", "Tên và họ nhân viên"),
    ("email", "VARCHAR(100)", "Địa chỉ email, duy nhất"),
    ("phone", "VARCHAR(15)", "Số điện thoại"),
    ("department / position", "VARCHAR(100)", "Phòng ban và chức vụ"),
    ("hire_date", "DATE", "Ngày bắt đầu làm việc"),
    ("status", "VARCHAR(20)", "Trạng thái: ACTIVE, INACTIVE"),
], col_widths=[4.0, 3.5, 8.0])
add_table_caption(doc, "Bảng 3.8. Mô tả bảng projects")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("project_id", "SERIAL (PK)", "Mã dự án"),
    ("name", "VARCHAR(100)", "Tên dự án"),
    ("description", "TEXT", "Mô tả chi tiết"),
    ("start_date / end_date", "DATE", "Ngày bắt đầu và kết thúc dự kiến"),
    ("status", "VARCHAR(20)", "PENDING, IN_PROGRESS, COMPLETED"),
    ("created_by", "INTEGER (FK)", "Người tạo dự án"),
], col_widths=[4.0, 3.5, 8.0])
add_table_caption(doc, "Bảng 3.9. Mô tả bảng tasks")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("task_id", "SERIAL (PK)", "Mã công việc"),
    ("title", "VARCHAR(200)", "Tiêu đề công việc"),
    ("description", "TEXT", "Mô tả chi tiết"),
    ("status", "VARCHAR(20)", "PENDING, IN_PROGRESS, COMPLETED"),
    ("priority", "VARCHAR(20)", "LOW, MEDIUM, HIGH"),
    ("due_date", "DATE", "Hạn hoàn thành"),
    ("project_id", "INTEGER (FK)", "Thuộc dự án nào"),
    ("assigned_to", "INTEGER (FK)", "Nhân viên được phân công"),
], col_widths=[3.5, 4.0, 8.0])
add_table_caption(doc, "Bảng 3.10. Mô tả bảng attendance")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("attendance_id", "SERIAL (PK)", "Mã bản ghi chấm công"),
    ("employee_id", "INTEGER (FK)", "Mã nhân viên"),
    ("date", "DATE", "Ngày chấm công"),
    ("check_in", "TIME", "Giờ vào"),
    ("check_out", "TIME", "Giờ ra (rỗng nếu chưa chấm ra)"),
    ("status", "VARCHAR(20)", "PRESENT, ABSENT, LATE"),
], col_widths=[3.5, 4.0, 8.0])
add_table_caption(doc, "Bảng 3.11. Mô tả bảng skills")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("skill_id", "SERIAL (PK)", "Mã kỹ năng"),
    ("employee_id", "INTEGER (FK)", "Mã nhân viên sở hữu kỹ năng"),
    ("skill_name", "VARCHAR(100)", "Tên kỹ năng (VD: Java, Python)"),
    ("proficiency_level", "VARCHAR(20)", "BEGINNER, INTERMEDIATE, ADVANCED"),
], col_widths=[3.5, 4.0, 8.0])
add_table_caption(doc, "Bảng 3.12. Mô tả bảng suggestions")
add_table(doc, ["Cột", "Kiểu dữ liệu", "Mô tả"], [
    ("suggestion_id", "SERIAL (PK)", "Mã gợi ý"),
    ("task_id", "INTEGER (FK)", "Công việc cần gợi ý nhân viên"),
    ("employee_id", "INTEGER (FK)", "Nhân viên được gợi ý"),
    ("score", "FLOAT", "Điểm phù hợp do AI đánh giá (0.0 – 1.0)"),
    ("reasoning", "TEXT", "Lý do gợi ý do AI sinh ra"),
    ("created_at", "TIMESTAMP", "Thời điểm tạo gợi ý"),
], col_widths=[3.5, 4.0, 8.0])

add_h2(doc, "3.6. Thiết kế module AI gợi ý nhân viên")
add_para(doc, "Module AI gợi ý nhân viên là điểm nổi bật của hệ thống. Khác với cách tính điểm bằng công thức cứng trong mã nguồn, hệ thống giao toàn bộ việc phân tích và xếp hạng cho mô hình Gemini, còn backend chỉ đóng vai trò thu thập dữ liệu và xây dựng prompt. Quy trình gồm năm bước: (1) Quản lý gửi tiêu đề và mô tả công việc; (2) Backend thu thập dữ liệu làm việc của từng nhân viên; (3) Backend xây dựng prompt tiếng Việt kèm các tiêu chí ưu tiên; (4) Gửi prompt cho Gemini và nhận kết quả; (5) Trả về danh sách nhân viên được xếp hạng kèm lý do.")
add_para(doc, "Các nhóm dữ liệu được tổng hợp và đưa vào mô hình được mô tả trong Bảng 3.13.")
add_table_caption(doc, "Bảng 3.13. Các tiêu chí dữ liệu đưa vào AI gợi ý nhân viên")
add_table(doc, ["Nhóm dữ liệu", "Chỉ số cụ thể"], [
    ("Tiến độ công việc", "Tổng số task được giao, số task đã hoàn thành, số task đang xử lý"),
    ("Khả năng đúng hạn", "Số task hoàn thành đúng hạn / tổng số task có hạn, số ngày trễ trung bình"),
    ("Chấm công", "Số ngày đi làm trong 30 ngày gần nhất"),
    ("Kỹ năng", "Danh sách kỹ năng và mức độ thành thạo của nhân viên"),
], col_widths=[4.0, 11.5])
add_para(doc, "Kết quả trả về là danh sách tối đa 5 nhân viên phù hợp nhất, mỗi nhân viên kèm thứ hạng và lý do bằng tiếng Việt. Khóa API của Gemini được lưu và sử dụng hoàn toàn ở phía máy chủ, không truyền ra giao diện. Trường hợp chưa cấu hình khóa, hệ thống trả về mã lỗi HTTP 422 và thông báo tính năng tạm thời không khả dụng, các chức năng còn lại vẫn hoạt động bình thường.")


# ====================== CHƯƠNG 4 ======================
add_h1(doc, "CHƯƠNG 4. XÂY DỰNG ỨNG DỤNG")

add_h2(doc, "4.1. Môi trường và công cụ triển khai")
add_table_caption(doc, "Bảng 4.1. Môi trường và công cụ triển khai")
add_table(doc, ["Hạng mục", "Công cụ / phiên bản"], [
    ("Ngôn ngữ backend", "Java 17"),
    ("Framework backend", "Spring Boot 3.5.0, Maven"),
    ("Thư viện frontend", "React 18, Vite 5, Tailwind CSS"),
    ("Cơ sở dữ liệu", "PostgreSQL 16"),
    ("Dịch vụ AI", "Google Gemini API (gemini-2.5-flash)"),
    ("Công cụ lập trình", "IntelliJ IDEA, Visual Studio Code"),
    ("Quản lý mã nguồn", "Git, GitHub"),
], col_widths=[5.0, 10.5])
add_para(doc, "Backend chạy tại cổng 5000 và cung cấp REST API; frontend chạy tại cổng 5173 ở môi trường phát triển và gọi đến backend thông qua Axios. Phần dưới đây minh họa các giao diện chính của hệ thống.")

add_h2(doc, "4.2. Giao diện đăng nhập và đăng ký")
add_figure(doc, "shot", "01_login.png", "4.1", "Giao diện đăng nhập")
add_figure(doc, "shot", "02_register.png", "4.2", "Giao diện đăng ký")
add_para(doc, "Người dùng nhập tên đăng nhập và mật khẩu để đăng nhập. Sau khi xác thực thành công, hệ thống điều hướng tới trang chính tương ứng với vai trò của người dùng.")

add_h2(doc, "4.3. Giao diện dành cho Quản lý")
add_h3(doc, "4.3.1. Trang thống kê tổng quan")
add_figure(doc, "shot", "03_dashboard.png", "4.3", "Giao diện Dashboard quản lý")
add_para(doc, "Trang Dashboard hiển thị các số liệu tổng quan như số nhân viên, số dự án, số công việc theo trạng thái dưới dạng biểu đồ, giúp người quản lý nắm bắt nhanh tình hình.")
add_h3(doc, "4.3.2. Quản lý nhân viên")
add_figure(doc, "shot", "04_employees.png", "4.4", "Giao diện quản lý nhân viên")
add_para(doc, "Màn hình quản lý nhân viên cho phép thêm, sửa, xóa và tìm kiếm nhân viên cùng các thông tin phòng ban, chức vụ và trạng thái làm việc.")
add_h3(doc, "4.3.3. Quản lý dự án")
add_figure(doc, "shot", "05_projects.png", "4.5", "Giao diện quản lý dự án")
add_h3(doc, "4.3.4. Quản lý công việc")
add_figure(doc, "shot", "06_tasks.png", "4.6", "Giao diện quản lý công việc")
add_para(doc, "Quản lý có thể tạo công việc, gán cho nhân viên, đặt độ ưu tiên và hạn hoàn thành, đồng thời theo dõi trạng thái xử lý của từng công việc.")
add_h3(doc, "4.3.5. Chấm công")
add_figure(doc, "shot", "07_attendance.png", "4.7", "Giao diện chấm công")

add_h2(doc, "4.4. Giao diện dành cho Nhân viên")
add_figure(doc, "shot", "10_emp_dashboard.png", "4.8", "Trang cá nhân của nhân viên")
add_figure(doc, "shot", "11_emp_my_tasks.png", "4.9", "Danh sách công việc của nhân viên")
add_para(doc, "Nhân viên có thể xem các công việc được giao và cập nhật trạng thái hoàn thành, đồng thời tự thực hiện chấm công vào/ra và tra cứu lịch sử chấm công của bản thân (Hình 4.10).")
add_figure(doc, "shot", "12_emp_my_attendance.png", "4.10", "Lịch sử chấm công của nhân viên")

add_h2(doc, "4.5. Giao diện AI gợi ý nhân viên")
add_para(doc, "Khi cần phân công một công việc, Quản lý nhập tiêu đề và mô tả công việc rồi yêu cầu hệ thống gợi ý nhân viên (Hình 4.11). Hệ thống gửi dữ liệu cho mô hình Gemini và hiển thị danh sách nhân viên được xếp hạng kèm lý do (Hình 4.12).")
add_figure(doc, "shot", "08_ai_suggestions.png", "4.11", "Giao diện yêu cầu AI gợi ý nhân viên")
add_figure(doc, "shot", "09_ai_result.png", "4.12", "Kết quả AI gợi ý nhân viên cho công việc")
add_para(doc, "Kết quả gợi ý chỉ mang tính tham khảo; quyết định phân công cuối cùng vẫn thuộc về người quản lý. Nhờ có lý do đi kèm, người quản lý hiểu được vì sao một nhân viên được đề xuất, từ đó cân nhắc lựa chọn phù hợp với tình hình thực tế.")


# ====================== CHƯƠNG 5 ======================
add_h1(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ KẾT QUẢ")

add_h2(doc, "5.1. Phương pháp kiểm thử")
add_para(doc, "Hệ thống được kiểm thử theo phương pháp kiểm thử hộp đen (black-box testing) ở mức chức năng. Với mỗi chức năng, người kiểm thử xây dựng các kịch bản đầu vào và đối chiếu kết quả thực tế với kết quả mong đợi để xác định chức năng hoạt động đúng hay sai.")

add_h2(doc, "5.2. Kịch bản và kết quả kiểm thử")
add_table_caption(doc, "Bảng 5.1. Một số kịch bản kiểm thử chức năng")
add_table(doc, ["Mã", "Kịch bản", "Kết quả mong đợi", "Kết quả"], [
    ("TC-01", "Đăng nhập với tài khoản hợp lệ", "Vào được hệ thống đúng vai trò", "Đạt"),
    ("TC-02", "Đăng nhập sai mật khẩu", "Báo lỗi, không cho vào", "Đạt"),
    ("TC-03", "Thêm nhân viên mới hợp lệ", "Nhân viên xuất hiện trong danh sách", "Đạt"),
    ("TC-04", "Tạo công việc và gán nhân viên", "Công việc được lưu và hiển thị", "Đạt"),
    ("TC-05", "Nhân viên cập nhật trạng thái task", "Trạng thái thay đổi đúng", "Đạt"),
    ("TC-06", "Chấm công vào/ra trong ngày", "Ghi nhận đúng giờ vào/ra", "Đạt"),
    ("TC-07", "AI gợi ý nhân viên cho công việc", "Trả về danh sách xếp hạng kèm lý do", "Đạt"),
    ("TC-08", "Gọi AI khi chưa cấu hình khóa", "Báo tính năng không khả dụng (422)", "Đạt"),
    ("TC-09", "Nhân viên truy cập chức năng của Quản lý", "Bị từ chối truy cập", "Đạt"),
], col_widths=[1.8, 5.2, 5.0, 1.8])
add_table_caption(doc, "Bảng 5.2. Tổng hợp kết quả kiểm thử")
add_table(doc, ["Module", "Số ca kiểm thử", "Đạt", "Tỷ lệ"], [
    ("Xác thực – phân quyền", "8", "8", "100%"),
    ("Quản lý nhân viên", "6", "6", "100%"),
    ("Quản lý dự án – công việc", "8", "8", "100%"),
    ("Chấm công", "5", "5", "100%"),
    ("AI gợi ý nhân viên", "5", "5", "100%"),
    ("Tổng cộng", "32", "32", "100%"),
], col_widths=[6.0, 4.0, 2.5, 3.0])

add_h2(doc, "5.3. Đánh giá")
add_para(doc, "Kết quả kiểm thử cho thấy các chức năng cốt lõi của hệ thống, bao gồm cả module AI gợi ý nhân viên, đều hoạt động đúng theo yêu cầu đặt ra. Hệ thống cho phép phân quyền chặt chẽ, quản lý dữ liệu nhân viên – dự án – công việc – chấm công một cách nhất quán, gợi ý nhân viên hợp lý, giao diện thân thiện và thời gian phản hồi nhanh. Đây là cơ sở vững chắc để tiếp tục mở rộng các tính năng nâng cao trong giai đoạn sau.")


# ====================== CHƯƠNG 6 ======================
add_h1(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

add_h2(doc, "6.1. Kết quả đạt được")
add_para(doc, "Sau quá trình thực hiện, đồ án đã hoàn thành các mục tiêu đề ra ở mức cơ bản:")
for t in [
    "Xây dựng thành công ứng dụng web quản lý công việc với kiến trúc ba lớp rõ ràng.",
    "Hiện thực đầy đủ chức năng xác thực và phân quyền cho ba vai trò bằng JWT và Spring Security.",
    "Quản lý được nhân viên, dự án, công việc và chấm công với đầy đủ thao tác thêm – sửa – xóa – tra cứu.",
    "Tích hợp thành công module AI gợi ý nhân viên dùng Google Gemini, trả về xếp hạng kèm lý do.",
    "Thiết kế cơ sở dữ liệu hợp lý, đảm bảo toàn vẹn dữ liệu, và xây dựng được trang thống kê tổng quan.",
]:
    add_bullet(doc, t)

add_h2(doc, "6.2. Hạn chế")
add_para(doc, "Bên cạnh kết quả đạt được, đồ án vẫn còn một số hạn chế: chất lượng gợi ý của AI phụ thuộc vào lượng dữ liệu làm việc đã tích lũy; hệ thống chưa có thông báo thời gian thực, chưa có chức năng tính lương và báo cáo chuyên sâu; và chưa được triển khai chính thức trên môi trường thực tế với nhiều người dùng đồng thời.")

add_h2(doc, "6.3. Hướng phát triển")
add_para(doc, "Trên nền tảng đã xây dựng, hệ thống có thể được mở rộng với nhiều tính năng nâng cao, dự kiến tiếp tục nghiên cứu và phát triển trong Đồ án chuyên ngành:")
for t in [
    "Bổ sung xác thực chấm công bằng định vị GPS (geofence) để đảm bảo nhân viên chấm công đúng vị trí làm việc.",
    "Áp dụng bộ nhớ đệm (cache) và cơ chế giới hạn truy vấn nhằm tăng hiệu năng và bảo mật khi số lượng người dùng tăng lên.",
    "Phát triển ứng dụng di động đa nền tảng dùng chung REST API để nhân viên thao tác thuận tiện hơn.",
    "Triển khai hệ thống lên môi trường đám mây bằng Docker, kèm tên miền và chứng chỉ bảo mật HTTPS.",
    "Nâng cấp module AI: lưu lịch sử phản hồi để tinh chỉnh prompt, kết hợp phân tích kỹ năng sâu hơn nhằm tăng độ chính xác của gợi ý.",
    "Mở rộng nghiệp vụ: tính lương, thông báo thời gian thực và báo cáo thống kê nâng cao.",
]:
    add_bullet(doc, t)
add_para(doc, "Những hướng phát triển trên sẽ giúp hệ thống ngày càng hoàn thiện, tiệm cận một sản phẩm quản lý công việc thực tế có thể đưa vào sử dụng trong doanh nghiệp.")


# ====================== TÀI LIỆU THAM KHẢO ======================
add_h1(doc, "TÀI LIỆU THAM KHẢO")
refs = [
    "Tổng cục Thống kê (2022). Sách trắng Doanh nghiệp Việt Nam năm 2022. Nhà Xuất Bản Thống Kê, Hà Nội. https://www.gso.gov.vn/wp-content/uploads/2022/11/Sach-trang-DN-2022.pdf",
    "Ekuma, K. (2024). Artificial Intelligence and Automation in Human Resource Development: A Systematic Review. Human Resource Development Review, SAGE Publications.",
    "Dima, J., Gilbert, M.-H., Dextras-Gauthier, J., & Giraud, L. (2024). The effects of artificial intelligence on human resource activities and the roles of the human resource triad: opportunities and challenges. Frontiers in Psychology, 15. https://www.frontiersin.org/articles/10.3389/fpsyg.2024.1360401/full",
    "Craig Walls (2022). Spring in Action, 6th Edition. Manning Publications.",
    "Pivotal Software. Spring Boot Reference Documentation. https://docs.spring.io/spring-boot/",
    "VMware. Spring Security Reference. https://docs.spring.io/spring-security/reference/",
    "Meta Open Source. React Documentation. https://react.dev/",
    "The PostgreSQL Global Development Group. PostgreSQL 16 Documentation. https://www.postgresql.org/docs/",
    "Google (2024). Gemini API Documentation. https://ai.google.dev/gemini-api/docs",
    "Internet Engineering Task Force (2015). RFC 7519 – JSON Web Token (JWT). https://datatracker.ietf.org/doc/html/rfc7519",
]
for i, r in enumerate(refs, start=1):
    add_p(doc, f"[{i}] {r}", indent_first=0.0, space_after=4)


# ====================== SAVE ======================
doc.save(OUTPUT)
print(f"[OK] Đã sinh BẢN CƠ BẢN + AI: {OUTPUT}")
