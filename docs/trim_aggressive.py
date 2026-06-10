"""
Trim mạnh tay hơn:
- Chương 4: Bỏ hoàn toàn Heading 3 + content trong 4.3, 4.4, 4.5, 4.6, 4.7
  Giữ Heading 2 + 1 đoạn tóm tắt
- Chương 2: Mỗi Heading 2 chỉ giữ 1 đoạn tóm tắt, bỏ Heading 3 con
- Cắt mọi line là XML/JSON tag, brackets, key:value lẻ
"""
import re
import docx

PATH = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'
d = docx.Document(PATH)

def hl(p):
    n = p.style.name
    if n.startswith('Heading '):
        try: return int(n.split(' ')[1])
        except: return 0
    return 0

def chap(text):
    text = text.upper()
    m = re.match(r'CHƯƠNG (\d)', text)
    if m: return int(m.group(1))
    if 'KẾT LUẬN' in text: return 6
    return None

def is_code_fragment(t):
    t = t.strip()
    if not t: return False
    # XML/HTML tags
    if re.match(r'^<[/!?]?[a-zA-Z]', t): return True
    if re.match(r'^</', t): return True
    # Brackets / braces lone
    if re.match(r'^[\{\}\[\]()<>;,]+\s*$', t): return True
    # Indented strings/values
    if re.match(r'^\s*"[^"]+"\s*[:,]?\s*$', t): return True
    # JSON key-value
    if re.match(r'^"[^"]+"\s*:\s*', t): return True
    # YAML/properties
    if re.match(r'^[a-zA-Z_]\w*\s*[:=]\s', t): return True
    # All caps constants / shell vars
    if re.match(r'^[A-Z_][A-Z0-9_]*\s*=', t): return True
    # Java imports / annotations
    if re.match(r'^(import|package|@\w+|public|private|protected|static|final|void|return|if\b|else|throw|new\b|class|interface|extends|implements|this\.|super\.)', t): return True
    # Common bracket close
    if t in ('});',  '};', '});', '))', '} else {', '} catch (Exception e) {', '}', '})', '],', '},', '});'): return True
    # Dependencies XML inside
    if '<groupId>' in t or '<artifactId>' in t or '<version>' in t or '<scope>' in t: return True
    # Stage / Dockerfile / npm
    if re.match(r'^(FROM|RUN|COPY|CMD|ENV|EXPOSE|WORKDIR|ADD|ENTRYPOINT|ARG|VOLUME|LABEL)\s', t): return True
    return False

# ============================================================
# Pass 1: cắt mọi code fragment trong Ch4
# ============================================================
in_ch4 = False
removed = 0
for p in list(d.paragraphs):
    t = p.text.strip()
    if hl(p) == 1:
        c = chap(t)
        in_ch4 = (c == 4)
        continue
    if not in_ch4: continue
    if hl(p): continue
    if not t: continue
    if p.style.name.startswith('List'): continue
    if is_code_fragment(t):
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el)
            removed += 1
print(f'Pass 1 (Ch4 code fragments): {removed} removed')

# ============================================================
# Pass 2: Bỏ hết Heading 3 + nội dung trong 4.3 -> 4.7
# (giữ Heading 2 + 1 đoạn tóm tắt phía dưới)
# ============================================================
# We'll iterate, when we hit Heading 2 in Ch4 = "4.X. ...", mark as such.
# Then on next Heading 3 inside it, remove that Heading 3 and everything until next Heading 2 or H1.
in_ch4 = False
inside_h2_to_trim = False
removed_h3 = 0

paragraphs = list(d.paragraphs)
i = 0
while i < len(paragraphs):
    p = paragraphs[i]
    t = p.text.strip()
    lvl = hl(p)
    if lvl == 1:
        c = chap(t)
        in_ch4 = (c == 4)
        inside_h2_to_trim = False
        i += 1
        continue
    if in_ch4 and lvl == 2:
        # Trim subsections in 4.3-4.7
        m = re.match(r'4\.([3-7])\.', t)
        inside_h2_to_trim = bool(m)
        i += 1
        continue
    if inside_h2_to_trim and lvl == 3:
        # Remove this H3 and everything until next H2 / H1
        # First, keep just the H3 if it's important — actually drop entirely
        j = i
        while j < len(paragraphs):
            np = paragraphs[j]
            nlvl = hl(np)
            if j > i and nlvl in (1, 2): break
            el = np._element
            if el.getparent() is not None:
                el.getparent().remove(el)
                removed_h3 += 1
            j += 1
        i = j
        continue
    i += 1

print(f'Pass 2 (Ch4 H3 subsections in 4.3-4.7): {removed_h3} paragraphs removed')

# ============================================================
# Pass 3: Chương 2 — mỗi Heading 2 chỉ giữ 1 đoạn tóm tắt, bỏ Heading 3
# ============================================================
in_ch2 = False
inside_h2 = False
removed_ch2 = 0
para_count = 0

paragraphs = list(d.paragraphs)
i = 0
while i < len(paragraphs):
    p = paragraphs[i]
    t = p.text.strip()
    lvl = hl(p)
    if lvl == 1:
        c = chap(t)
        in_ch2 = (c == 2)
        inside_h2 = False
        para_count = 0
        i += 1; continue
    if not in_ch2:
        i += 1; continue
    if lvl == 2:
        inside_h2 = True
        para_count = 0
        i += 1; continue
    if inside_h2 and lvl == 3:
        # Remove H3 + content until next H2/H3
        j = i
        while j < len(paragraphs):
            np = paragraphs[j]
            nlvl = hl(np)
            if j > i and nlvl in (1, 2): break
            if j > i and nlvl == 3: break  # next H3 keeps the loop boundary
            el = np._element
            if el.getparent() is not None:
                el.getparent().remove(el)
                removed_ch2 += 1
            j += 1
        i = j; continue
    if inside_h2 and t and not p.style.name.startswith('List'):
        para_count += 1
        if para_count > 1:
            el = p._element
            if el.getparent() is not None:
                el.getparent().remove(el)
                removed_ch2 += 1
    i += 1

print(f'Pass 3 (Ch2 H3 + extra paras): {removed_ch2} removed')

d.save(PATH)
print(f'\nSaved → {PATH}')
