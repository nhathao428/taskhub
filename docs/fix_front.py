"""
Fix front matter trống:
1. Thêm bảng từ viết tắt vào DANH MỤC CÁC TỪ VIẾT TẮT
2. Bỏ các trang trống giữa Mục lục / Danh mục
3. Bỏ paragraph trống thừa giữa các section
"""
import docx
from docx.shared import Pt, Cm
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

PATH = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'
d = docx.Document(PATH)

# ============================================================
# 1. Thêm content vào DANH MỤC CÁC TỪ VIẾT TẮT
# ============================================================
ABBREVIATIONS = [
    ('Viết tắt', 'Diễn giải'),
    ('API', 'Application Programming Interface'),
    ('JWT', 'JSON Web Token'),
    ('JPA', 'Java Persistence API'),
    ('ORM', 'Object Relational Mapping'),
    ('REST', 'Representational State Transfer'),
    ('CRUD', 'Create – Read – Update – Delete'),
    ('SPA', 'Single Page Application'),
    ('MVC', 'Model – View – Controller'),
    ('CSDL', 'Cơ sở dữ liệu'),
    ('AI', 'Artificial Intelligence – Trí tuệ nhân tạo'),
    ('LLM', 'Large Language Model'),
    ('UI', 'User Interface'),
    ('UX', 'User Experience'),
    ('GPS', 'Global Positioning System'),
    ('PWA', 'Progressive Web Application'),
    ('VPS', 'Virtual Private Server'),
]

# Find heading "DANH MỤC CÁC TỪ VIẾT TẮT"
target_p = None
for p in d.paragraphs:
    if p.style.name == 'Heading 1' and 'TỪ VIẾT TẮT' in p.text.upper():
        target_p = p
        break

if target_p is not None:
    # Add table immediately after the heading
    anchor = target_p._element
    # Create a 2-column table
    tbl_xml = '<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    tbl_xml += '<w:tblPr><w:tblW w:w="5000" w:type="pct"/><w:tblBorders>'
    for b in ['top','bottom','left','right','insideH','insideV']:
        tbl_xml += f'<w:{b} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    tbl_xml += '</w:tblBorders><w:tblLayout w:type="autofit"/></w:tblPr>'
    tbl_xml += '<w:tblGrid><w:gridCol w:w="2200"/><w:gridCol w:w="6800"/></w:tblGrid>'
    for i, (abbr, mean) in enumerate(ABBREVIATIONS):
        tbl_xml += '<w:tr>'
        for col_text in (abbr, mean):
            bold = ' w:val="bold"' if i == 0 else ''
            tbl_xml += f'<w:tc><w:tcPr><w:tcW w:w="0" w:type="auto"/></w:tcPr>'
            tbl_xml += f'<w:p><w:r>'
            if i == 0:
                tbl_xml += '<w:rPr><w:b/></w:rPr>'
            tbl_xml += f'<w:t xml:space="preserve">{col_text}</w:t></w:r></w:p></w:tc>'
        tbl_xml += '</w:tr>'
    tbl_xml += '</w:tbl>'
    from lxml import etree
    tbl_el = etree.fromstring(tbl_xml)
    anchor.addnext(tbl_el)
    print(f'Inserted abbreviation table with {len(ABBREVIATIONS)-1} entries')

# ============================================================
# 2. Bỏ empty paragraphs liên tiếp (trong front matter + toàn bộ)
# ============================================================
removed_empty = 0
prev_empty = False
for p in list(d.paragraphs):
    has_image = bool(p._element.findall(f'.//{qn("w:drawing")}'))
    has_break = bool(p._element.findall(f'.//{qn("w:br")}[@{qn("w:type")}="page"]'))
    is_empty = (not p.text.strip()) and (not has_image) and (not has_break)
    if is_empty and prev_empty:
        el = p._element
        if el.getparent() is not None:
            el.getparent().remove(el); removed_empty += 1
    prev_empty = is_empty
print(f'Removed {removed_empty} consecutive empty paragraphs')

# ============================================================
# 3. Bỏ "page break before" trên các Heading 1 nhỏ ở front matter
# (để DANH MỤC CÁC TỪ VIẾT TẮT / DANH MỤC CÁC BẢNG... không bắt đầu trang mới riêng)
# ============================================================
# Look at Heading 1 paragraphs in early sections and remove pageBreakBefore
front_h1_to_inline = ('LỜI CAM ĐOAN', 'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC CÁC BẢNG', 'DANH MỤC CÁC HÌNH')
inlined = 0
for p in d.paragraphs:
    if p.style.name == 'Heading 1' and any(k in p.text.upper() for k in front_h1_to_inline):
        # Remove pageBreakBefore directly from paragraph
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            pbb = pPr.find(qn('w:pageBreakBefore'))
            if pbb is not None:
                pPr.remove(pbb)
                inlined += 1
        # Also remove preceding page break
        prev = p._element.getprevious()
        if prev is not None and prev.tag.endswith('}p'):
            prev_p = Paragraph(prev, d)
            brs = prev_p._element.findall(f'.//{qn("w:br")}[@{qn("w:type")}="page"]')
            for b in brs:
                b.getparent().remove(b)

print(f'Inlined {inlined} front-matter headings (removed pageBreakBefore)')

d.save(PATH)
print('Saved')
