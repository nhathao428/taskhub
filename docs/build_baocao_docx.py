"""
Sinh file Word báo cáo ĐỒ ÁN CƠ SỞ theo mẫu Khoa Công nghệ Thông tin.
Format: Times New Roman 13, line 1.5, lề T2 B2 L3 R2 cm.
Page numbering: La Mã từ Mục lục → Danh mục; Ả Rập từ Chương 1 → hết.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

import os as _os
OUTPUT = r"C:\Users\Admin\task-management-system\docs\BAO_CAO_DO_AN_CO_SO.docx"
SHOTS_DIR = r"C:\Users\Admin\task-management-system\docs\screenshots"
UML_DIR = r"C:\Users\Admin\task-management-system\docs\uml\png"

FONT = "Times New Roman"
SIZE_BODY = 13
SIZE_H1 = 14
SIZE_H_COVER = 16

# ------------------------------------------------------------------
# HELPERS
# ------------------------------------------------------------------
def set_run(run, size=SIZE_BODY, bold=False, italic=False, upper=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_p(doc, text="", size=SIZE_BODY, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=0.0,
          space_before=0, space_after=0, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def _style_set_font(style, size, bold=False, italic=False, color=None, upper=False):
    """Thiết lập font Times New Roman cho built-in style để TOC nhận diện được."""
    f = style.font
    f.name = FONT
    f.size = Pt(size)
    f.bold = bold
    f.italic = italic
    if color:
        f.color.rgb = RGBColor(*color)
    # Đảm bảo eastAsia/cs cũng dùng Times New Roman
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    if upper:
        caps = OxmlElement("w:caps")
        caps.set(qn("w:val"), "1")
        rPr.append(caps)


def configure_heading_styles(doc):
    """Cấu hình built-in Heading 1/2/3 theo mẫu Khoa CNTT để TOC tự lấy được."""
    # Heading 1: CHƯƠNG – in hoa, đậm, size 14, căn giữa
    h1 = doc.styles["Heading 1"]
    _style_set_font(h1, size=SIZE_H1, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True   # mỗi chương sang trang mới

    # Heading 2: 1.1 – đậm, đứng, size 13, căn trái
    h2 = doc.styles["Heading 2"]
    _style_set_font(h2, size=SIZE_BODY, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 1.5
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3: 1.1.1 – nghiêng, đậm, size 13
    h3 = doc.styles["Heading 3"]
    _style_set_font(h3, size=SIZE_BODY, bold=True, italic=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = 1.5
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Heading 4: 1.1.1.1 – nghiêng, size 13
    h4 = doc.styles["Heading 4"]
    _style_set_font(h4, size=SIZE_BODY, italic=True)
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.line_spacing = 1.5
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(2)
    h4.paragraph_format.keep_with_next = True


def add_h1(doc, text, numbered=True):
    """Tiêu đề chương (Heading 1)."""
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text.upper())
    return p


def add_h2(doc, text):
    """Tiêu đề mục 1.1 (Heading 2)."""
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_h3(doc, text):
    """Tiêu đề nhóm tiểu mục 1.1.1 (Heading 3)."""
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(text)
    return p


def add_h4(doc, text):
    """Tiểu mục 1.1.1.1 (Heading 4)."""
    p = doc.add_paragraph(style="Heading 4")
    p.add_run(text)
    return p


def add_toc_field(doc, levels="1-3", title=None):
    """Chèn TOC field thật – Word sẽ tự cập nhật khi mở/F9."""
    if title:
        p_title = doc.add_paragraph(style="Heading 1")
        p_title.add_run(title)
    p = doc.add_paragraph()
    run = p.add_run()
    # <w:fldChar w:fldCharType="begin" w:dirty="true"/>
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    fldChar_begin.set(qn("w:dirty"), "true")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f' TOC \\o "{levels}" \\h \\z \\u '
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("Hãy nhấn F9 (hoặc chuột phải → Update Field) "
                        "trong Word để cập nhật mục lục.")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    r_el = run._element
    r_el.append(fldChar_begin)
    r_el.append(instrText)
    r_el.append(fldChar_sep)
    r_el.append(placeholder)
    r_el.append(fldChar_end)
    # Báo Word phải tự cập nhật field khi mở file
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    return p


def add_para(doc, text, indent=0.75):
    """Đoạn văn body, justify, thụt dòng đầu 0.75cm."""
    return add_p(doc, text, indent_first=indent)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0] if p.runs else p.add_run("")
    # clear and replace
    p.text = ""
    run = p.add_run(text)
    set_run(run, size=SIZE_BODY)
    return p


def add_image(doc, filename, width_cm=15.5):
    """Chèn ảnh screenshot căn giữa, scale theo chiều rộng cố định.
    filename có thể là tên file (đặt trong SHOTS_DIR) hoặc đường dẫn con
    (ví dụ "mobile/m01_login.png")."""
    path = _os.path.join(SHOTS_DIR, filename)
    if not _os.path.exists(path):
        add_p(doc, f"[Không tìm thấy ảnh: {filename}]",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p


def add_uml_image(doc, filename, width_cm=15.5):
    """Chèn sơ đồ UML (PNG do PlantUML sinh) từ docs/uml/png/."""
    path = _os.path.join(UML_DIR, filename)
    if not _os.path.exists(path):
        add_p(doc, f"[Không tìm thấy sơ đồ UML: {filename}]",
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    return p


def add_mobile_pair(doc, left_file, left_cap, right_file, right_cap,
                    width_cm=6.2):
    """Chèn 2 ảnh chụp màn hình mobile cạnh nhau trong bảng 1 hàng × 2 cột.
    Bên dưới ảnh có caption nhỏ (italic, bold, căn giữa)."""
    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col in range(2):
        tbl.columns[col].width = Cm(8.0)

    for col, (fname, cap) in enumerate([(left_file, left_cap),
                                        (right_file, right_cap)]):
        if fname is None:
            continue
        cell = tbl.rows[0].cells[col]
        cell.width = Cm(8.0)
        # Xóa paragraph mặc định trong cell
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = _os.path.join(SHOTS_DIR, fname)
        if _os.path.exists(path):
            p.add_run().add_picture(path, width=Cm(width_cm))
        else:
            r = p.add_run(f"[Không tìm thấy ảnh: {fname}]")
            set_run(r, italic=True)

        cap_cell = tbl.rows[1].cells[col]
        cap_cell.text = ""
        cp = cap_cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(cap)
        set_run(cr, size=SIZE_BODY, italic=True, bold=True)
    add_p(doc, "", space_after=4)
    return tbl


def add_mobile_image(doc, filename, caption, width_cm=7.0):
    """Chèn 1 ảnh mobile đơn lẻ + caption ngay bên dưới."""
    add_image(doc, filename, width_cm=width_cm)
    add_caption(doc, caption, kind="figure")


def add_caption(doc, text, kind="figure"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run(run, size=SIZE_BODY, italic=True, bold=True)
    return p


def add_code(doc, code):
    """Khối code: font Consolas, size 10, line 1.0."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:cs"), "Consolas")
    add_p(doc, "", space_after=4)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # header
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=SIZE_BODY, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run(run, size=SIZE_BODY)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    # spacing after
    add_p(doc, "", space_after=4)
    return tbl


def add_page_break(doc):
    doc.add_page_break()


def set_page_number(section, fmt="decimal", start=None):
    """fmt: 'decimal' | 'upperRoman' | 'lowerRoman'"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def add_footer_page_number(section, show=True):
    """Đặt số trang ở footer căn giữa."""
    footer = section.footer
    footer.is_linked_to_previous = False
    # clear existing
    for p in list(footer.paragraphs):
        p.clear()
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if show:
        run = p.add_run()
        set_run(run, size=SIZE_BODY)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)


def new_section(doc, page_num_fmt="decimal", start=None, show_pn=True):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_number(section, fmt=page_num_fmt, start=start)
    add_footer_page_number(section, show=show_pn)
    return section


# ------------------------------------------------------------------
# DOC INIT
# ------------------------------------------------------------------
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(SIZE_BODY)
style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

# Cấu hình Heading 1/2/3/4 để TOC field tự nhận diện
configure_heading_styles(doc)

# First section: cover + forms (no page number)
sec1 = doc.sections[0]
sec1.top_margin = Cm(2)
sec1.bottom_margin = Cm(2)
sec1.left_margin = Cm(3)
sec1.right_margin = Cm(2)
add_footer_page_number(sec1, show=False)


# ==================================================================
# 1) TRANG BÌA CHÍNH
# ==================================================================
def add_cover_page(doc, is_main=True):
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(6)
    logo_p.add_run().add_picture("hutech_logo.png", width=Cm(2.8))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO")
    set_run(run, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP. HỒ CHÍ MINH")
    set_run(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run(run, size=13, bold=True)

    add_p(doc, "─" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("ĐỒ ÁN MÔN HỌC:")
    set_run(run, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("ĐỒ ÁN CƠ SỞ")
    set_run(run, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TÊN ĐỀ TÀI:")
    set_run(run, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("HỆ THỐNG QUẢN LÝ CÔNG VIỆC")
    set_run(run, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run("CHO DOANH NGHIỆP NHỎ ĐA NGÀNH TÍCH HỢP AI")
    set_run(run, size=18, bold=True)

    # Khối thông tin
    info_lines = [
        ("Ngành:", "CÔNG NGHỆ THÔNG TIN"),
        ("Lớp:", "23DTHC1"),
        ("Giảng viên hướng dẫn:", "ThS. DƯƠNG THÀNH PHẾT"),
        ("Sinh viên thực hiện:", "NGUYỄN NHẬT HẢO"),
        ("Mã số sinh viên:", "2380612688"),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(label.ljust(28))
        set_run(r1, size=14, bold=True)
        r2 = p.add_run(value)
        set_run(r2, size=14, bold=True)

    add_p(doc, "", space_after=18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TP. Hồ Chí Minh, tháng 5 năm 2026")
    set_run(run, size=14, bold=True, italic=True)


add_cover_page(doc, is_main=True)
doc.add_page_break()
add_cover_page(doc, is_main=False)

# Lưu ý: 4 phiếu hành chính (Giao đề tài / Theo dõi tiến độ / Chấm điểm GVHD /
# Chấm điểm GVPB) đã được tách ra thành các file .docx độc lập do
# docs/build_phieu_docx.py sinh ra. Báo cáo này chỉ còn nội dung chuyên môn.

print("[OK] Cover only (4 phiếu đã tách thành file riêng)")


# ==================================================================
# SECTION 2: FRONT MATTER (Lời cảm ơn → Danh mục) – đánh số La Mã i, ii, iii
# ==================================================================
sec2 = new_section(doc, page_num_fmt="lowerRoman", start=1, show_pn=True)

# --------- LỜI CẢM ƠN ---------
add_h1(doc, "LỜI CẢM ƠN")

paragraphs_thanks = [
    "Lời đầu tiên, em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến ThS. Dương Thành Phết – giảng viên hướng dẫn đồ án môn Đồ án cơ sở. Trong suốt quá trình thực hiện đề tài, thầy đã tận tình hướng dẫn, định hướng giải pháp, giải đáp thắc mắc, cung cấp tài liệu tham khảo có giá trị và đưa ra những góp ý sắc bén giúp em hoàn thành đồ án một cách tốt nhất. Sự hướng dẫn tận tâm, kiên nhẫn của thầy là nguồn động lực lớn giúp em vượt qua những khó khăn trong quá trình nghiên cứu và triển khai hệ thống.",
    "Em xin gửi lời cảm ơn đến Ban Giám hiệu Trường Đại học Công nghệ TP.HCM và Khoa Công nghệ Thông tin đã tạo điều kiện thuận lợi, cung cấp cơ sở vật chất, phòng thực hành, thư viện và môi trường học tập hiện đại trong suốt quá trình học tập tại trường. Chương trình đào tạo bài bản và các môn học chuyên ngành của Khoa đã trang bị cho em nền tảng kiến thức vững chắc để thực hiện đề tài này.",
    "Em cũng xin gửi lời cảm ơn đến toàn thể quý thầy cô trong Khoa Công nghệ Thông tin đã truyền đạt kiến thức trong suốt những năm học vừa qua. Những kiến thức về Lập trình hướng đối tượng, Cấu trúc dữ liệu và giải thuật, Cơ sở dữ liệu, Mạng máy tính, Kỹ thuật phần mềm, Lập trình Web, An toàn thông tin và Trí tuệ nhân tạo mà quý thầy cô đã giảng dạy là nền tảng trực tiếp để em xây dựng hệ thống trong đồ án này.",
    "Em xin chân thành cảm ơn các anh chị khóa trên, bạn bè cùng lớp đã chia sẻ tài liệu, kinh nghiệm và hỗ trợ kỹ thuật trong quá trình em thực hiện đồ án. Sự trao đổi cởi mở trong các nhóm học tập đã giúp em có cái nhìn đa chiều và giải quyết nhiều vấn đề phát sinh trong quá trình phát triển hệ thống.",
    "Cuối cùng, em xin bày tỏ lòng biết ơn sâu sắc đến gia đình đã luôn động viên, ủng hộ và tạo mọi điều kiện tốt nhất để em học tập và hoàn thành đồ án. Sự quan tâm và khích lệ của gia đình là nguồn sức mạnh tinh thần vô giá giúp em vượt qua mọi khó khăn trong suốt quá trình học tập và thực hiện đề tài.",
    "Do thời gian thực hiện và kinh nghiệm còn hạn chế, đồ án chắc chắn không tránh khỏi những thiếu sót về cả nội dung lẫn hình thức. Em rất mong nhận được sự góp ý, nhận xét chân thành từ quý thầy cô để em có thể hoàn thiện hơn trong tương lai.",
    "Em xin chân thành cảm ơn!",
]
for t in paragraphs_thanks:
    add_para(doc, t)

add_p(doc, "", space_after=24)
add_p(doc, "TP. Hồ Chí Minh, tháng 5 năm 2026",
      align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=2)
add_p(doc, "Sinh viên thực hiện",
      align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=36)
add_p(doc, "Nguyễn Nhật Hảo",
      align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)


# --------- LỜI MỞ ĐẦU ---------
add_h1(doc, "LỜI MỞ ĐẦU")

opening = [
    "Trong bối cảnh cách mạng công nghiệp lần thứ tư đang diễn ra mạnh mẽ trên phạm vi toàn cầu, công nghệ thông tin đã trở thành động lực then chốt thúc đẩy sự phát triển kinh tế – xã hội của mọi quốc gia. Tại Việt Nam, Chính phủ đã ban hành nhiều chính sách quan trọng nhằm thúc đẩy chuyển đổi số trong các doanh nghiệp, đặc biệt là doanh nghiệp nhỏ và vừa – khu vực kinh tế chiếm gần 98% tổng số doanh nghiệp và đóng góp khoảng 45% GDP cả nước.",
    "Tuy nhiên, một thực tế đáng quan tâm là phần lớn doanh nghiệp nhỏ và vừa, đặc biệt là các doanh nghiệp nhỏ đa ngành, vẫn đang sử dụng các phương pháp quản lý truyền thống như bảng tính Excel, sổ ghi chép tay hoặc các nhóm chat trên các ứng dụng nhắn tin để quản lý công việc, phân công nhân sự và theo dõi tiến độ. Cách làm này dẫn đến nhiều bất cập như thiếu minh bạch, phân công không tối ưu, khó tổng hợp dữ liệu, không có cơ sở khoa học để ra quyết định và đặc biệt là chưa tận dụng được sức mạnh của trí tuệ nhân tạo trong việc gợi ý nhân viên phù hợp cho từng công việc.",
    "Nhận thức được những vấn đề thực tế đó, đề tài “Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ Đa ngành Tích hợp AI” được lựa chọn nhằm xây dựng một nền tảng phần mềm toàn diện, hiện đại, đáp ứng các nhu cầu cốt lõi của doanh nghiệp nhỏ trong việc quản lý nhân viên, dự án, công việc và chấm công. Đặc biệt, hệ thống tích hợp module AI gợi ý nhân viên dùng Google Gemini gemini-2.5-flash xếp hạng định tính theo kỹ năng, tiến độ, đúng hạn và chấm công, cho phép người quản lý nhanh chóng tìm được nhân sự phù hợp nhất cho từng công việc cụ thể.",
    "Đồ án được thực hiện trong khuôn khổ môn học Đồ án cơ sở thuộc chương trình đào tạo Công nghệ thông tin của Khoa Công nghệ Thông tin. Báo cáo này được tổ chức thành 6 chương, trình bày một cách có hệ thống từ tổng quan đề tài, cơ sở lý thuyết, phân tích – thiết kế hệ thống, triển khai ứng dụng đến kiểm thử, demo và kết luận. Toàn bộ mã nguồn, tài liệu kỹ thuật và sản phẩm cuối cùng đều được công bố tại kho lưu trữ của tác giả để phục vụ mục đích học tập và phát triển tiếp theo.",
    "Em xin trân trọng giới thiệu báo cáo và mong nhận được những ý kiến đóng góp quý báu của quý thầy cô.",
]
for t in opening:
    add_para(doc, t)


# --------- MỤC LỤC ---------
add_h1(doc, "MỤC LỤC")
# Chèn TOC field thật – Word tự cập nhật bằng F9 hoặc khi mở file
add_toc_field(doc, levels="1-3")

# Mục lục thủ công không còn cần thiết – TOC field tự dựng từ Heading 1/2/3.


# --------- DANH MỤC CÁC TỪ VIẾT TẮT ---------
add_h1(doc, "DANH MỤC CÁC TỪ VIẾT TẮT")

abbr_rows = [
    ("AI", "Artificial Intelligence – Trí tuệ nhân tạo"),
    ("API", "Application Programming Interface – Giao diện lập trình ứng dụng"),
    ("BCrypt", "Thuật toán băm mật khẩu dựa trên Blowfish"),
    ("CRUD", "Create, Read, Update, Delete – Thao tác cơ bản với dữ liệu"),
    ("CSDL", "Cơ sở dữ liệu"),
    ("DTO", "Data Transfer Object – Đối tượng truyền dữ liệu"),
    ("ERD", "Entity Relationship Diagram – Sơ đồ quan hệ thực thể"),
    ("GVHD", "Giảng viên hướng dẫn"),
    ("GVPB", "Giảng viên phản biện"),
    ("HTTP", "HyperText Transfer Protocol – Giao thức truyền siêu văn bản"),
    ("HTTPS", "HTTP Secure – HTTP có mã hóa TLS/SSL"),
    ("JPA", "Java Persistence API"),
    ("JSON", "JavaScript Object Notation – Định dạng dữ liệu chuẩn"),
    ("JWT", "JSON Web Token – Chuẩn xác thực dạng token"),
    ("LLM", "Large Language Model – Mô hình ngôn ngữ lớn"),
    ("MSSV", "Mã số sinh viên"),
    ("MVC", "Model – View – Controller"),
    ("ORM", "Object–Relational Mapping – Ánh xạ đối tượng – quan hệ"),
    ("REST", "Representational State Transfer – Phong cách kiến trúc API"),
    ("SDK", "Software Development Kit – Bộ công cụ phát triển phần mềm"),
    ("SME", "Small and Medium-sized Enterprise – Doanh nghiệp nhỏ và vừa"),
    ("SPA", "Single Page Application – Ứng dụng web một trang"),
    ("SQL", "Structured Query Language – Ngôn ngữ truy vấn cấu trúc"),
    ("UML", "Unified Modeling Language – Ngôn ngữ mô hình hóa thống nhất"),
    ("UI", "User Interface – Giao diện người dùng"),
    ("UX", "User Experience – Trải nghiệm người dùng"),
]
add_table(doc, headers=["Từ viết tắt", "Diễn giải"],
          rows=abbr_rows, col_widths=[3.0, 12.5])


# --------- DANH MỤC CÁC BẢNG ---------
add_h1(doc, "DANH MỤC CÁC BẢNG")

table_list = [
    ("Bảng 2.1", "So sánh các framework backend Java phổ biến", "16"),
    ("Bảng 2.2", "Cấu trúc 3 phần của JSON Web Token", "21"),
    ("Bảng 2.3", "So sánh PostgreSQL với một số RDBMS khác", "27"),
    ("Bảng 2.4", "So sánh Redis với các giải pháp cache khác", "30"),
    ("Bảng 3.1", "Bảng phân tích ưu/nhược điểm của phương pháp quản lý hiện hành", "51"),
    ("Bảng 3.2", "Danh sách yêu cầu chức năng của hệ thống", "53"),
    ("Bảng 3.3", "Danh sách yêu cầu phi chức năng", "56"),
    ("Bảng 3.4", "Đặc tả use case UC-01: Đăng nhập", "58"),
    ("Bảng 3.5", "Đặc tả use case UC-08: Tạo công việc và gán nhân viên", "60"),
    ("Bảng 3.6", "Đặc tả use case UC-11: AI gợi ý nhân viên", "61"),
    ("Bảng 3.7", "Mô tả bảng users", "65"),
    ("Bảng 3.8", "Mô tả bảng employees", "66"),
    ("Bảng 3.9", "Mô tả bảng projects", "67"),
    ("Bảng 3.10", "Mô tả bảng tasks", "68"),
    ("Bảng 3.11", "Mô tả bảng attendances", "69"),
    ("Bảng 3.12", "Mô tả bảng suggestions", "69"),
    ("Bảng 4.1", "Yêu cầu phần cứng và phần mềm cần thiết", "82"),
    ("Bảng 4.2", "Danh sách các REST endpoint của backend", "96"),
    ("Bảng 4.3", "Tiêu chí xếp hạng AI gợi ý nhân viên", "100"),
    ("Bảng 5.1", "Ma trận test case theo module", "115"),
    ("Bảng 5.2", "Kịch bản kiểm thử module Auth", "117"),
    ("Bảng 5.3", "Kịch bản kiểm thử module Employee", "118"),
    ("Bảng 5.4", "Kịch bản kiểm thử module Task & Project", "119"),
    ("Bảng 5.5", "Kịch bản kiểm thử module AI Suggestion", "120"),
    ("Bảng 5.6", "Tổng hợp kết quả kiểm thử 42 test cases", "122"),
    ("Bảng 5.7", "Kịch bản kiểm thử module Geofence (GPS chấm công)", "123"),
    ("Bảng 5.8", "Lược đồ dữ liệu mở rộng cho xác thực chấm công GPS", "125"),
]
add_table(doc, headers=["Số hiệu", "Tên bảng", "Trang"],
          rows=table_list, col_widths=[2.5, 11.0, 2.0])


# --------- DANH MỤC HÌNH ---------
add_h1(doc, "DANH MỤC CÁC HÌNH ẢNH, SƠ ĐỒ")

fig_list = [
    ("Hình 2.1", "Mô hình kiến trúc Client–Server", "11"),
    ("Hình 2.2", "Mô hình 3 tầng (3-tier architecture)", "12"),
    ("Hình 2.3", "Luồng xử lý request RESTful API", "14"),
    ("Hình 2.4", "Vòng đời Spring Bean trong Spring Boot", "17"),
    ("Hình 2.5", "Quy trình xác thực JWT trong hệ thống", "22"),
    ("Hình 2.6", "Ánh xạ Entity Java – Bảng PostgreSQL qua JPA", "24"),
    ("Hình 2.7", "Mô hình caching write-through với Redis", "30"),
    ("Hình 2.8", "Vòng đời component React và Virtual DOM", "32"),
    ("Hình 2.9", "Mô hình widget tree của Flutter", "37"),
    ("Hình 2.10", "So sánh ứng dụng truyền thống và ứng dụng container hóa", "40"),
    ("Hình 2.11", "Quy trình gọi Google Gemini API", "43"),
    ("Hình 3.1", "Sơ đồ Use Case tổng thể của hệ thống", "57"),
    ("Hình 3.2", "Sơ đồ ERD của hệ thống", "64"),
    ("Hình 3.3", "Sơ đồ lớp (Class Diagram)", "71"),
    ("Hình 3.4", "Sơ đồ tuần tự – Đăng nhập", "73"),
    ("Hình 3.5", "Sơ đồ tuần tự – AI gợi ý nhân viên", "74"),
    ("Hình 3.6", "Sơ đồ tuần tự – Tạo công việc và gán nhân viên", "75"),
    ("Hình 3.7", "Sơ đồ hoạt động – Chấm công ngày", "76"),
    ("Hình 3.8", "Sơ đồ kiến trúc tổng thể hệ thống", "79"),
    ("Hình 3.9", "Wireframe trang Dashboard", "80"),
    ("Hình 3.10", "Wireframe trang AI Suggestion", "81"),
    ("Hình 4.1", "Cấu trúc thư mục backend Spring Boot", "85"),
    ("Hình 4.2", "Cấu trúc thư mục frontend React", "85"),
    ("Hình 4.3", "Sơ đồ luồng module AiSuggestionService", "101"),
    ("Hình 4.4", "Sơ đồ Docker Compose", "110"),
    ("Hình 5.1", "Màn hình Đăng nhập", "124"),
    ("Hình 5.2", "Màn hình Đăng ký tài khoản", "125"),
    ("Hình 5.3", "Màn hình Dashboard – tổng quan hệ thống", "126"),
    ("Hình 5.4", "Màn hình Quản lý nhân viên", "127"),
    ("Hình 5.5", "Màn hình Quản lý dự án", "128"),
    ("Hình 5.6", "Màn hình Quản lý công việc", "129"),
    ("Hình 5.7", "Màn hình Chấm công", "130"),
    ("Hình 5.8", "Màn hình AI gợi ý nhân viên – form nhập", "131"),
    ("Hình 5.9", "Màn hình AI gợi ý – kết quả top 5 nhân viên", "132"),
    ("Hình 5.10", "Dashboard cá nhân của nhân viên", "133"),
    ("Hình 5.11", "Trang Công việc của tôi (My Tasks)", "133"),
    ("Hình 5.12", "Trang Chấm công của tôi (My Attendance)", "134"),
    ("Hình 5.13", "Trang Dự án – góc nhìn nhân viên (chỉ đọc)", "134"),
]
add_table(doc, headers=["Số hiệu", "Tên hình", "Trang"],
          rows=fig_list, col_widths=[2.5, 11.0, 2.0])


# ==================================================================
# SECTION 3: NỘI DUNG CHÍNH – đánh số Ả Rập 1, 2, 3
# ==================================================================
sec3 = new_section(doc, page_num_fmt="decimal", start=1, show_pn=True)

# ==================================================================
# CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI
# ==================================================================
add_h1(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI")

add_h2(doc, "1.1. Đặt vấn đề")
for t in [
    "Trong xu thế hội nhập kinh tế quốc tế và cuộc cách mạng công nghiệp lần thứ tư (Industry 4.0), chuyển đổi số đã trở thành yêu cầu tất yếu đối với mọi doanh nghiệp, không phân biệt quy mô lớn nhỏ hay lĩnh vực kinh doanh. Theo số liệu của Bộ Kế hoạch và Đầu tư, tính đến ngày 31/12/2024, cả nước có khoảng 940.000 doanh nghiệp đang hoạt động, trong đó doanh nghiệp nhỏ và vừa (Small and Medium-sized Enterprise – SME) chiếm gần 98%. Khu vực này đóng góp khoảng 45% GDP và thu hút khoảng 5,5 triệu lao động, giữ vai trò quan trọng đối với nền kinh tế.",
    "Tuy đóng vai trò quan trọng như vậy, nhưng các doanh nghiệp nhỏ tại Việt Nam – đặc biệt là doanh nghiệp nhỏ đa ngành – vẫn đang gặp rất nhiều khó khăn trong việc ứng dụng công nghệ thông tin vào quản lý hoạt động nội bộ. Khảo sát của Phòng Thương mại và Công nghiệp Việt Nam (VCCI) trên 1.000 doanh nghiệp nhỏ và vừa cho thấy phần mềm kế toán là loại được ứng dụng phổ biến nhất với 748/1.000 doanh nghiệp, trong khi phần mềm quản lý nhân sự chỉ có 146/1.000 doanh nghiệp sử dụng, và gần như không doanh nghiệp nào dùng hệ thống quản lý công việc hay phê duyệt nội bộ. Phần lớn doanh nghiệp nhỏ vẫn dựa vào bảng tính Excel, sổ ghi chép tay và các nhóm chat trên Zalo, Facebook Messenger để phân công công việc, theo dõi tiến độ và quản lý nhân sự.",
    "Việc quản lý theo phương pháp truyền thống dẫn đến hàng loạt bất cập nghiêm trọng, ảnh hưởng trực tiếp đến năng suất lao động và khả năng cạnh tranh của doanh nghiệp:",
]:
    add_para(doc, t)

for s in [
    "Thiếu minh bạch và khó theo dõi: Khi công việc được giao qua tin nhắn hoặc email, rất khó để người quản lý cũng như nhân viên có cái nhìn tổng quan về tất cả các đầu việc đang triển khai. Tiến độ công việc bị phân tán ở nhiều nơi, không có nguồn thông tin duy nhất (single source of truth).",
    "Phân công công việc chưa tối ưu: Người quản lý thường phải dựa vào kinh nghiệm, trí nhớ và cảm tính để phân công công việc. Họ khó nắm bắt được chính xác kỹ năng, khối lượng công việc hiện tại, hiệu suất quá khứ và tình hình chuyên cần của từng nhân viên. Hệ quả là công việc có thể bị giao cho người không phù hợp, hoặc dồn quá nhiều vào một vài cá nhân.",
    "Quản lý nhân sự kém hiệu quả: Việc chấm công, theo dõi giờ làm, đánh giá hiệu suất và lưu trữ hồ sơ nhân viên đòi hỏi rất nhiều công sức thủ công. Nhiều doanh nghiệp nhỏ phải thuê dịch vụ kế toán bên ngoài để xử lý lương thưởng dựa trên bảng chấm công Excel, dễ phát sinh sai sót và gian lận.",
    "Không có dữ liệu để ra quyết định: Vì dữ liệu nằm rải rác ở nhiều nơi, doanh nghiệp khó tổng hợp được các báo cáo tổng quan như: nhân viên nào hiệu suất cao nhất, dự án nào đang chậm tiến độ, kỹ năng nào đang thiếu hụt trong tổ chức. Việc ra quyết định nhân sự, tuyển dụng, đào tạo do đó thiếu cơ sở khoa học.",
    "Khó mở rộng quy mô: Khi doanh nghiệp tăng trưởng, số lượng nhân viên và dự án tăng lên, mô hình quản lý truyền thống nhanh chóng bộc lộ giới hạn. Việc chuyển đổi sang phần mềm chuyên dụng vào thời điểm này thường tốn kém và rủi ro vì dữ liệu lịch sử không được số hóa.",
    "Chưa khai thác được sức mạnh AI: Trong khi các tập đoàn lớn đã ứng dụng trí tuệ nhân tạo để gợi ý phân công công việc, dự đoán tiến độ và tự động hóa quy trình, các doanh nghiệp nhỏ hoàn toàn đứng ngoài cuộc do thiếu công cụ và chi phí triển khai cao.",
]:
    add_bullet(doc, s)

for t in [
    "Đặc thù của doanh nghiệp nhỏ đa ngành – tức là các doanh nghiệp cung cấp nhiều loại sản phẩm/dịch vụ thuộc các lĩnh vực khác nhau (ví dụ một công ty vừa làm phần mềm, vừa làm marketing, vừa làm tổ chức sự kiện) – càng làm cho bài toán quản lý trở nên phức tạp. Nhân viên trong các doanh nghiệp này thường có kỹ năng đa dạng và phải tham gia vào nhiều dự án thuộc nhiều ngành khác nhau. Việc lựa chọn đúng người, đúng việc trở thành thách thức lớn của người quản lý.",
    "Mặt khác, sự phát triển nhanh chóng của các công nghệ web hiện đại (Spring Boot, React, Flutter), công nghệ container (Docker), và đặc biệt là sự bùng nổ của các Large Language Model (LLM) như Google Gemini đã mở ra cơ hội xây dựng các phần mềm quản lý tinh gọn, chi phí thấp nhưng tích hợp được trí tuệ nhân tạo. Đây chính là động lực để đề tài này được lựa chọn.",
]:
    add_para(doc, t)


add_h2(doc, "1.2. Lý do chọn đề tài")
add_para(doc,
    "Đề tài “Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ Đa ngành Tích hợp AI” được "
    "lựa chọn vì nhiều lý do thiết thực, vừa đáp ứng nhu cầu của xã hội, vừa phù hợp "
    "với mục tiêu học tập của môn Đồ án cơ sở:")
for t in [
    "Thứ nhất, đề tài có tính thực tiễn rất cao. Khác với những đồ án nặng tính lý thuyết hoặc mô phỏng, hệ thống xây dựng trong đề tài này hoàn toàn có thể được triển khai và sử dụng thực tế cho các doanh nghiệp nhỏ với chi phí rất thấp (chủ yếu là chi phí hạ tầng cloud và API Gemini). Điều này tạo ra giá trị thực sự cho xã hội.",
    "Thứ hai, đề tài cho phép vận dụng đầy đủ các kiến thức đã học trong chương trình đào tạo Công nghệ thông tin: lập trình hướng đối tượng (Java), cơ sở dữ liệu (PostgreSQL), mạng máy tính (HTTP, REST), an toàn thông tin (BCrypt, JWT, HTTPS), trí tuệ nhân tạo (Machine Learning, LLM), kỹ thuật phần mềm (UML, Agile), và các công nghệ web hiện đại (React, Flutter, Docker).",
    "Thứ ba, đề tài tích hợp trực tiếp công nghệ trí tuệ nhân tạo: backend gom dữ liệu lịch sử thô (task, chấm công, kỹ năng tự nhập của nhân viên), sau đó để Google Gemini gemini-2.5-flash xếp hạng định tính TOP 5 nhân viên cho từng task kèm `reasoning` bằng tiếng Việt. Đây là một chủ đề rất “nóng” trong giai đoạn 2024–2026: dùng LLM thay cho công thức trọng số cố định, vừa linh hoạt với dữ liệu tự do vừa giúp giải thích được kết quả cho người quản lý.",
    "Thứ tư, kiến trúc đa nền tảng (multi-platform) của hệ thống – gồm web (React), mobile (Flutter) và backend (Spring Boot) – là kiến trúc đang được sử dụng phổ biến trong các sản phẩm thương mại hiện nay. Việc triển khai thành công đề tài giúp sinh viên có kinh nghiệm thực tế khi đi làm.",
    "Thứ năm, đề tài phù hợp với phạm vi và thời lượng của môn Đồ án cơ sở (12 tuần). Mỗi module chức năng có thể được triển khai trong một vài tuần, các module được tích hợp dần và kiểm thử song song. Mức độ phức tạp vừa phải, không quá đơn giản như một CRUD app cơ bản, cũng không quá tham vọng như các hệ thống phân tán quy mô lớn.",
    "Thứ sáu, hệ thống có thể được mở rộng tự nhiên thành một sản phẩm thương mại (SaaS) trong tương lai. Điều này khích lệ tinh thần khởi nghiệp và đổi mới sáng tạo trong sinh viên.",
]:
    add_para(doc, t)


add_h2(doc, "1.3. Mục tiêu đề tài")
add_h3(doc, "1.3.1. Mục tiêu tổng quát")
add_para(doc,
    "Xây dựng một hệ thống phần mềm hoàn chỉnh giúp doanh nghiệp nhỏ đa ngành quản "
    "lý hiệu quả nhân sự, dự án, công việc và chấm công, đồng thời tích hợp trí tuệ "
    "nhân tạo để gợi ý nhân viên phù hợp cho từng công việc, nâng cao năng suất lao "
    "động và chất lượng ra quyết định của người quản lý.")
add_h3(doc, "1.3.2. Mục tiêu cụ thể")
for s in [
    "Xây dựng kiến trúc backend RESTful API bằng Spring Boot 3.5.0 với bảo mật JWT và phân quyền theo 3 role (ADMIN/MANAGER/EMPLOYEE).",
    "Thiết kế cơ sở dữ liệu chuẩn hóa PostgreSQL 16 gồm 6 bảng chính: users, employees, projects, tasks, attendances, suggestions.",
    "Triển khai 7 nhóm chức năng nghiệp vụ: Xác thực & phân quyền, Quản lý nhân viên (kèm trường kỹ năng tự nhập), Quản lý dự án, Quản lý công việc, Chấm công, AI Gợi ý, Dashboard.",
    "Xây dựng frontend Single Page Application bằng React 18 + Vite 5 + Tailwind CSS, sử dụng React Router v6 để điều hướng; giao diện thích ứng theo role.",
    "Phát triển ứng dụng mobile Flutter chạy được trên cả Android và iOS với các tính năng cốt lõi.",
    "Tích hợp Google Gemini gemini-2.5-flash để xếp hạng nhân viên phù hợp cho từng task dựa trên kỹ năng, tiến độ, đúng hạn và chấm công — KHÔNG dùng công thức trọng số cố định mà để LLM đánh giá định tính.",
    "Sử dụng Redis làm tầng cache để tối ưu hiệu năng cho các truy vấn nặng (AI Suggestion).",
    "Đóng gói toàn bộ hệ thống bằng Docker Compose để dễ dàng triển khai và vận hành.",
    "Kiểm thử toàn diện với ít nhất 30 test cases, đảm bảo các luồng nghiệp vụ chính hoạt động đúng.",
    "Viết tài liệu đầy đủ: báo cáo đồ án, API Specification, Database Schema, Setup Guide, UML Diagrams.",
]:
    add_bullet(doc, s)


add_h2(doc, "1.4. Nội dung đề tài")
add_para(doc,
    "Đề tài tập trung vào ba khối nội dung chính: (1) nghiên cứu cơ sở lý thuyết và "
    "công nghệ, (2) phân tích – thiết kế hệ thống, (3) triển khai – kiểm thử sản phẩm. "
    "Cụ thể các nội dung được thực hiện như sau:")
add_h3(doc, "1.4.1. Nghiên cứu cơ sở lý thuyết")
for s in [
    "Nghiên cứu kiến trúc Client–Server, mô hình 3 tầng và RESTful API.",
    "Nghiên cứu framework Spring Boot 3.x: cấu hình tự động, Spring MVC, Spring Data JPA, Spring Security, Spring Cache.",
    "Nghiên cứu chuẩn JWT (RFC 7519) và cơ chế xác thực stateless.",
    "Nghiên cứu hệ quản trị PostgreSQL 16 và đặc thù khi dùng với JPA/Hibernate.",
    "Nghiên cứu Redis 7 và các pattern caching phổ biến.",
    "Nghiên cứu React 18 (hooks, Context API, React Router v6) và build tool Vite 5.",
    "Nghiên cứu Flutter 3.x – widget tree, state management, gọi HTTP qua dio.",
    "Nghiên cứu Docker và Docker Compose – container hóa multi-service applications.",
    "Nghiên cứu Google Gemini API và các kỹ thuật prompt engineering cơ bản.",
    "Nghiên cứu các cách xếp hạng đa tiêu chí: Weighted Scoring (thiết kế ban đầu), so sánh với cách dùng LLM (phương án triển khai thực tế).",
]:
    add_bullet(doc, s)
add_h3(doc, "1.4.2. Phân tích – thiết kế hệ thống")
for s in [
    "Khảo sát hiện trạng quản lý công việc tại doanh nghiệp nhỏ Việt Nam qua số liệu thống kê, khảo sát ngành và tài liệu chuyên ngành.",
    "Tổng hợp 14 yêu cầu chức năng và 8 yêu cầu phi chức năng.",
    "Vẽ Use Case Diagram (14 use case) và đặc tả chi tiết 3 use case tiêu biểu theo mẫu Cockburn.",
    "Thiết kế ERD với 6 bảng, mô tả chi tiết các trường và ràng buộc.",
    "Vẽ Class Diagram thể hiện mối quan hệ giữa các Entity và Service.",
    "Vẽ Sequence Diagram cho 5 luồng nghiệp vụ chính.",
    "Vẽ Activity Diagram cho luồng chấm công và luồng AI gợi ý.",
    "Thiết kế kiến trúc tổng thể và sơ đồ triển khai.",
    "Thiết kế wireframe cho các trang quan trọng (Dashboard, AI Suggestion).",
]:
    add_bullet(doc, s)
add_h3(doc, "1.4.3. Triển khai – kiểm thử")
for s in [
    "Triển khai backend Spring Boot với 25+ class chia thành các tầng controller, service, repository, entity, dto, security, config.",
    "Triển khai frontend React với 10 trang chính, tích hợp Chart.js cho dashboard.",
    "Triển khai app Flutter với 6 màn hình chính.",
    "Triển khai thuật toán AiSuggestionService với cache Redis.",
    "Tích hợp Gemini API trực tiếp trong AiSuggestionService qua RestClient của Spring.",
    "Viết Dockerfile cho từng service và file docker-compose.yml.",
    "Kiểm thử với 42 test cases bao phủ 9 module nghiệp vụ: Auth, Employee, Project, Task, Attendance, AI Suggestion, Dashboard, Self-service và Geofence.",
    "Viết tài liệu API bằng Swagger UI và tài liệu kỹ thuật bằng Markdown.",
]:
    add_bullet(doc, s)


add_h2(doc, "1.5. Phương pháp nghiên cứu")
for t in [
    "Để hoàn thành các mục tiêu đề ra, đề tài sử dụng kết hợp nhiều phương pháp nghiên cứu, vừa định tính vừa định lượng, vừa lý thuyết vừa thực nghiệm:",
]:
    add_para(doc, t)
add_h3(doc, "1.5.1. Phương pháp nghiên cứu tài liệu")
add_para(doc,
    "Tiến hành thu thập, đọc và phân tích các tài liệu chính thức của các công nghệ "
    "sử dụng trong đề tài (Spring Boot Reference Documentation, React Documentation, "
    "Flutter Documentation, PostgreSQL Manual, Docker Documentation, Gemini API "
    "Documentation). Đồng thời tham khảo các sách chuyên ngành về kỹ thuật phần mềm, "
    "kiến trúc phần mềm và trí tuệ nhân tạo. Các bài báo khoa học trên Google Scholar "
    "về hệ thống gợi ý đa tiêu chí cũng được khảo cứu để xây dựng thuật toán phù hợp.")
add_h3(doc, "1.5.2. Phương pháp khảo sát hiện trạng")
add_para(doc,
    "Khảo sát hiện trạng quản lý công việc tại các doanh nghiệp nhỏ đa ngành ở "
    "Việt Nam thông qua nghiên cứu tài liệu thứ cấp: số liệu thống kê của Bộ Kế "
    "hoạch và Đầu tư, khảo sát của Phòng Thương mại và Công nghiệp Việt Nam (VCCI) "
    "về mức độ ứng dụng phần mềm, cùng các bài viết chuyên ngành về chuyển đổi số "
    "trong doanh nghiệp nhỏ và vừa. Đồng thời tiến hành so sánh, phân tích các công "
    "cụ quản lý công việc đang phổ biến (Excel, Trello, phần mềm ERP) nhằm xác định "
    "khoảng trống mà đề tài hướng tới giải quyết.")
add_h3(doc, "1.5.3. Phương pháp phân tích thiết kế hướng đối tượng")
add_para(doc,
    "Sử dụng ngôn ngữ mô hình hóa thống nhất (UML) để phân tích và thiết kế hệ thống. "
    "Các loại sơ đồ được sử dụng bao gồm Use Case Diagram, Class Diagram, Sequence "
    "Diagram, Activity Diagram và ERD. Phương pháp này giúp tách bạch giai đoạn "
    "phân tích – thiết kế khỏi giai đoạn triển khai, đảm bảo kiến trúc hệ thống "
    "rõ ràng trước khi bắt tay vào viết code.")
add_h3(doc, "1.5.4. Phương pháp phát triển phần mềm Agile")
add_para(doc,
    "Áp dụng phương pháp Agile/Scrum với chu kỳ sprint 2 tuần. Sau mỗi sprint, một "
    "nhóm chức năng được hoàn thành và demo cho giảng viên hướng dẫn để nhận góp ý. "
    "Cách làm này giúp phát hiện và khắc phục sớm các vấn đề về thiết kế cũng như "
    "triển khai, tránh tình trạng đến cuối kỳ mới phát hiện sai sót lớn.")
add_h3(doc, "1.5.5. Phương pháp kiểm thử thực nghiệm")
add_para(doc,
    "Sau khi triển khai xong, hệ thống được kiểm thử theo phương pháp hộp đen "
    "(black-box testing) với 42 test cases bao phủ tất cả các luồng nghiệp vụ "
    "chính. Sử dụng Postman để kiểm thử API và trình duyệt Chrome để kiểm thử "
    "giao diện. Kết quả kiểm thử được ghi nhận và phân tích trong Chương 5.")


add_h2(doc, "1.6. Phạm vi đề tài")
add_h3(doc, "1.6.1. Phạm vi chức năng")
add_para(doc,
    "Hệ thống tập trung vào các chức năng cốt lõi của quản lý công việc, không "
    "mở rộng sang các nghiệp vụ kế toán – tài chính:")
add_table(
    doc,
    headers=["Nhóm chức năng", "Mô tả ngắn"],
    rows=[
        ("Xác thực & phân quyền",
         "Đăng ký, đăng nhập, JWT, phân quyền ADMIN / MANAGER / EMPLOYEE."),
        ("Quản lý nhân viên",
         "CRUD nhân viên kèm trường `skills` (TEXT) do quản lý nhập tự do."),
        ("Quản lý dự án",
         "CRUD dự án, gán nhân viên vào dự án, theo dõi trạng thái dự án."),
        ("Quản lý công việc",
         "CRUD công việc, gán nhân viên, theo dõi tiến độ, deadline, mức ưu tiên."),
        ("Chấm công",
         "Ghi nhận chấm công ngày (PRESENT/ABSENT/LATE), xem lịch sử."),
        ("AI gợi ý nhân viên",
         "Gợi ý top 5 nhân viên phù hợp dựa trên 4 tiêu chí, kèm Gemini summary."),
        ("Dashboard & thống kê",
         "Tổng quan số nhân viên, dự án, task, biểu đồ trạng thái task."),
        ("API & tài liệu",
         "Swagger UI tự sinh tài liệu API, hỗ trợ developer test nhanh."),
    ],
    col_widths=[4.5, 11.0],
)
add_h3(doc, "1.6.2. Phạm vi loại trừ")
for s in [
    "Không bao gồm module kế toán, lương thưởng, hóa đơn, thuế.",
    "Không tích hợp với các phần mềm ERP / CRM bên ngoài.",
    "Không có tính năng video call hay chat nội bộ thời gian thực.",
    "Không có module hỗ trợ thiết bị chấm công vân tay/QR (chỉ chấm công thủ công).",
]:
    add_bullet(doc, s)
add_h3(doc, "1.6.3. Phạm vi kỹ thuật")
for s in [
    "Backend: Java 17+, Spring Boot 3.5.0, Maven 3.9+.",
    "Frontend: Node.js 18+, React 18, Vite 5, Tailwind CSS 3.",
    "Mobile: Flutter 3.x, Dart ≥ 3.0.",
    "Database: PostgreSQL 16.",
    "Cache: Redis 7.",
    "Container: Docker 24+, Docker Compose 2.x.",
    "AI: Google Gemini gemini-2.5-flash qua HTTP API.",
    "Hệ điều hành test: Windows 11, Ubuntu 22.04 LTS.",
    "Trình duyệt hỗ trợ: Chrome 120+, Edge 120+, Firefox 120+.",
]:
    add_bullet(doc, s)


add_h2(doc, "1.7. Kết cấu báo cáo")
add_para(doc, "Báo cáo đồ án được tổ chức thành 6 chương, không kể phần Lời cảm ơn, Mục lục và Tài liệu tham khảo:")
chapters_overview = [
    ("Chương 1 – Tổng quan đề tài",
     "Trình bày bối cảnh, lý do chọn đề tài, mục tiêu, nội dung, phương pháp nghiên cứu và phạm vi đề tài."),
    ("Chương 2 – Cơ sở lý thuyết",
     "Trình bày các kiến thức nền tảng về kiến trúc client–server, REST API, các công nghệ Spring Boot, React, Flutter, PostgreSQL, Redis, Docker và Gemini."),
    ("Chương 3 – Phân tích và thiết kế hệ thống",
     "Trình bày kết quả khảo sát hiện trạng, đặc tả yêu cầu, các sơ đồ UML (Use Case, Class, Sequence, Activity), thiết kế CSDL và kiến trúc tổng thể."),
    ("Chương 4 – Xây dựng ứng dụng",
     "Trình bày chi tiết quá trình triển khai backend, frontend, mobile, tích hợp Redis, Gemini và đóng gói bằng Docker."),
    ("Chương 5 – Kiểm thử và đánh giá kết quả",
     "Trình bày kế hoạch kiểm thử, 42 test cases, kết quả kiểm thử, demo giao diện thực tế (web + mobile) và quy trình khởi chạy hệ thống trên mọi nền tảng."),
    ("Chương 6 – Kết luận và hướng phát triển",
     "Tổng kết kết quả đạt được, các hạn chế của đề tài và đề xuất hướng phát triển trong tương lai."),
]
for title, desc in chapters_overview:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.75)
    r1 = p.add_run(title + ": ")
    set_run(r1, bold=True)
    r2 = p.add_run(desc)
    set_run(r2)


# ==================================================================
# CHƯƠNG 2: CƠ SỞ LÝ THUYẾT
# ==================================================================
add_h1(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT")

add_h2(doc, "2.1. Kiến trúc Client–Server và mô hình 3 tầng")
add_h3(doc, "2.1.1. Khái niệm kiến trúc Client–Server")
for t in [
    "Kiến trúc Client–Server là mô hình tính toán phân tán, trong đó các tác vụ và khối lượng công việc được phân chia giữa các thành phần đóng vai trò Server (máy phục vụ – cung cấp dịch vụ) và các thành phần đóng vai trò Client (máy khách – yêu cầu dịch vụ). Đây là mô hình kiến trúc phổ biến nhất hiện nay cho các ứng dụng web, mobile, doanh nghiệp và đa số các dịch vụ Internet.",
    "Trong mô hình này, Client gửi một yêu cầu (request) đến Server qua mạng (thường là Internet). Server nhận yêu cầu, xử lý logic, có thể truy vấn cơ sở dữ liệu hoặc gọi các dịch vụ bên thứ ba, sau đó trả lại phản hồi (response) cho Client. Giao tiếp giữa Client và Server diễn ra theo cơ chế request–response và thường sử dụng các giao thức tiêu chuẩn như HTTP/HTTPS, WebSocket, gRPC.",
    "Trong hệ thống quản lý công việc được xây dựng trong đề tài này, Server là ứng dụng Spring Boot chạy trên cổng 5000, đảm nhiệm toàn bộ logic nghiệp vụ và truy cập dữ liệu. Client gồm hai phía: (1) trình duyệt web với ứng dụng React chạy trên cổng 5173 và (2) ứng dụng Flutter chạy trên thiết bị di động Android/iOS. Cả hai phía Client đều giao tiếp với Server qua giao thức HTTP với định dạng dữ liệu JSON.",
]:
    add_para(doc, t)
add_h3(doc, "2.1.2. Mô hình 3 tầng (Three-tier Architecture)")
add_para(doc,
    "Mô hình 3 tầng là một biến thể nâng cao của kiến trúc Client–Server, trong đó "
    "logic của ứng dụng được tách thành ba tầng riêng biệt, mỗi tầng có vai trò "
    "rõ ràng và có thể được triển khai trên các máy chủ vật lý khác nhau:")
for s in [
    "Tầng trình bày (Presentation Tier): Đây là tầng giao tiếp với người dùng cuối, hiển thị giao diện và nhận thao tác từ người dùng. Trong đề tài, tầng này được hiện thực bởi React (web) và Flutter (mobile).",
    "Tầng nghiệp vụ (Business Logic Tier hay Application Tier): Đây là tầng xử lý logic nghiệp vụ – tính toán, kiểm tra ràng buộc, áp dụng quy tắc kinh doanh, gọi các dịch vụ khác. Trong đề tài, tầng này được hiện thực bởi backend Spring Boot.",
    "Tầng dữ liệu (Data Tier): Đây là tầng lưu trữ và quản lý dữ liệu lâu dài. Trong đề tài, tầng này gồm PostgreSQL 16 (lưu trữ dữ liệu chính) và Redis 7 (lưu trữ cache).",
]:
    add_bullet(doc, s)
add_para(doc,
    "Việc tách biệt 3 tầng mang lại nhiều lợi ích quan trọng: dễ bảo trì vì thay đổi "
    "ở một tầng không ảnh hưởng đến các tầng khác (miễn là giữ nguyên giao diện "
    "giao tiếp); dễ mở rộng (scale) vì có thể nhân bản từng tầng độc lập theo nhu "
    "cầu (ví dụ chạy nhiều instance backend nhưng chỉ một database master); dễ "
    "phân công công việc cho các đội phát triển khác nhau (frontend team, backend "
    "team, database team); và dễ tái sử dụng tầng backend cho nhiều loại Client "
    "(web, mobile, desktop, IoT, API cho đối tác).")
add_h3(doc, "2.1.3. Ưu và nhược điểm")
add_para(doc, "Mô hình Client–Server 3 tầng có những ưu điểm nổi bật như tính module hóa cao, khả năng mở rộng tốt, dễ bảo trì, hỗ trợ nhiều loại Client cùng lúc và phân quyền truy cập dữ liệu chặt chẽ ở tầng backend. Tuy nhiên, mô hình này cũng có một số nhược điểm cần lưu ý: (i) độ trễ mạng giữa các tầng, đặc biệt khi triển khai trên các máy chủ vật lý khác nhau; (ii) sự phụ thuộc vào kết nối mạng – khi mạng có sự cố thì Client không thể sử dụng dịch vụ; (iii) chi phí vận hành cao hơn so với ứng dụng đơn khối (monolithic) chạy trên một máy.")


add_h2(doc, "2.2. RESTful API và HTTP")
add_h3(doc, "2.2.1. Giao thức HTTP")
add_para(doc,
    "HyperText Transfer Protocol (HTTP) là giao thức truyền thông tin tầng ứng dụng "
    "được sử dụng cho World Wide Web. HTTP là giao thức không trạng thái "
    "(stateless) hoạt động theo mô hình request–response: Client gửi một HTTP "
    "request đến Server, Server xử lý và trả về một HTTP response.")
add_para(doc,
    "Một HTTP request gồm: (1) dòng start line chứa method (GET, POST, PUT, DELETE, "
    "PATCH...) và URL; (2) các header chứa metadata như Content-Type, "
    "Authorization, User-Agent; và (3) body (tuỳ chọn) chứa dữ liệu gửi đi. Một "
    "HTTP response gồm: (1) status line chứa HTTP version và status code (200 OK, "
    "201 Created, 400 Bad Request, 401 Unauthorized, 404 Not Found, 500 Internal "
    "Server Error...); (2) các response header; và (3) body chứa nội dung trả về.")
add_h3(doc, "2.2.2. Phong cách kiến trúc REST")
add_para(doc,
    "REST (Representational State Transfer) là phong cách kiến trúc do Roy Fielding "
    "đề xuất trong luận án tiến sĩ năm 2000. Một API tuân thủ phong cách REST "
    "(RESTful API) phải đáp ứng các nguyên tắc:")
for s in [
    "Client–Server: Tách biệt Client và Server, cho phép tiến hóa độc lập.",
    "Stateless: Server không lưu trạng thái phiên làm việc của Client; mỗi request phải chứa đầy đủ thông tin cần thiết để Server xử lý.",
    "Cacheable: Response có thể được cache để giảm tải cho Server.",
    "Uniform Interface: Sử dụng một giao diện thống nhất – tài nguyên được xác định qua URL, thao tác trên tài nguyên qua HTTP method.",
    "Layered System: Cho phép có các tầng trung gian (proxy, load balancer, gateway) mà không ảnh hưởng đến Client.",
    "Code on demand (tuỳ chọn): Server có thể gửi code thực thi xuống Client (ví dụ JavaScript).",
]:
    add_bullet(doc, s)
add_h3(doc, "2.2.3. RESTful API trong hệ thống")
add_para(doc,
    "Toàn bộ backend Spring Boot trong đề tài tuân thủ phong cách REST. Mỗi tài "
    "nguyên (employee, project, task, attendance) có URL riêng và được thao tác "
    "qua các HTTP method tương ứng: GET để đọc, POST để tạo mới, PUT để cập nhật "
    "toàn bộ, PATCH để cập nhật một phần, DELETE để xóa. Dữ liệu trao đổi giữa "
    "Client và Server đều ở định dạng JSON. Server trả về HTTP status code phù "
    "hợp cho từng kết quả xử lý.")


add_h2(doc, "2.3. Ngôn ngữ Java và Spring Boot Framework")
add_h3(doc, "2.3.1. Ngôn ngữ Java")
add_para(doc,
    "Java là ngôn ngữ lập trình hướng đối tượng do Sun Microsystems (nay là Oracle) "
    "phát triển và phát hành lần đầu vào năm 1995. Java có triết lý nổi tiếng "
    "“Write Once, Run Anywhere” – viết một lần, chạy mọi nơi – nhờ máy ảo Java "
    "(Java Virtual Machine – JVM) cho phép chương trình biên dịch ra bytecode "
    "chạy được trên mọi nền tảng có JVM cài đặt.")
add_para(doc,
    "Java có hệ sinh thái rất phong phú với hàng nghìn thư viện mã nguồn mở và "
    "được sử dụng rộng rãi trong các ứng dụng doanh nghiệp lớn, hệ thống ngân "
    "hàng, viễn thông và Android. Phiên bản Java 17 LTS (Long Term Support) "
    "được sử dụng trong đề tài là phiên bản hỗ trợ dài hạn của Oracle, đảm bảo "
    "tính ổn định và bảo mật.")
add_h3(doc, "2.3.2. Spring Framework và Spring Boot")
add_para(doc,
    "Spring Framework là framework Java mã nguồn mở phổ biến nhất, ra mắt năm "
    "2003. Spring giải quyết vấn đề lớn nhất của Java EE truyền thống là sự "
    "phức tạp khi cấu hình và phát triển ứng dụng doanh nghiệp. Spring áp dụng "
    "các nguyên tắc Dependency Injection (DI) và Inversion of Control (IoC) để "
    "tách rời các thành phần và tăng tính linh hoạt.")
add_para(doc,
    "Spring Boot là một dự án con của Spring Framework, ra mắt năm 2014. Spring "
    "Boot cung cấp cấu hình tự động (auto-configuration), starter dependencies "
    "(các nhóm thư viện được cấu hình sẵn theo từng tác vụ phổ biến) và Spring "
    "Boot Actuator (giám sát ứng dụng). Spring Boot giúp tạo ra các ứng dụng "
    "Spring chạy độc lập (standalone) bằng một câu lệnh duy nhất, không cần "
    "deploy lên application server.")
add_para(doc, "So sánh các framework backend Java phổ biến:")
add_table(
    doc,
    headers=["Framework", "Năm ra mắt", "Đặc điểm nổi bật"],
    rows=[
        ("Spring Boot", "2014", "Cấu hình tự động, hệ sinh thái lớn, được dùng rộng rãi."),
        ("Quarkus", "2019", "Tối ưu cho cloud-native, native image với GraalVM."),
        ("Micronaut", "2018", "Khởi động nhanh, ít tốn RAM, compile-time DI."),
        ("Helidon", "2018", "Do Oracle phát triển, hỗ trợ cả MicroProfile và reactive."),
        ("Play Framework", "2007", "Theo phong cách reactive, hỗ trợ Java và Scala."),
    ],
    col_widths=[3.5, 2.5, 9.5],
)
add_caption(doc, "Bảng 2.1: So sánh các framework backend Java phổ biến", kind="table")
add_h3(doc, "2.3.3. Các thành phần chính của Spring Boot sử dụng trong đề tài")
for s in [
    "Spring MVC: Xây dựng các REST Controller xử lý HTTP request qua các annotation @RestController, @GetMapping, @PostMapping, @PutMapping, @DeleteMapping.",
    "Spring Data JPA: Tầng trừu tượng hóa truy cập cơ sở dữ liệu trên cơ sở JPA (Hibernate). Chỉ cần khai báo Repository interface là có ngay các phương thức CRUD.",
    "Spring Security: Xử lý xác thực và phân quyền cho REST API, tích hợp với JWT.",
    "Spring Cache: Cung cấp các annotation @Cacheable, @CacheEvict, @CachePut để dễ dàng cache kết quả method.",
    "Spring Boot Actuator: Cung cấp các endpoint giám sát ứng dụng (health, metrics, env).",
    "Spring Web Configuration: Cấu hình CORS, ConverterFactory, HandlerInterceptor.",
]:
    add_bullet(doc, s)
add_h3(doc, "2.3.4. Vòng đời Spring Bean và Dependency Injection")
add_para(doc,
    "Trong Spring, các đối tượng được quản lý bởi container gọi là Spring Bean. "
    "Vòng đời của một Spring Bean gồm các giai đoạn: instantiation (khởi tạo "
    "đối tượng), population of properties (gán giá trị các thuộc tính), "
    "BeanNameAware/BeanFactoryAware (gán tên và reference container), "
    "BeanPostProcessor before initialization, "
    "afterPropertiesSet() và init method, BeanPostProcessor after initialization, "
    "ready to use, và cuối cùng là destroy method khi container shutdown.")
add_para(doc,
    "Dependency Injection (DI) là cơ chế trong đó các phụ thuộc (dependencies) "
    "của một đối tượng được container tự động truyền vào (inject) thay vì đối "
    "tượng tự khởi tạo. Spring hỗ trợ ba kiểu DI: constructor injection, setter "
    "injection và field injection. Trong đề tài, cách dùng phổ biến nhất là "
    "constructor injection thông qua annotation @Autowired hoặc Lombok "
    "@RequiredArgsConstructor để giảm boilerplate code.")


add_h2(doc, "2.4. Spring Security và JSON Web Token (JWT)")
add_h3(doc, "2.4.1. Tổng quan Spring Security")
add_para(doc,
    "Spring Security là framework bảo mật toàn diện cho ứng dụng Spring, "
    "cung cấp hai chức năng chính: Authentication (xác thực – kiểm tra danh "
    "tính người dùng) và Authorization (phân quyền – kiểm tra người dùng có "
    "được phép thực hiện hành động nào đó hay không). Spring Security hoạt "
    "động dựa trên chuỗi filter (filter chain) gắn vào Servlet container.")
add_para(doc,
    "Trong đề tài, Spring Security được cấu hình theo phong cách stateless – "
    "không sử dụng HttpSession – để phù hợp với kiến trúc RESTful API và để "
    "dễ scale theo chiều ngang (horizontal scaling). Toàn bộ trạng thái xác "
    "thực được mã hóa trong JWT và truyền qua header Authorization.")
add_h3(doc, "2.4.2. JSON Web Token (JWT)")
add_para(doc,
    "JWT (JSON Web Token) là chuẩn mở RFC 7519 định nghĩa cách truyền thông "
    "tin an toàn giữa các bên dưới dạng đối tượng JSON. JWT có ba phần được "
    "phân tách bằng dấu chấm: Header.Payload.Signature.")
add_table(
    doc,
    headers=["Phần", "Mô tả"],
    rows=[
        ("Header",
         "Chứa metadata: loại token (typ=JWT) và thuật toán ký (alg=HS256, RS256...)."),
        ("Payload",
         "Chứa các claim – thông tin về user (sub, iss, exp, iat, custom claims)."),
        ("Signature",
         "Chữ ký số: HMAC-SHA256(base64(Header) + '.' + base64(Payload), secret)."),
    ],
    col_widths=[3.0, 12.5],
)
add_caption(doc, "Bảng 2.2: Cấu trúc 3 phần của JSON Web Token", kind="table")
add_para(doc,
    "Mã hóa Base64URL được dùng cho Header và Payload, sau đó chữ ký HMAC-SHA256 "
    "(hoặc thuật toán bất đối xứng RS256, ES256) đảm bảo token không bị giả mạo. "
    "Chỉ bên nắm giữ secret key (server) mới có thể tạo và xác thực được chữ ký.")
add_h3(doc, "2.4.3. Quy trình xác thực JWT trong hệ thống")
for s in [
    "Bước 1: Client gửi POST /api/auth/login với username và password.",
    "Bước 2: Server xác thực credential. Nếu hợp lệ, Server tạo JWT chứa các claim (sub=username, roles, exp=2h sau) ký bằng HMAC-SHA256 với secret key.",
    "Bước 3: Server trả JWT cho Client. Client lưu vào localStorage hoặc SecureStorage (mobile).",
    "Bước 4: Mỗi request sau đó, Client gắn JWT vào header Authorization: Bearer <token>.",
    "Bước 5: JwtAuthenticationFilter ở backend đọc header, giải mã JWT, kiểm tra signature và thời hạn.",
    "Bước 6: Nếu hợp lệ, Filter tạo Authentication object và đẩy vào SecurityContext.",
    "Bước 7: Spring Security cho phép request đi tiếp đến Controller. Quyền truy cập từng method được kiểm soát bằng @PreAuthorize.",
    "Bước 8: Khi token hết hạn, Client phải đăng nhập lại (đề tài đơn giản, chưa triển khai refresh token).",
]:
    add_bullet(doc, s)
add_h3(doc, "2.4.4. So sánh JWT với session truyền thống")
add_para(doc,
    "Trong cơ chế xác thực truyền thống dựa trên session, server phải lưu trạng "
    "thái phiên của mỗi user (thường trong RAM hoặc database). Khi user gửi "
    "request, server tra session ID trong cookie để xác định danh tính. Cách "
    "này đơn giản nhưng có nhược điểm: khó scale ngang vì các instance phải "
    "chia sẻ session store (Redis, sticky session); không phù hợp cho mobile "
    "app vì mobile thường không quản lý cookie tự động.")
add_para(doc,
    "Ngược lại, JWT là stateless – server không cần lưu gì cả, tất cả thông tin "
    "đều nằm trong token. Điều này giúp dễ scale ngang, dễ tích hợp với mobile "
    "và microservices. Tuy nhiên JWT cũng có nhược điểm: khó thu hồi token "
    "trước hạn (cần cơ chế blacklist), kích thước token lớn hơn session ID, "
    "và nếu secret key bị lộ thì có thể giả mạo toàn bộ token.")
add_h3(doc, "2.4.5. Mã hóa mật khẩu bằng BCrypt")
add_para(doc,
    "Mật khẩu của user không bao giờ được lưu dưới dạng văn bản thuần. Trong "
    "đề tài, mật khẩu được băm bằng thuật toán BCrypt – một hàm băm chậm có "
    "thể điều chỉnh độ phức tạp (work factor). BCrypt tự động sinh salt cho "
    "mỗi mật khẩu, ngăn chặn tấn công rainbow table. Spring Security cung cấp "
    "sẵn lớp BCryptPasswordEncoder để mã hóa và so khớp mật khẩu.")

add_h3(doc, "2.4.6. Giới hạn tần suất (Rate Limiting) chống spam")
add_para(doc,
    "Để chống lạm dụng, hệ thống áp dụng cơ chế giới hạn tần suất request theo "
    "địa chỉ IP. Một servlet filter (`RateLimitFilter`) chạy trước chuỗi lọc "
    "của Spring Security, đếm số request mỗi IP trong cửa sổ thời gian cố định "
    "1 phút. Hai nhóm endpoint dễ bị lạm dụng được bảo vệ: nhóm xác thực "
    "`/api/auth/**` (chống tấn công dò mật khẩu – brute force) và nhóm gợi ý AI "
    "`/api/suggestions/**` (bảo vệ hạn mức gọi Gemini API). Khi một IP vượt "
    "ngưỡng cho phép, server trả về mã HTTP 429 (Too Many Requests) kèm thông "
    "báo rõ ràng thay vì xử lý tiếp. Ngưỡng được cấu hình qua biến môi trường "
    "`RATELIMIT_AUTH` (mặc định 20 request/phút) và `RATELIMIT_AI` (mặc định "
    "10 request/phút). Bộ đếm lưu trong bộ nhớ, đủ cho mô hình một instance; "
    "khi mở rộng nhiều instance có thể chuyển sang lưu trên Redis.")


add_h2(doc, "2.5. Spring Data JPA và Hibernate ORM")
add_h3(doc, "2.5.1. JPA và Hibernate")
add_para(doc,
    "Java Persistence API (JPA) là chuẩn của Java EE định nghĩa cách ánh xạ "
    "các đối tượng Java sang bảng cơ sở dữ liệu quan hệ (Object–Relational "
    "Mapping – ORM). JPA chỉ là đặc tả, không phải implementation. Hibernate "
    "là implementation phổ biến nhất của JPA, ra mắt từ năm 2001 và đến nay "
    "vẫn được sử dụng rộng rãi.")
add_para(doc,
    "Với JPA/Hibernate, lập trình viên không cần viết SQL thủ công cho các "
    "thao tác CRUD đơn giản. Thay vào đó, họ định nghĩa các Entity class "
    "Java (chú thích bằng @Entity, @Table, @Column, @Id...) và Hibernate "
    "sẽ tự sinh ra các câu lệnh SQL tương ứng khi thực thi.")
add_h3(doc, "2.5.2. Spring Data JPA")
add_para(doc,
    "Spring Data JPA xây dựng trên nền JPA, cung cấp tầng trừu tượng hóa "
    "Repository giúp đơn giản hóa hơn nữa việc truy cập dữ liệu. Lập trình "
    "viên chỉ cần khai báo một interface kế thừa JpaRepository<Entity, ID> "
    "là tự động có các phương thức findAll(), findById(), save(), delete(). "
    "Có thể bổ sung các query tùy chỉnh bằng cách đặt tên method theo quy "
    "ước (findByUsername, findByDepartmentAndStatus) hoặc sử dụng @Query.")
add_h3(doc, "2.5.3. Các annotation JPA quan trọng")
for s in [
    "@Entity: Đánh dấu class là một entity, được ánh xạ sang bảng database.",
    "@Table(name=\"users\"): Chỉ định tên bảng tương ứng.",
    "@Id, @GeneratedValue: Đánh dấu trường khóa chính và cách sinh giá trị (IDENTITY, SEQUENCE, AUTO).",
    "@Column: Cấu hình cột (name, nullable, unique, length).",
    "@OneToMany, @ManyToOne, @OneToOne, @ManyToMany: Khai báo các loại quan hệ.",
    "@JoinColumn: Chỉ định cột khóa ngoại trong quan hệ.",
    "@FetchType.LAZY / EAGER: Cách tải các quan hệ – lười (khi cần mới truy vấn) hay nóng (truy vấn ngay).",
    "@Transactional: Đánh dấu một method là transaction – nhiều câu lệnh SQL được thực hiện trong một transaction duy nhất.",
]:
    add_bullet(doc, s)
add_h3(doc, "2.5.4. Vấn đề N+1 và cách giải quyết")
add_para(doc,
    "Một vấn đề kinh điển khi dùng ORM là truy vấn N+1: khi load danh sách N "
    "đối tượng cha, mỗi đối tượng cha lại trigger một query để load các đối "
    "tượng con (LAZY loading). Kết quả là 1+N câu query – ảnh hưởng nghiêm "
    "trọng đến hiệu năng. Để giải quyết, có thể dùng JOIN FETCH trong JPQL, "
    "@EntityGraph để định nghĩa graph load, hoặc Hibernate's batch fetching. "
    "Trong AiSuggestionService của đề tài, vấn đề N+1 được giải quyết bằng "
    "cách gọi `findByAssignedToEmployeeIdIn(ids)` và "
    "`findByEmployeeEmployeeIdInAndDateBetween(...)` chỉ một lần — gom toàn "
    "bộ task và attendance của tất cả nhân viên trong vài câu query, sau đó "
    "groupBy ở tầng Java.")


add_h2(doc, "2.6. Hệ quản trị cơ sở dữ liệu PostgreSQL 16")
add_h3(doc, "2.6.1. Giới thiệu PostgreSQL")
add_para(doc,
    "PostgreSQL (thường gọi tắt là Postgres) là hệ quản trị cơ sở dữ liệu "
    "quan hệ – đối tượng mã nguồn mở, phát triển từ năm 1986 tại Đại học "
    "California, Berkeley. PostgreSQL nổi tiếng về độ ổn định, tuân thủ "
    "chuẩn SQL nghiêm ngặt, hỗ trợ nhiều tính năng nâng cao như JSON/JSONB, "
    "Full-text search, GIS (qua PostGIS), Materialized View, Common Table "
    "Expression (CTE) và Window Function.")
add_para(doc,
    "Phiên bản PostgreSQL 16 được phát hành tháng 9/2023, mang lại nhiều "
    "cải tiến: logical replication song song, parallel query nâng cao, hỗ "
    "trợ SQL/JSON path expressions, cải thiện hiệu năng cho khối lượng "
    "công việc OLTP và OLAP.")
add_h3(doc, "2.6.2. So sánh PostgreSQL với các RDBMS khác")
add_table(
    doc,
    headers=["Tiêu chí", "PostgreSQL", "MySQL", "SQL Server"],
    rows=[
        ("License", "Open source (PostgreSQL License)",
         "Open source (GPL) + Commercial", "Commercial (Microsoft)"),
        ("Tuân thủ SQL", "Rất nghiêm ngặt", "Vừa phải", "Nghiêm ngặt"),
        ("Kiểu JSON", "JSONB (binary, index được)",
         "JSON (text-based)", "JSON (text-based)"),
        ("MVCC", "Có (mặc định)", "Có (InnoDB)", "Có"),
        ("Stored procedure", "PL/pgSQL, PL/Python...",
         "MySQL stored procedure", "T-SQL"),
        ("Mở rộng", "Rất nhiều extension", "Hạn chế", "CLR, R, Python"),
        ("Performance OLTP", "Tốt", "Rất tốt", "Rất tốt"),
        ("Performance OLAP", "Rất tốt", "Khá", "Rất tốt"),
    ],
    col_widths=[3.5, 4.5, 4.0, 3.5],
)
add_caption(doc, "Bảng 2.3: So sánh PostgreSQL với một số RDBMS khác", kind="table")
add_h3(doc, "2.6.3. Vai trò PostgreSQL trong hệ thống")
add_para(doc,
    "PostgreSQL 16 trong đề tài đóng vai trò là cơ sở dữ liệu chính, lưu "
    "trữ toàn bộ dữ liệu nghiệp vụ: thông tin user, employee, project, "
    "task, attendance, suggestion. Hibernate được cấu hình dialect "
    "PostgreSQLDialect để sinh SQL phù hợp. Quy tắc đặt tên bảng và cột "
    "theo snake_case theo chuẩn PostgreSQL, ánh xạ qua @Column(name=...) "
    "với các Entity sử dụng camelCase của Java.")
add_para(doc,
    "Connection pooling được quản lý bởi HikariCP (mặc định của Spring "
    "Boot) với cấu hình maximum-pool-size=10 cho môi trường phát triển. "
    "Migration database có thể được thực hiện thủ công bằng script SQL "
    "đặt trong thư mục backend/src/main/resources/db/init.sql, được "
    "Docker Compose mount khi khởi tạo container PostgreSQL.")


add_h2(doc, "2.7. Redis và cơ chế caching")
add_h3(doc, "2.7.1. Tổng quan Redis")
add_para(doc,
    "Redis (REmote DIctionary Server) là một kho dữ liệu key–value chạy "
    "trong bộ nhớ (in-memory), ra mắt năm 2009 bởi Salvatore Sanfilippo. "
    "Redis hỗ trợ nhiều kiểu dữ liệu nâng cao: String, List, Hash, Set, "
    "Sorted Set, Bitmap, HyperLogLog, Stream, Geospatial. Redis có thể "
    "persist dữ liệu xuống đĩa qua hai cơ chế RDB (snapshot) và AOF "
    "(append-only file).")
add_para(doc,
    "Redis 7 (phát hành 2022) bổ sung tính năng Functions (kế nhiệm Lua "
    "scripting), Sharded Pub/Sub, cải thiện ACL và Cluster mode.")
add_h3(doc, "2.7.2. So sánh Redis với các giải pháp caching khác")
add_table(
    doc,
    headers=["Tiêu chí", "Redis", "Memcached", "Caffeine"],
    rows=[
        ("Loại", "Remote in-memory store", "Remote in-memory store", "Local (in-JVM) cache"),
        ("Kiểu dữ liệu", "Phong phú (10+)", "Chỉ string", "Chỉ object Java"),
        ("Persistence", "Có (RDB/AOF)", "Không", "Không"),
        ("Cluster", "Có (Redis Cluster, Sentinel)", "Hỗ trợ partition cơ bản", "Không"),
        ("Replication", "Có (master-replica)", "Không", "Không"),
        ("Sử dụng", "Cache + queue + pub/sub", "Cache đơn thuần", "Cache local trong app"),
    ],
    col_widths=[3.0, 4.5, 4.0, 4.0],
)
add_caption(doc, "Bảng 2.4: So sánh Redis với các giải pháp cache khác", kind="table")
add_h3(doc, "2.7.3. Vai trò Redis trong hệ thống")
add_para(doc,
    "Trong đề tài, Redis 7 đóng vai trò là tầng cache cho backend Spring "
    "Boot. Cấu hình spring.cache.type=redis được thiết lập trong "
    "application.yml. Module AiSuggestionService dùng @Cacheable với "
    "key dựa trên `taskId` (hoặc tiêu đề task), TTL 5 phút. Khi quản lý "
    "yêu cầu gợi ý cho cùng một task, kết quả được lấy từ cache thay vì "
    "gọi Gemini lần nữa – giảm thời gian phản hồi từ ~1500ms xuống ~5ms "
    "và tiết kiệm chi phí API.")
add_para(doc,
    "Pattern caching được sử dụng là cache-aside (look-aside): trước "
    "khi truy vấn database, kiểm tra cache; nếu cache hit thì trả về, "
    "nếu cache miss thì truy vấn database, lưu kết quả vào cache rồi "
    "trả về. Khi dữ liệu thay đổi (employee/task được thêm/sửa/xóa), "
    "@CacheEvict được dùng để xóa cache liên quan, đảm bảo tính nhất "
    "quán của dữ liệu.")


add_h2(doc, "2.8. ReactJS, Vite và kiến trúc SPA")
add_h3(doc, "2.8.1. ReactJS")
add_para(doc,
    "React (hay ReactJS) là thư viện JavaScript do Meta (Facebook) phát "
    "triển và mã nguồn mở từ năm 2013. React cho phép xây dựng giao diện "
    "người dùng theo mô hình component-based: mỗi phần của UI là một "
    "component có thể tái sử dụng được, kết hợp lại thành cây component. "
    "React sử dụng Virtual DOM (cây ảo trong bộ nhớ) để tối ưu hóa việc "
    "cập nhật giao diện thật chỉ ở những phần thay đổi.")
add_para(doc,
    "React 18 (phát hành 2022) giới thiệu nhiều tính năng quan trọng: "
    "Automatic Batching, Concurrent Rendering, Suspense for data fetching, "
    "useTransition, useDeferredValue, hỗ trợ Server Components.")
add_h3(doc, "2.8.2. React Hooks")
add_para(doc,
    "React Hooks được giới thiệu từ React 16.8 (2019), cho phép sử dụng "
    "state và các React feature khác trong functional component. Các "
    "hook cơ bản:")
for s in [
    "useState: Khai báo state trong functional component.",
    "useEffect: Thực hiện side effect (gọi API, đăng ký sự kiện, ...) sau khi component render.",
    "useContext: Truy cập giá trị của một React Context.",
    "useRef: Lưu giá trị mutable không gây re-render hoặc reference DOM element.",
    "useMemo, useCallback: Memoize giá trị / hàm để tránh tính toán lại không cần thiết.",
    "useReducer: Quản lý state phức tạp theo mô hình reducer (giống Redux thu nhỏ).",
]:
    add_bullet(doc, s)
add_h3(doc, "2.8.3. Vite")
add_para(doc,
    "Vite là build tool thế hệ mới do Evan You (cha đẻ Vue.js) phát triển, "
    "ra mắt 2020. Khác với các build tool truyền thống như Webpack (đóng "
    "gói toàn bộ ứng dụng trước khi chạy), Vite tận dụng ES Modules native "
    "của trình duyệt: trong môi trường phát triển, code được serve trực "
    "tiếp dưới dạng ESM. Điều này giúp khởi động dev server gần như tức "
    "thời, Hot Module Replacement (HMR) nhanh hơn nhiều lần. Khi build "
    "production, Vite dùng Rollup để bundle.")
add_h3(doc, "2.8.4. Kiến trúc Single Page Application (SPA)")
add_para(doc,
    "SPA là kiến trúc web trong đó toàn bộ ứng dụng được tải một lần đầu "
    "tiên dưới dạng HTML/CSS/JS, sau đó các thao tác chuyển trang được "
    "thực hiện bằng JavaScript mà không cần tải lại trang. Việc chuyển "
    "trang nhanh, mượt như ứng dụng desktop. React Router v6 đóng vai "
    "trò quản lý URL trong SPA. Trong đề tài, toàn bộ frontend web là "
    "một SPA gồm 10 trang chính, sử dụng client-side routing.")


add_h2(doc, "2.9. Tailwind CSS và Chart.js")
add_h3(doc, "2.9.1. Tailwind CSS")
add_para(doc,
    "Tailwind CSS là framework CSS theo hướng utility-first, ra mắt năm "
    "2017 bởi Adam Wathan. Khác với các framework UI truyền thống "
    "(Bootstrap, Material UI) cung cấp các component đã được style sẵn, "
    "Tailwind cung cấp một tập hợp các class tiện ích (utility classes) "
    "ở cấp thấp như flex, p-4 (padding 1rem), bg-blue-500, "
    "rounded-lg, text-xl. Lập trình viên kết hợp các class này trực "
    "tiếp trong JSX/HTML để tạo ra giao diện tùy chỉnh.")
add_para(doc,
    "Ưu điểm của Tailwind: (1) không cần đặt tên class CSS – một bài toán "
    "khó trong các dự án lớn; (2) tránh trùng lặp CSS toàn cục; "
    "(3) build production nhỏ nhờ PurgeCSS loại bỏ các class không dùng; "
    "(4) dễ nhất quán thiết kế khi cả team dùng cùng một bộ utility.")
add_h3(doc, "2.9.2. Chart.js và react-chartjs-2")
add_para(doc,
    "Chart.js là thư viện vẽ biểu đồ JavaScript phổ biến, hỗ trợ nhiều "
    "loại biểu đồ: line, bar, pie, doughnut, radar, polar area, bubble, "
    "scatter. Chart.js sử dụng HTML5 Canvas để render, đảm bảo hiệu năng "
    "tốt với lượng dữ liệu lớn. react-chartjs-2 là wrapper React của "
    "Chart.js, giúp tích hợp dễ dàng vào component React. Trong đề tài, "
    "Chart.js được sử dụng ở trang Dashboard để vẽ biểu đồ cột (số task "
    "theo trạng thái), biểu đồ tròn (phân bổ task theo nhân viên) và "
    "biểu đồ đường (tiến độ hoàn thành theo tuần).")


add_h2(doc, "2.10. Flutter và ngôn ngữ Dart")
add_h3(doc, "2.10.1. Ngôn ngữ Dart")
add_para(doc,
    "Dart là ngôn ngữ lập trình do Google phát triển và phát hành lần "
    "đầu năm 2011. Dart có cú pháp tương tự Java/C# nhưng được tối ưu "
    "cho việc xây dựng giao diện người dùng. Dart hỗ trợ cả hai chế độ "
    "biên dịch: JIT (Just-In-Time) cho phát triển (hot reload) và AOT "
    "(Ahead-Of-Time) cho production. Từ Dart 2.12, null safety trở thành "
    "tính năng mặc định, giúp loại bỏ một trong những lỗi phổ biến nhất "
    "trong lập trình – null reference exception.")
add_h3(doc, "2.10.2. Flutter framework")
add_para(doc,
    "Flutter là framework UI đa nền tảng do Google phát triển, ra mắt "
    "phiên bản 1.0 năm 2018. Flutter cho phép viết một codebase Dart "
    "duy nhất, biên dịch ra ứng dụng native chạy trên iOS, Android, "
    "Web, Windows, macOS và Linux. Khác với React Native (dùng cầu "
    "JavaScript–native), Flutter có engine render riêng (Skia/Impeller) "
    "vẽ trực tiếp các widget lên màn hình – cho phép hiệu năng gần như "
    "native và giao diện đồng nhất 100% giữa các nền tảng.")
add_h3(doc, "2.10.3. Widget tree và state management")
add_para(doc,
    "Triết lý của Flutter là “mọi thứ đều là widget” – từ một nút bấm, "
    "đoạn chữ, layout, đến cả app. Widget được tổ chức thành cây widget "
    "tree. Flutter phân biệt StatelessWidget (không có state, chỉ "
    "render dựa trên props) và StatefulWidget (có state riêng, có thể "
    "thay đổi và trigger re-render). Có nhiều giải pháp quản lý state "
    "phức tạp hơn: setState, Provider, Riverpod, Bloc, GetX. Trong đề "
    "tài sử dụng Provider – một giải pháp đơn giản, được Flutter team "
    "khuyến nghị cho dự án vừa và nhỏ.")
add_h3(doc, "2.10.4. Gọi HTTP API từ Flutter")
add_para(doc,
    "Flutter có hai package phổ biến để gọi HTTP: http (chính thức) và "
    "dio (cộng đồng). Trong đề tài sử dụng dio vì hỗ trợ interceptor "
    "(dễ gắn JWT vào mọi request), retry policy, transform request/response, "
    "và xử lý lỗi tốt hơn. Một DioClient class được tạo riêng làm "
    "singleton, được tất cả các service gọi qua.")


add_h2(doc, "2.11. Docker và Docker Compose")
add_h3(doc, "2.11.1. Container hóa với Docker")
add_para(doc,
    "Docker là nền tảng container hóa phát hành năm 2013, đã cách mạng "
    "hóa cách triển khai phần mềm. Container là một đơn vị đóng gói "
    "phần mềm bao gồm ứng dụng và toàn bộ dependency (thư viện, runtime, "
    "biến môi trường) chạy trên một image. Khác với máy ảo (VM) đóng "
    "gói cả OS, container chia sẻ kernel với host nên rất nhẹ – một "
    "container Linux thường chỉ vài chục MB.")
add_para(doc,
    "Lợi ích của Docker: (1) “build once, run anywhere” – cùng image "
    "chạy giống nhau trên dev, staging, production; (2) cô lập ứng dụng – "
    "tránh xung đột thư viện; (3) dễ scale ngang – chạy nhiều container "
    "song song; (4) hỗ trợ CI/CD – tích hợp với Jenkins, GitHub Actions, "
    "GitLab CI.")
add_h3(doc, "2.11.2. Dockerfile")
add_para(doc,
    "Dockerfile là file text mô tả các bước xây dựng image. Một "
    "Dockerfile tiêu biểu gồm: FROM (image gốc, ví dụ openjdk:17-slim), "
    "WORKDIR, COPY, RUN (cài đặt thêm gì đó), EXPOSE (mở port), CMD "
    "hoặc ENTRYPOINT (lệnh chạy khi container start). Trong đề tài có 3 "
    "Dockerfile: cho backend (multi-stage build – stage 1 build bằng "
    "Maven, stage 2 chạy bằng Eclipse Temurin), cho frontend (build bằng "
    "node, serve bằng nginx), và cho mobile thì không cần Docker.")
add_h3(doc, "2.11.3. Docker Compose")
add_para(doc,
    "Docker Compose là công cụ định nghĩa và quản lý ứng dụng multi-container "
    "qua một file YAML duy nhất (docker-compose.yml). File này khai báo "
    "các service (backend, frontend, postgres, redis), network, volume, "
    "biến môi trường, dependency giữa các service. Chỉ cần một câu lệnh "
    "docker-compose up -d là toàn bộ hệ thống khởi động.")


add_h2(doc, "2.12. Gemini API và tích hợp LLM vào ứng dụng nội bộ")
add_h3(doc, "2.12.1. Large Language Model (LLM)")
add_para(doc,
    "Large Language Model (LLM) là loại mô hình học sâu được huấn luyện "
    "trên khối lượng văn bản khổng lồ, có khả năng hiểu và sinh văn bản "
    "tự nhiên. Các LLM nổi tiếng gồm: Gemini của Google, GPT của "
    "OpenAI, Claude của Anthropic, Llama của Meta, Mistral. Phía sau "
    "các LLM là kiến trúc Transformer (Vaswani et al., 2017) với cơ chế "
    "self-attention. LLM mở ra khả năng xây dựng các ứng dụng AI mà "
    "trước đây tưởng chừng rất khó: chatbot, dịch thuật, tóm tắt văn "
    "bản, code generation, semantic search...")
add_h3(doc, "2.12.2. Google Gemini API")
add_para(doc,
    "Google cung cấp Gemini API — giao diện HTTP để gọi các mô hình "
    "Gemini từ ứng dụng bên thứ ba. Endpoint chính là POST "
    "https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent. "
    "Request gửi đi gồm: contents (mảng các phần nội dung, mỗi phần chứa "
    "parts với trường text), generationConfig (temperature điều khiển "
    "mức ngẫu nhiên, responseMimeType ép định dạng đầu ra — ví dụ "
    "application/json để model trả JSON thuần). Response trả về "
    "candidates, trong đó candidates[0].content.parts[0].text là nội "
    "dung model sinh ra; trường promptFeedback.blockReason cho biết "
    "lý do nếu prompt bị bộ lọc an toàn chặn. API key có thể truyền "
    "qua tham số truy vấn ?key=<API_KEY> hoặc header x-goog-api-key; "
    "đề tài dùng header để key không xuất hiện trong URL và log.")
add_h3(doc, "2.12.3. Mô hình gemini-2.5-flash")
add_para(doc,
    "gemini-2.5-flash là phiên bản nhanh, nhẹ trong họ mô hình Gemini "
    "2.0 của Google, tối ưu cho độ trễ thấp và chi phí thấp. Mô hình "
    "duy trì khả năng hiểu ngữ cảnh tiếng Việt tốt và xử lý được các "
    "tác vụ suy luận tầm trung. Lý do chọn gemini-2.5-flash trong đề "
    "tài: có gói miễn phí (free tier) hào phóng ~1500 request/ngày, "
    "không cần khai báo thẻ thanh toán — phù hợp với một đồ án; tốc độ "
    "phản hồi nhanh; đủ thông minh cho tác vụ xếp hạng nhân viên và "
    "sinh reasoning bằng tiếng Việt.")
add_h3(doc, "2.12.4. Tích hợp Gemini vào hệ thống")
add_para(doc,
    "Trong đề tài, Gemini được gọi trực tiếp từ `AiSuggestionService` "
    "trong backend Spring Boot, dùng `RestClient` của Spring 6. Service "
    "này build prompt tiếng Việt bao gồm thông tin task + danh sách "
    "nhân viên kèm số liệu thô (lịch sử task, chấm công, kỹ năng), gửi "
    "lên endpoint `/v1beta/models/gemini-2.5-flash:generateContent` của Gemini API, "
    "nhận về JSON mảng `[{employeeId, rank, reasoning}]`. API key được "
    "đọc từ biến môi trường `GEMINI_API_KEY` (qua file `.env`), không "
    "bao giờ commit lên git.")
add_para(doc,
    "Khía cạnh bảo mật khi tích hợp LLM: (1) API key chỉ tồn tại ở "
    "backend, không expose ra frontend hay mobile; (2) prompt được "
    "kiểm tra/đánh dấu để giảm rủi ro prompt injection; (3) hạn mức "
    "sử dụng được giám sát qua Google AI Studio để không vượt quota "
    "free tier.")


add_h2(doc, "2.13. Mô hình MVC và kiến trúc phân tầng")
add_h3(doc, "2.13.1. Mô hình MVC")
add_para(doc,
    "Model–View–Controller (MVC) là một mẫu thiết kế (design pattern) "
    "phân chia ứng dụng thành ba thành phần: Model (dữ liệu và logic "
    "nghiệp vụ), View (giao diện hiển thị), Controller (điều phối giữa "
    "Model và View). MVC giúp tách biệt mối quan tâm (separation of "
    "concerns), dễ kiểm thử và bảo trì.")
add_h3(doc, "2.13.2. Kiến trúc phân tầng trong backend Spring Boot")
add_para(doc,
    "Backend của đề tài tổ chức theo 4 tầng:")
for s in [
    "Tầng Controller: Nhận HTTP request, validate đầu vào, gọi service tương ứng, format kết quả và trả về response. Các class này không chứa business logic.",
    "Tầng Service: Chứa toàn bộ business logic – ràng buộc nghiệp vụ, tính toán, gọi các repository và service khác. Tầng service không quan tâm đến HTTP.",
    "Tầng Repository: Tiếp giáp với database qua JPA. Mỗi repository tương ứng với một entity.",
    "Tầng Entity/Model: Đại diện cho dữ liệu nghiệp vụ – mapping 1-1 với bảng database.",
]:
    add_bullet(doc, s)
add_para(doc,
    "Ngoài ra còn các thành phần ngang hàng: DTO (Data Transfer Object – "
    "đối tượng truyền dữ liệu giữa các tầng, đặc biệt giữa Controller và "
    "Client), Mapper (chuyển đổi Entity ↔ DTO), Exception (xử lý lỗi), "
    "Config (cấu hình các bean).")


add_h2(doc, "2.14. Phương pháp phát triển phần mềm Agile")
add_h3(doc, "2.14.1. Triết lý Agile")
add_para(doc,
    "Agile là một triết lý phát triển phần mềm được công bố trong "
    "Tuyên ngôn Agile (Agile Manifesto, 2001) với 4 giá trị cốt lõi: "
    "(1) cá nhân và sự tương tác trên quy trình và công cụ; "
    "(2) phần mềm chạy được trên tài liệu chi tiết; "
    "(3) hợp tác với khách hàng trên đàm phán hợp đồng; "
    "(4) đáp ứng thay đổi trên việc tuân thủ kế hoạch.")
add_h3(doc, "2.14.2. Scrum framework")
add_para(doc,
    "Scrum là framework Agile phổ biến nhất, gồm các thành phần: "
    "Product Owner (đại diện cho khách hàng), Scrum Master (hỗ trợ "
    "team), Development Team (3–9 người thực hiện), Product Backlog "
    "(danh sách yêu cầu được sắp xếp ưu tiên), Sprint Backlog (các "
    "task được chọn để thực hiện trong sprint), Increment (sản phẩm "
    "có thể chạy được sau sprint). Mỗi sprint kéo dài 1–4 tuần với "
    "các sự kiện: Sprint Planning, Daily Standup, Sprint Review, "
    "Sprint Retrospective.")
add_h3(doc, "2.14.3. Áp dụng Agile trong đề tài")
add_para(doc,
    "Trong đề tài, một mình sinh viên đóng cả 3 vai trò Product Owner, "
    "Scrum Master và Developer. Giảng viên hướng dẫn đóng vai trò "
    "khách hàng – nhận demo và đưa feedback sau mỗi sprint 2 tuần. "
    "Kế hoạch 12 tuần được chia thành 6 sprint, mỗi sprint hoàn thành "
    "một nhóm chức năng và demo cho thầy hướng dẫn để nhận góp ý.")


# ==================================================================
# CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
# ==================================================================
add_h1(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

add_h2(doc, "3.1. Khảo sát hiện trạng")
add_h3(doc, "3.1.1. Nguồn dữ liệu khảo sát")
add_para(doc,
    "Để hiểu rõ thực trạng quản lý công việc tại các doanh nghiệp nhỏ "
    "đa ngành ở Việt Nam, đề tài tiến hành khảo sát dựa trên các nguồn "
    "dữ liệu thứ cấp đã được công bố:")
for s in [
    "Số liệu thống kê chính thức: báo cáo tình hình đăng ký doanh nghiệp và các công bố của Bộ Kế hoạch và Đầu tư, phản ánh quy mô và đặc điểm của khu vực doanh nghiệp nhỏ và vừa.",
    "Khảo sát ngành: kết quả khảo sát của Phòng Thương mại và Công nghiệp Việt Nam (VCCI) về mức độ ứng dụng phần mềm quản lý trong doanh nghiệp nhỏ và vừa.",
    "Bài viết và báo cáo chuyên ngành về thực trạng chuyển đổi số của doanh nghiệp nhỏ tại Việt Nam, đăng trên các tạp chí và báo điện tử uy tín.",
]:
    add_bullet(doc, s)
add_h3(doc, "3.1.2. Phương pháp khảo sát")
add_para(doc,
    "Đề tài sử dụng phương pháp nghiên cứu tài liệu thứ cấp (desk "
    "research): thu thập, tổng hợp và phân tích các số liệu, báo cáo nêu "
    "trên để rút ra đặc điểm chung về cách thức quản lý công việc tại "
    "doanh nghiệp nhỏ. Bên cạnh đó, đề tài so sánh trực tiếp các công cụ "
    "quản lý công việc hiện có (bảng tính Excel, sổ ghi chép, Trello, "
    "phần mềm ERP) trên các tiêu chí: chi phí, mức độ trực quan, khả năng "
    "hỗ trợ nghiệp vụ nhân sự – chấm công và độ phù hợp với doanh nghiệp "
    "nhỏ tại Việt Nam.")
add_h3(doc, "3.1.3. Kết quả khảo sát")
add_para(doc, "Tổng hợp các phát hiện chính từ quá trình khảo sát:")
for s in [
    "Phần lớn doanh nghiệp nhỏ vẫn dùng bảng tính Excel kết hợp Zalo/Messenger để giao việc; rất ít doanh nghiệp sử dụng phần mềm quản lý công việc chuyên dụng một cách nhất quán.",
    "Người quản lý khó nắm bắt khối lượng công việc thực tế của từng nhân viên, dẫn đến việc phân công thiếu cân đối.",
    "Tại các doanh nghiệp nhỏ đa ngành, nhân viên thường kiêm nhiệm nhiều kỹ năng (multi-skill) và bị xáo trộn giữa các dự án, khiến việc bố trí nhân sự trở nên phức tạp.",
    "Việc chấm công chủ yếu được thực hiện thủ công bằng giấy hoặc Excel, tốn thời gian tổng hợp vào cuối tháng.",
    "Doanh nghiệp nhỏ quan tâm đến các tính năng ứng dụng AI hỗ trợ ra quyết định, với điều kiện chi phí thấp và dễ sử dụng.",
    "Yêu cầu phổ biến đối với một phần mềm quản lý: dễ sử dụng (không cần đào tạo), chạy được trên cả web và điện thoại, hỗ trợ tiếng Việt, tốc độ phản hồi nhanh.",
]:
    add_bullet(doc, s)
add_table(
    doc,
    headers=["Phương pháp hiện tại", "Ưu điểm", "Nhược điểm"],
    rows=[
        ("Excel + Zalo/Messenger",
         "Quen thuộc, không tốn chi phí, dễ bắt đầu.",
         "Dữ liệu phân tán, không theo dõi được tiến độ realtime, khó tổng hợp báo cáo."),
        ("Sổ ghi chép tay",
         "Đơn giản, không phụ thuộc thiết bị.",
         "Khó tra cứu, dễ mất mát, không chia sẻ được."),
        ("Trello miễn phí",
         "Trực quan kéo thả, có app mobile.",
         "Không có module nhân sự, chấm công; tiếng Anh có rào cản với nhân viên lớn tuổi."),
        ("Phần mềm ERP",
         "Tích hợp đầy đủ, có hỗ trợ chuyên gia.",
         "Chi phí cao (vài trăm triệu), thời gian triển khai dài (6–12 tháng), khó tùy biến."),
    ],
    col_widths=[3.8, 5.5, 6.2],
)
add_caption(doc, "Bảng 3.1: Bảng phân tích ưu/nhược điểm của phương pháp quản lý hiện hành", kind="table")


add_h2(doc, "3.2. Yêu cầu chức năng")
add_para(doc,
    "Dựa trên kết quả khảo sát, hệ thống được thiết kế đáp ứng 14 yêu "
    "cầu chức năng, chia thành 6 nhóm chính:")
add_table(
    doc,
    headers=["Mã", "Tên chức năng", "Mô tả", "Ưu tiên"],
    rows=[
        ("YC-01", "Đăng ký tài khoản",
         "Người dùng tạo tài khoản với username, email, password.", "Cao"),
        ("YC-02", "Đăng nhập",
         "Xác thực username/password, trả về JWT token.", "Cao"),
        ("YC-03", "Đăng xuất",
         "Xóa token phía client.", "Cao"),
        ("YC-04", "Đổi mật khẩu",
         "Người dùng đã đăng nhập đổi mật khẩu của mình.", "Trung bình"),
        ("YC-05", "Quản lý nhân viên (CRUD)",
         "Thêm/sửa/xóa/xem nhân viên kèm thông tin liên hệ và phòng ban.", "Cao"),
        ("YC-06", "Quản lý kỹ năng",
         "Gán hoặc xóa kỹ năng cho nhân viên.", "Cao"),
        ("YC-07", "Quản lý dự án (CRUD)",
         "Tạo dự án, đặt trạng thái, gán nhân viên.", "Cao"),
        ("YC-08", "Quản lý công việc (CRUD)",
         "Tạo task, gán nhân viên, theo dõi tiến độ, mức ưu tiên, deadline.", "Cao"),
        ("YC-09", "Chấm công ngày",
         "Ghi nhận trạng thái PRESENT/ABSENT/LATE, giờ vào/ra.", "Cao"),
        ("YC-10", "Xem lịch sử chấm công",
         "Lọc theo nhân viên và khoảng thời gian.", "Trung bình"),
        ("YC-11", "AI gợi ý nhân viên",
         "Nhập danh sách kỹ năng → trả top 5 nhân viên phù hợp + tóm tắt Gemini.",
         "Cao"),
        ("YC-12", "Xem lịch sử gợi ý",
         "Lưu lịch sử các lần gợi ý AI để tham khảo sau.", "Thấp"),
        ("YC-13", "Dashboard tổng quan",
         "Biểu đồ số nhân viên, dự án, task, tỷ lệ task đúng hạn.", "Trung bình"),
        ("YC-14", "Tài liệu API",
         "Swagger UI tự sinh, hỗ trợ developer kiểm thử.", "Thấp"),
    ],
    col_widths=[1.5, 4.0, 8.0, 2.0],
)
add_caption(doc, "Bảng 3.2: Danh sách yêu cầu chức năng của hệ thống", kind="table")


add_h2(doc, "3.3. Yêu cầu phi chức năng")
add_table(
    doc,
    headers=["Mã", "Loại", "Yêu cầu"],
    rows=[
        ("NF-01", "Hiệu năng",
         "API thông thường phản hồi trong 500ms; API AI Gợi ý ≤ 2 giây với 100 nhân viên."),
        ("NF-02", "Bảo mật",
         "Mật khẩu băm BCrypt cost=10; JWT HS256 expires 2h; HTTPS cho production; CORS giới hạn; "
         "rate limiting theo IP chống spam/brute-force."),
        ("NF-03", "Khả dụng",
         "Hệ thống hoạt động 99% thời gian; phục hồi sau lỗi container trong 1 phút."),
        ("NF-04", "Khả năng mở rộng",
         "Kiến trúc cho phép tăng số người dùng đến 1000 mà không cần thiết kế lại."),
        ("NF-05", "Tính nhất quán",
         "Dữ liệu đồng bộ giữa web và mobile do dùng chung backend."),
        ("NF-06", "Tính sử dụng",
         "Giao diện tiếng Việt, responsive trên desktop/mobile, tuân thủ một số tiêu chí WCAG."),
        ("NF-07", "Khả năng bảo trì",
         "Code có cấu trúc tầng rõ ràng; tuân thủ Java Code Conventions và Airbnb React Style Guide."),
        ("NF-08", "Triển khai",
         "Hệ thống đóng gói bằng Docker Compose, cài đặt được trên Linux/Windows/macOS với 1 lệnh."),
    ],
    col_widths=[1.5, 3.0, 11.0],
)
add_caption(doc, "Bảng 3.3: Danh sách yêu cầu phi chức năng", kind="table")


add_h2(doc, "3.4. Sơ đồ Use Case và đặc tả use case")
add_h3(doc, "3.4.1. Sơ đồ Use Case tổng thể")
add_para(doc,
    "Hệ thống có 3 actor chính: User (nhân viên thông thường), Manager "
    "(người quản lý), và Admin (quản trị hệ thống). Sơ đồ Use Case "
    "tổng thể trình bày dưới đây mô tả mối quan hệ giữa các actor và "
    "các use case chính.")
add_para(doc,
    "Hình 3.1 thể hiện sơ đồ Use Case bằng ký pháp UML. Người dùng "
    "(actor) ở bên trái, hệ thống (system boundary) là khung chữ nhật "
    "ở giữa, các use case là các ellipse bên trong khung. Các mũi tên "
    "thể hiện quan hệ giữa actor và use case (association), hoặc giữa "
    "các use case với nhau (include, extend).")
add_uml_image(doc, "use-case-tong-the.png", width_cm=12.0)
add_caption(doc, "Hình 3.1: Sơ đồ Use Case tổng thể của hệ thống", kind="figure")
add_para(doc,
    "Sơ đồ gồm 14 use case. Do nhiều use case thuộc nhóm CRUD (xem, thêm, "
    "sửa, xóa) có cấu trúc tương tự nhau, phần dưới đặc tả chi tiết theo "
    "mẫu Cockburn cho 3 use case tiêu biểu nhất — đại diện cho ba nhóm "
    "chức năng quan trọng: đăng nhập (UC-01), tạo công việc và gán nhân "
    "viên (UC-08), và AI gợi ý nhân viên (UC-11).")
add_h3(doc, "3.4.2. Đặc tả use case UC-01: Đăng nhập")
add_table(
    doc,
    headers=["Mục", "Nội dung"],
    rows=[
        ("Mã use case", "UC-01"),
        ("Tên use case", "Đăng nhập"),
        ("Actor", "User, Manager, Admin"),
        ("Mục tiêu", "Cho phép người dùng đã có tài khoản truy cập hệ thống."),
        ("Tiền điều kiện", "Người dùng đã đăng ký tài khoản; trình duyệt/app hoạt động."),
        ("Hậu điều kiện thành công",
         "Người dùng được điều hướng tới Dashboard; JWT lưu vào localStorage."),
        ("Hậu điều kiện thất bại", "Hiển thị thông báo lỗi tương ứng."),
        ("Luồng chính",
         "1. Người dùng truy cập /login. "
         "2. Người dùng nhập username và password. "
         "3. Hệ thống gửi POST /api/auth/login. "
         "4. Hệ thống xác thực, tạo JWT, trả về client. "
         "5. Client lưu JWT, điều hướng đến /dashboard."),
        ("Luồng phụ – Sai mật khẩu",
         "Tại bước 4, nếu username hoặc password sai → trả HTTP 401, "
         "hiển thị thông báo “Tên đăng nhập hoặc mật khẩu không đúng”."),
        ("Luồng phụ – Tài khoản bị khóa",
         "Nếu user.status = LOCKED → trả HTTP 403, hiển thị thông báo “Tài khoản đã bị khóa”."),
        ("Luồng phụ – Thiếu trường",
         "Nếu username hoặc password rỗng → trả HTTP 400, hiển thị “Vui lòng nhập đầy đủ thông tin”."),
    ],
    col_widths=[3.5, 12.0],
)
add_caption(doc, "Bảng 3.4: Đặc tả use case UC-01: Đăng nhập", kind="table")

add_h3(doc, "3.4.3. Đặc tả use case UC-08: Tạo công việc và gán nhân viên")
add_table(
    doc,
    headers=["Mục", "Nội dung"],
    rows=[
        ("Mã use case", "UC-08"),
        ("Tên use case", "Tạo công việc và gán nhân viên"),
        ("Actor", "Manager"),
        ("Mục tiêu", "Tạo một công việc mới và gán cho nhân viên trong dự án."),
        ("Tiền điều kiện",
         "Người dùng đã đăng nhập với vai trò MANAGER hoặc ADMIN; dự án và nhân viên đã tồn tại."),
        ("Hậu điều kiện thành công",
         "Task được lưu vào DB với trạng thái TODO; nhân viên có thể thấy task trong danh sách."),
        ("Luồng chính",
         "1. Manager vào trang /tasks. "
         "2. Manager nhấn “+ Tạo task”. "
         "3. Modal mở ra với các trường: title, description, projectId, assignedEmployeeId, dueDate, priority. "
         "4. Manager nhập đầy đủ thông tin và nhấn “Lưu”. "
         "5. Hệ thống validate dữ liệu, gọi POST /api/tasks. "
         "6. Backend lưu task vào DB, trả về task DTO. "
         "7. UI cập nhật danh sách, hiển thị toast “Tạo task thành công”."),
        ("Luồng phụ – Thiếu trường bắt buộc",
         "Nếu title rỗng → hiển thị inline error, không gửi request."),
        ("Luồng phụ – Nhân viên không phù hợp",
         "Nếu nhân viên đã có ≥ 5 task IN_PROGRESS → cảnh báo nhưng vẫn cho phép tiếp tục."),
    ],
    col_widths=[3.5, 12.0],
)
add_caption(doc, "Bảng 3.5: Đặc tả use case UC-08: Tạo công việc và gán nhân viên", kind="table")

add_h3(doc, "3.4.4. Đặc tả use case UC-11: AI gợi ý nhân viên")
add_table(
    doc,
    headers=["Mục", "Nội dung"],
    rows=[
        ("Mã use case", "UC-11"),
        ("Tên use case", "AI gợi ý nhân viên phù hợp"),
        ("Actor", "Manager"),
        ("Mục tiêu",
         "Tìm top 5 nhân viên phù hợp nhất cho một công việc dựa trên các tiêu chí khách quan."),
        ("Tiền điều kiện",
         "Manager đã đăng nhập; có ít nhất một nhân viên trong hệ thống."),
        ("Hậu điều kiện thành công",
         "Hệ thống hiển thị tối đa 5 thẻ nhân viên kèm thứ hạng và lý do "
         "(reasoning) do Gemini sinh ra."),
        ("Luồng chính",
         "1. Manager vào trang /ai-suggestions. "
         "2. Nhập tiêu đề + mô tả công việc + kỹ năng yêu cầu (text tự do). "
         "3. Nhấn “Gợi ý ngay”. "
         "4. Frontend gửi POST /api/suggestions/recommend với body {taskTitle, taskDescription, requiredSkills}. "
         "5. Backend kiểm tra cache Redis với key = taskId hoặc tiêu đề. "
         "6. Nếu cache miss: backend gọi `findByAssignedToEmployeeIdIn(...)` và "
         "`findByEmployeeEmployeeIdInAndDateBetween(...)` để gom lịch sử + chấm công của tất cả nhân viên. "
         "7. Backend xây prompt tiếng Việt cho Google Gemini gemini-2.5-flash, gồm thông tin task + danh sách nhân viên "
         "kèm số liệu thô (KHÔNG tính điểm số). "
         "8. Nhận về JSON [{employeeId, rank, reasoning}] từ Gemini, lưu vào cache 5 phút. "
         "9. Trả về client; UI hiển thị 5 thẻ nhân viên kèm thứ hạng và lý do bằng tiếng Việt."),
        ("Luồng phụ – Không có nhân viên",
         "Nếu hệ thống chưa có nhân viên nào → trả về danh sách rỗng, UI hiển thị "
         "“Chưa có nhân viên trong hệ thống, vui lòng thêm trước.”"),
        ("Luồng phụ – Thiếu API key",
         "Nếu chưa cấu hình GEMINI_API_KEY → backend trả HTTP 422 kèm thông báo "
         "“AI suggestion is unavailable: GEMINI_API_KEY is not configured”; UI hiển "
         "thị “Tính năng AI đang không khả dụng, vui lòng thử lại sau.”"),
        ("Luồng phụ – Gemini lỗi",
         "Nếu gọi Gemini thất bại (mạng lỗi, key sai, hết quota) hoặc prompt bị bộ "
         "lọc an toàn chặn → backend trả HTTP 422/500 kèm lý do; UI hiển thị thông "
         "báo lỗi tương ứng."),
    ],
    col_widths=[3.5, 12.0],
)
add_caption(doc, "Bảng 3.6: Đặc tả use case UC-11: AI gợi ý nhân viên", kind="table")


add_h2(doc, "3.5. Thiết kế cơ sở dữ liệu – Sơ đồ ERD")
add_para(doc,
    "Cơ sở dữ liệu được thiết kế gồm 6 bảng chính, đáp ứng đầy đủ các "
    "yêu cầu nghiệp vụ và đảm bảo nguyên tắc chuẩn hóa (3NF). Kỹ năng "
    "nhân viên được lưu dưới dạng cột TEXT do quản lý nhập tự do, "
    "không tách thành bảng riêng. Sơ đồ ERD dưới đây thể hiện các "
    "thực thể (entity) và mối quan hệ.")
add_uml_image(doc, "erd.png", width_cm=13.0)
add_caption(doc, "Hình 3.2: Sơ đồ ERD của hệ thống", kind="figure")


add_h2(doc, "3.6. Mô tả chi tiết các bảng")
add_h3(doc, "3.6.1. Bảng users")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("id", "BIGSERIAL", "PK", "Khóa chính, tự tăng."),
        ("username", "VARCHAR(50)", "NOT NULL, UNIQUE", "Tên đăng nhập."),
        ("email", "VARCHAR(100)", "NOT NULL, UNIQUE", "Địa chỉ email."),
        ("password", "VARCHAR(255)", "NOT NULL", "Mật khẩu băm BCrypt."),
        ("role", "VARCHAR(20)", "NOT NULL", "ADMIN | MANAGER | EMPLOYEE."),
        ("status", "VARCHAR(20)", "DEFAULT 'ACTIVE'", "ACTIVE | LOCKED."),
        ("created_at", "TIMESTAMP", "NOT NULL DEFAULT NOW()", "Thời điểm tạo tài khoản."),
    ],
    col_widths=[3.0, 3.0, 4.0, 5.5],
)
add_caption(doc, "Bảng 3.7: Mô tả bảng users", kind="table")

add_h3(doc, "3.6.2. Bảng employees")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("employee_id", "BIGSERIAL", "PK", "Khóa chính."),
        ("user_id", "BIGINT", "FK → users.id", "Liên kết với tài khoản đăng nhập (nếu có)."),
        ("first_name", "VARCHAR(50)", "NOT NULL", "Họ."),
        ("last_name", "VARCHAR(50)", "NOT NULL", "Tên."),
        ("position", "VARCHAR(50)", "", "Chức danh."),
        ("department", "VARCHAR(50)", "", "Phòng ban / bộ phận."),
        ("employee_group", "VARCHAR(100)", "", "Nhóm/team trực thuộc."),
        ("skills", "TEXT", "", "Kỹ năng (quản lý nhập tự do, ngăn cách bởi dấu phẩy)."),
        ("hired_at", "TIMESTAMP", "", "Thời điểm vào làm."),
    ],
    col_widths=[3.0, 3.0, 4.0, 5.5],
)
add_caption(doc, "Bảng 3.8: Mô tả bảng employees", kind="table")

add_h3(doc, "3.6.3. Bảng projects")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("id", "BIGSERIAL", "PK", "Khóa chính."),
        ("name", "VARCHAR(200)", "NOT NULL", "Tên dự án."),
        ("description", "TEXT", "", "Mô tả chi tiết."),
        ("start_date", "DATE", "", "Ngày bắt đầu."),
        ("end_date", "DATE", "", "Ngày kết thúc dự kiến."),
        ("status", "VARCHAR(20)", "DEFAULT 'PLANNING'",
         "PLANNING | ACTIVE | COMPLETED | CANCELLED."),
        ("manager_id", "BIGINT", "FK → employees.id", "Trưởng dự án."),
    ],
    col_widths=[3.0, 3.0, 4.5, 5.0],
)
add_caption(doc, "Bảng 3.9: Mô tả bảng projects", kind="table")

add_h3(doc, "3.6.4. Bảng tasks")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("task_id", "BIGSERIAL", "PK", "Khóa chính."),
        ("title", "VARCHAR(100)", "NOT NULL", "Tiêu đề công việc."),
        ("description", "TEXT", "", "Mô tả công việc."),
        ("required_skills", "TEXT", "",
         "Kỹ năng yêu cầu (quản lý nhập tự do, dùng làm input cho AI gợi ý)."),
        ("status", "VARCHAR(50)", "DEFAULT 'pending'",
         "pending | in_progress | completed."),
        ("due_date", "DATE", "", "Hạn hoàn thành."),
        ("completed_at", "TIMESTAMP", "", "Thời điểm hoàn thành thực tế."),
        ("project_id", "BIGINT", "FK → projects.id", "Dự án chứa task."),
        ("assigned_to", "BIGINT", "FK → employees.employee_id", "Nhân viên được giao."),
    ],
    col_widths=[3.0, 3.0, 4.5, 5.0],
)
add_caption(doc, "Bảng 3.10: Mô tả bảng tasks", kind="table")

add_h3(doc, "3.6.5. Bảng attendances")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("id", "BIGSERIAL", "PK", "Khóa chính."),
        ("attendance_date", "DATE", "NOT NULL", "Ngày chấm công."),
        ("status", "VARCHAR(20)", "NOT NULL",
         "PRESENT | ABSENT | LATE | LEAVE."),
        ("check_in", "TIME", "", "Giờ vào."),
        ("check_out", "TIME", "", "Giờ ra."),
        ("notes", "TEXT", "", "Ghi chú: nghỉ phép, công tác, v.v."),
        ("employee_id", "BIGINT", "FK → employees.id, NOT NULL", "Nhân viên chấm công."),
    ],
    col_widths=[3.0, 3.0, 4.5, 5.0],
)
add_caption(doc, "Bảng 3.11: Mô tả bảng attendances", kind="table")

add_h3(doc, "3.6.6. Bảng suggestions")
add_table(
    doc,
    headers=["Cột", "Kiểu dữ liệu", "Ràng buộc", "Mô tả"],
    rows=[
        ("id", "BIGSERIAL", "PK", "Khóa chính."),
        ("required_skills", "VARCHAR(500)", "",
         "Danh sách kỹ năng yêu cầu lưu lại từ lần gợi ý."),
        ("task_title", "VARCHAR(200)", "", "Tiêu đề task được gợi ý."),
        ("user_id", "BIGINT", "FK → users.id", "Người tạo gợi ý."),
        ("created_at", "TIMESTAMP", "DEFAULT NOW()", "Thời điểm tạo."),
    ],
    col_widths=[3.0, 3.0, 4.5, 5.0],
)
add_caption(doc, "Bảng 3.12: Mô tả bảng suggestions", kind="table")


add_h2(doc, "3.7. Sơ đồ lớp (Class Diagram)")
add_para(doc,
    "Sơ đồ lớp dưới đây thể hiện các Entity class chính, các Service "
    "tầng nghiệp vụ và mối quan hệ giữa chúng. Đối tượng DTO và "
    "Controller được giản lược để tập trung vào domain model.")
add_uml_image(doc, "class-diagram-thuc-the.png", width_cm=13.0)
add_caption(doc, "Hình 3.3a: Sơ đồ lớp – Entity (Domain Model)", kind="figure")
add_para(doc,
    "Bên cạnh các lớp Entity, hệ thống còn có tầng Controller–Service–"
    "Repository tổ chức theo kiến trúc Spring Boot, được mô tả ở Hình 3.3b "
    "dưới đây. Lớp đặc biệt `AiSuggestionService` gọi tới nhiều repository "
    "và Gemini API thông qua `GeminiClient`.")
add_uml_image(doc, "class-diagram-kien-truc.png", width_cm=15.0)
add_caption(doc, "Hình 3.3b: Sơ đồ lớp – Kiến trúc Controller / Service / Repository",
            kind="figure")


add_h2(doc, "3.8. Sơ đồ tuần tự (Sequence Diagram)")
add_h3(doc, "3.8.1. Sơ đồ tuần tự – Đăng nhập")
add_uml_image(doc, "sequence-dang-nhap.png", width_cm=15.5)
add_caption(doc, "Hình 3.4: Sơ đồ tuần tự – Đăng nhập", kind="figure")

add_h3(doc, "3.8.2. Sơ đồ tuần tự – AI gợi ý nhân viên")
add_uml_image(doc, "sequence-ai-goi-y.png", width_cm=15.5)
add_caption(doc, "Hình 3.5: Sơ đồ tuần tự – AI gợi ý nhân viên", kind="figure")

add_h3(doc, "3.8.3. Sơ đồ tuần tự – Tạo công việc và gán nhân viên")
add_uml_image(doc, "sequence-quan-ly-cong-viec.png", width_cm=15.5)
add_caption(doc, "Hình 3.6: Sơ đồ tuần tự – Tạo công việc và gán nhân viên", kind="figure")


add_h2(doc, "3.9. Sơ đồ hoạt động (Activity Diagram)")
add_h3(doc, "3.9.1. Activity Diagram – Chấm công ngày")
add_uml_image(doc, "activity-cham-cong.png", width_cm=13.5)
add_caption(doc, "Hình 3.7: Sơ đồ hoạt động – Chấm công ngày", kind="figure")

add_h3(doc, "3.9.2. Activity Diagram – AI gợi ý nhân viên")
add_uml_image(doc, "activity-ai-goi-y.png", width_cm=14.0)
add_caption(doc, "Hình 3.7b: Sơ đồ hoạt động – AI gợi ý nhân viên", kind="figure")


add_h2(doc, "3.10. Kiến trúc tổng thể hệ thống")
add_para(doc,
    "Hệ thống được thiết kế theo kiến trúc 3 tầng kết hợp microservice "
    "nhẹ. Tất cả các thành phần được đóng gói trong Docker container và "
    "kết nối qua một mạng Docker bridge nội bộ `taskmgmt_net`.")
add_uml_image(doc, "architecture.png", width_cm=15.5)
add_caption(doc, "Hình 3.8: Sơ đồ kiến trúc tổng thể hệ thống", kind="figure")


add_h2(doc, "3.11. Thiết kế giao diện (Wireframe)")
add_h3(doc, "3.11.1. Wireframe trang Dashboard")
add_code(doc, """
+----------------------------------------------------------------------+
| Logo  Task Manager                                  [User v]  [Bell] |
+--------+-------------------------------------------------------------+
| Side   |  Dashboard                                                  |
| nav    |  +------------+ +------------+ +------------+ +-----------+ |
|        |  | Employees  | | Projects   | | Active Tasks| | Overdue  | |
| - Dash |  |    24      | |     8      | |    47       | |    3     | |
| - Emp  |  +------------+ +------------+ +------------+ +-----------+ |
| - Proj |                                                              |
| - Task |  Tasks by status (bar chart)        Tasks by employee (pie)  |
| - Att  |  +-----------------------+         +--------------------+   |
| - AI   |  |  ███ TODO        12   |         |   Employee A 25%   |   |
|        |  |  ██████ IN_PROG  20   |         |   Employee B 18%   |   |
|        |  |  ███ DONE        15   |         |   Employee C 14%   |   |
|        |  |                       |         |   Others    43%    |   |
|        |  +-----------------------+         +--------------------+   |
+--------+-------------------------------------------------------------+
""")
add_caption(doc, "Hình 3.9: Wireframe trang Dashboard", kind="figure")

add_h3(doc, "3.11.2. Wireframe trang AI Suggestion")
add_code(doc, """
+----------------------------------------------------------------------+
|  AI Gợi ý nhân viên                                                  |
+----------------------------------------------------------------------+
|  Tiêu đề công việc:                                                  |
|  [_______________________________________________________________ ] |
|                                                                      |
|  Mô tả công việc:                                                    |
|  [_______________________________________________________________ ] |
|                                                                      |
|  Kỹ năng yêu cầu (nhập tự do, ngăn cách bởi dấu phẩy):              |
|  [ Java, Spring Boot, PostgreSQL_________________________________ ] |
|                                                                      |
|                              [ Gợi ý ngay ]                          |
+----------------------------------------------------------------------+
|  Top 5 nhân viên phù hợp (do AI xếp hạng):                           |
|                                                                      |
|  #1  Nguyễn Văn A  – Backend Dev (Phòng Phát triển)                  |
|       Lý do: Có kỹ năng Java + Spring Boot đầy đủ, đã hoàn thành    |
|              23/25 task đúng hạn, đi làm 21/22 ngày qua.            |
|                                                                      |
|  #2  Trần Thị B   – Full-stack  (Phòng Phát triển)                  |
|       Lý do: ...                                                     |
|                                                                      |
|  #3 ... #5                                                           |
+----------------------------------------------------------------------+
""")
add_caption(doc, "Hình 3.10: Wireframe trang AI Suggestion", kind="figure")


# ==================================================================
# CHƯƠNG 4: XÂY DỰNG ỨNG DỤNG
# ==================================================================
add_h1(doc, "CHƯƠNG 4. XÂY DỰNG ỨNG DỤNG")

add_h2(doc, "4.1. Môi trường phát triển")
add_para(doc,
    "Để đảm bảo có thể tái lập kết quả của đề tài, môi trường phát "
    "triển và các phiên bản phần mềm sử dụng được mô tả chi tiết trong "
    "bảng dưới đây.")
add_table(
    doc,
    headers=["Phần mềm", "Phiên bản", "Mục đích sử dụng"],
    rows=[
        ("Java JDK (Eclipse Temurin)", "17.0.12", "Biên dịch và chạy backend Spring Boot."),
        ("Apache Maven", "3.9.6", "Quản lý dependency và build backend."),
        ("Spring Boot", "3.5.0", "Framework phát triển backend."),
        ("Node.js", "18.20.4 LTS", "Runtime JavaScript cho frontend."),
        ("npm", "10.8.2", "Quản lý package cho React."),
        ("Vite", "5.4.0", "Build tool cho frontend React."),
        ("React", "18.3.1", "Thư viện UI cho web."),
        ("Tailwind CSS", "3.4.10", "Framework CSS utility-first."),
        ("Flutter SDK", "3.24.0", "Phát triển ứng dụng mobile."),
        ("Dart", "3.5.0", "Ngôn ngữ Flutter."),
        ("PostgreSQL", "16.4", "Cơ sở dữ liệu chính."),
        ("Redis", "7.4", "Cache server."),
        ("Docker Engine", "27.0.3", "Container runtime."),
        ("Docker Compose", "2.29.1", "Orchestration multi-container."),
        ("Git", "2.46.0", "Quản lý mã nguồn."),
        ("IDE Backend", "IntelliJ IDEA 2024.2", "Phát triển Java."),
        ("IDE Frontend", "VS Code 1.92", "Phát triển React/Flutter."),
        ("API Tester", "Postman 11.7", "Kiểm thử API."),
        ("OS phát triển", "Windows 11 + WSL2 Ubuntu 22.04", "Hệ điều hành dev."),
    ],
    col_widths=[5.0, 3.5, 7.0],
)
add_caption(doc, "Bảng 4.1: Yêu cầu phần cứng và phần mềm cần thiết", kind="table")
add_para(doc,
    "Cấu hình máy phát triển: CPU Intel Core i5-12400 (6 core), RAM "
    "16GB, SSD NVMe 512GB. Cấu hình này đủ để chạy đồng thời backend, "
    "frontend dev server, PostgreSQL, Redis và Android emulator. Khi "
    "triển khai production có thể dùng VPS tối thiểu 2 vCPU, 4GB RAM "
    "với Ubuntu Server 22.04 LTS.")


add_h2(doc, "4.2. Cấu trúc thư mục dự án")
add_para(doc,
    "Toàn bộ mã nguồn của hệ thống được tổ chức trong một monorepo "
    "với cấu trúc thư mục như sau, đảm bảo tách biệt rõ ràng giữa "
    "backend, frontend và mobile.")
add_code(doc, """
task-management-system/
├── backend/                          # Spring Boot 3.5
│   ├── src/
│   │   ├── main/
│   │   │   ├── java/com/taskmanagement/
│   │   │   │   ├── TaskManagementApplication.java
│   │   │   │   ├── controller/
│   │   │   │   │   ├── AuthController.java
│   │   │   │   │   ├── EmployeeController.java
│   │   │   │   │   ├── ProjectController.java
│   │   │   │   │   ├── TaskController.java
│   │   │   │   │   ├── AttendanceController.java
│   │   │   │   │   └── SuggestionController.java
│   │   │   │   ├── service/
│   │   │   │   │   ├── UserService.java
│   │   │   │   │   ├── EmployeeService.java
│   │   │   │   │   ├── ProjectService.java
│   │   │   │   │   ├── TaskService.java
│   │   │   │   │   ├── AttendanceService.java
│   │   │   │   │   ├── AiSuggestionService.java
│   │   │   │   │   ├── SuggestionService.java
│   │   │   │   │   └── CurrentUserService.java
│   │   │   │   ├── repository/
│   │   │   │   │   ├── UserRepository.java
│   │   │   │   │   ├── EmployeeRepository.java
│   │   │   │   │   ├── ProjectRepository.java
│   │   │   │   │   ├── TaskRepository.java
│   │   │   │   │   ├── AttendanceRepository.java
│   │   │   │   │   └── SuggestionRepository.java
│   │   │   │   ├── entity/
│   │   │   │   │   ├── User.java
│   │   │   │   │   ├── Employee.java
│   │   │   │   │   ├── Project.java
│   │   │   │   │   ├── Task.java
│   │   │   │   │   ├── Attendance.java
│   │   │   │   │   └── Suggestion.java
│   │   │   │   ├── dto/
│   │   │   │   │   ├── LoginRequest.java
│   │   │   │   │   ├── LoginResponse.java
│   │   │   │   │   ├── EmployeeDTO.java
│   │   │   │   │   ├── TaskDTO.java
│   │   │   │   │   ├── EmployeeScoreDTO.java
│   │   │   │   │   ├── ApiResponse.java
│   │   │   │   │   └── PageResponse.java
│   │   │   │   ├── security/
│   │   │   │   │   ├── JwtTokenProvider.java
│   │   │   │   │   ├── JwtAuthenticationFilter.java
│   │   │   │   │   └── SecurityConfig.java
│   │   │   │   ├── config/
│   │   │   │   │   ├── RedisConfig.java
│   │   │   │   │   └── OpenApiConfig.java
│   │   │   │   └── exception/
│   │   │   │       ├── GlobalExceptionHandler.java
│   │   │   │       ├── ResourceNotFoundException.java
│   │   │   │       └── BadRequestException.java
│   │   │   └── resources/
│   │   │       ├── application.yml
│   │   │       ├── application-prod.yml
│   │   │       └── db/init.sql
│   │   └── test/
│   │       └── java/com/taskmanagement/
│   │           ├── service/AiSuggestionServiceTest.java
│   │           └── controller/AuthControllerTest.java
│   ├── pom.xml
│   └── Dockerfile
├── frontend/                         # React 18 + Vite 5
│   ├── src/
│   │   ├── api/
│   │   │   ├── axiosConfig.js
│   │   │   ├── authApi.js
│   │   │   ├── employeeApi.js
│   │   │   ├── taskApi.js
│   │   │   └── suggestionApi.js
│   │   ├── components/
│   │   │   ├── Layout.jsx
│   │   │   ├── Sidebar.jsx
│   │   │   ├── Header.jsx
│   │   │   ├── EmployeeFormModal.jsx
│   │   │   ├── TaskFormModal.jsx
│   │   │   └── ProtectedRoute.jsx
│   │   ├── context/
│   │   │   └── AuthContext.jsx
│   │   ├── pages/
│   │   │   ├── Login.jsx
│   │   │   ├── Register.jsx
│   │   │   ├── Dashboard.jsx
│   │   │   ├── Employees.jsx
│   │   │   ├── Projects.jsx
│   │   │   ├── Tasks.jsx
│   │   │   ├── Attendance.jsx
│   │   │   └── AiSuggestions.jsx
│   │   ├── App.jsx
│   │   ├── main.jsx
│   │   └── index.css
│   ├── package.json
│   ├── vite.config.js
│   ├── tailwind.config.js
│   └── Dockerfile
├── mobile/                           # Flutter 3.x
│   ├── lib/
│   │   ├── main.dart
│   │   ├── models/
│   │   ├── services/
│   │   ├── screens/
│   │   │   ├── login_screen.dart
│   │   │   ├── dashboard_screen.dart
│   │   │   ├── tasks_screen.dart
│   │   │   ├── attendance_screen.dart
│   │   │   └── ai_suggestions_screen.dart
│   │   └── widgets/
│   └── pubspec.yaml
├── docs/                             # Tài liệu kỹ thuật
│   ├── API_SPECIFICATION.md
│   ├── DATABASE_SCHEMA.md
│   ├── SETUP_GUIDE.md
│   ├── UML_DIAGRAMS.md
│   └── BAO_CAO_DO_AN_CO_SO.docx
├── docker-compose.yml
├── docker-compose.prod.yml
├── Caddyfile
├── .env.example
├── render.yaml
├── DEPLOY.md
├── README.md
└── CLAUDE.md
""")
add_caption(doc, "Hình 4.1: Cấu trúc thư mục backend Spring Boot", kind="figure")


add_h2(doc, "4.3. Triển khai Backend Spring Boot")
add_h3(doc, "4.3.1. Khai báo dependency trong pom.xml")
add_para(doc,
    "File pom.xml khai báo các dependency cần thiết. Spring Boot Starter "
    "đã đóng gói sẵn các thư viện liên quan, chỉ cần khai báo starter là "
    "có ngay toàn bộ tính năng cần thiết.")
add_code(doc, """
<dependencies>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-web</artifactId>
    </dependency>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-data-jpa</artifactId>
    </dependency>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-security</artifactId>
    </dependency>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-data-redis</artifactId>
    </dependency>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-cache</artifactId>
    </dependency>
    <dependency>
        <groupId>org.springframework.boot</groupId>
        <artifactId>spring-boot-starter-validation</artifactId>
    </dependency>
    <dependency>
        <groupId>org.postgresql</groupId>
        <artifactId>postgresql</artifactId>
        <scope>runtime</scope>
    </dependency>
    <dependency>
        <groupId>io.jsonwebtoken</groupId>
        <artifactId>jjwt-api</artifactId>
        <version>0.12.6</version>
    </dependency>
    <dependency>
        <groupId>io.jsonwebtoken</groupId>
        <artifactId>jjwt-impl</artifactId>
        <version>0.12.6</version>
        <scope>runtime</scope>
    </dependency>
    <dependency>
        <groupId>io.jsonwebtoken</groupId>
        <artifactId>jjwt-jackson</artifactId>
        <version>0.12.6</version>
        <scope>runtime</scope>
    </dependency>
    <dependency>
        <groupId>org.springdoc</groupId>
        <artifactId>springdoc-openapi-starter-webmvc-ui</artifactId>
        <version>2.6.0</version>
    </dependency>
    <dependency>
        <groupId>org.projectlombok</groupId>
        <artifactId>lombok</artifactId>
        <optional>true</optional>
    </dependency>
</dependencies>
""")

add_h3(doc, "4.3.2. Cấu hình application.yml")
add_code(doc, """
spring:
  application:
    name: task-management-system
  datasource:
    url: jdbc:postgresql://${DB_HOST:localhost}:5432/${DB_NAME:taskmgmt}
    username: ${DB_USER:postgres}
    password: ${DB_PASS:postgres}
    driver-class-name: org.postgresql.Driver
    hikari:
      maximum-pool-size: 10
      minimum-idle: 2
  jpa:
    hibernate:
      ddl-auto: update
    properties:
      hibernate:
        dialect: org.hibernate.dialect.PostgreSQLDialect
        format_sql: true
        jdbc.batch_size: 25
    show-sql: false
    open-in-view: false
  data:
    redis:
      host: ${REDIS_HOST:localhost}
      port: 6379
  cache:
    type: redis
    redis:
      time-to-live: 300000     # 5 phút

server:
  port: 5000
  servlet:
    encoding:
      charset: UTF-8

jwt:
  secret: ${JWT_SECRET:bXlfc3VwZXJfc2VjcmV0X2tleV9mb3JfaGFvX2Jhb19jYW9fZG9fYW5fY29fc28=}
  expiration: 7200000           # 2h

gemini:
  api:
    key: ${GEMINI_API_KEY:}             # trống thì tính năng AI tắt
    model: ${GEMINI_MODEL:gemini-2.5-flash}

springdoc:
  api-docs:
    path: /api-docs
  swagger-ui:
    path: /swagger-ui.html
""")

add_h3(doc, "4.3.3. Cấu hình bảo mật – SecurityConfig.java")
add_code(doc, """
@Configuration
@EnableWebSecurity
@EnableMethodSecurity
@RequiredArgsConstructor
public class SecurityConfig {

    private final JwtAuthenticationFilter jwtAuthFilter;
    private final UserDetailsService userDetailsService;

    @Bean
    public SecurityFilterChain filterChain(HttpSecurity http) throws Exception {
        http
            .csrf(AbstractHttpConfigurer::disable)
            .cors(cors -> cors.configurationSource(corsConfigurationSource()))
            .sessionManagement(s -> s.sessionCreationPolicy(SessionCreationPolicy.STATELESS))
            .authorizeHttpRequests(auth -> auth
                .requestMatchers("/api/auth/**", "/v3/api-docs/**", "/swagger-ui/**")
                    .permitAll()
                .requestMatchers(HttpMethod.GET, "/actuator/health").permitAll()
                .anyRequest().authenticated()
            )
            .exceptionHandling(eh -> eh
                .authenticationEntryPoint((req, res, e) -> {
                    res.setStatus(HttpStatus.UNAUTHORIZED.value());
                    res.setContentType("application/json");
                    res.getWriter().write(
                        "{\\\"success\\\":false,\\\"message\\\":\\\"Bạn cần đăng nhập\\\"}");
                }))
            .addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class);
        return http.build();
    }

    @Bean
    public PasswordEncoder passwordEncoder() {
        return new BCryptPasswordEncoder(10);
    }

    @Bean
    public AuthenticationManager authenticationManager(AuthenticationConfiguration cfg)
            throws Exception {
        return cfg.getAuthenticationManager();
    }

    @Bean
    public CorsConfigurationSource corsConfigurationSource() {
        CorsConfiguration config = new CorsConfiguration();
        config.setAllowedOrigins(List.of("http://localhost:5173", "http://localhost:80"));
        config.setAllowedMethods(List.of("GET", "POST", "PUT", "DELETE", "PATCH", "OPTIONS"));
        config.setAllowedHeaders(List.of("*"));
        config.setAllowCredentials(true);
        UrlBasedCorsConfigurationSource source = new UrlBasedCorsConfigurationSource();
        source.registerCorsConfiguration("/**", config);
        return source;
    }
}
""")

add_h3(doc, "4.3.4. Tiện ích JWT – JwtUtil.java")
add_code(doc, """
@Component
public class JwtUtil {

    @Value("${jwt.secret}")
    private String secret;

    @Value("${jwt.expiration}")
    private long expiration;

    private SecretKey getKey() {
        return Keys.hmacShaKeyFor(Decoders.BASE64.decode(secret));
    }

    public String generateToken(UserDetails user, String role) {
        return Jwts.builder()
            .subject(user.getUsername())
            .claim("role", role)
            .issuedAt(new Date())
            .expiration(new Date(System.currentTimeMillis() + expiration))
            .signWith(getKey(), Jwts.SIG.HS256)
            .compact();
    }

    public String extractUsername(String token) {
        return parse(token).getPayload().getSubject();
    }

    public String extractRole(String token) {
        return (String) parse(token).getPayload().get("role");
    }

    public boolean isValid(String token, UserDetails user) {
        try {
            String username = extractUsername(token);
            return username.equals(user.getUsername())
                    && !parse(token).getPayload().getExpiration().before(new Date());
        } catch (JwtException | IllegalArgumentException e) {
            return false;
        }
    }

    private Jws<Claims> parse(String token) {
        return Jwts.parser()
            .verifyWith(getKey())
            .build()
            .parseSignedClaims(token);
    }
}
""")

add_h3(doc, "4.3.5. Bộ lọc JWT – JwtAuthenticationFilter.java")
add_code(doc, """
@Component
@RequiredArgsConstructor
public class JwtAuthenticationFilter extends OncePerRequestFilter {

    private final JwtUtil jwtUtil;
    private final UserDetailsService userDetailsService;

    @Override
    protected void doFilterInternal(HttpServletRequest req, HttpServletResponse res,
                                    FilterChain chain) throws ServletException, IOException {
        String header = req.getHeader(HttpHeaders.AUTHORIZATION);
        if (header == null || !header.startsWith("Bearer ")) {
            chain.doFilter(req, res);
            return;
        }
        String token = header.substring(7);
        try {
            String username = jwtUtil.extractUsername(token);
            if (username != null
                    && SecurityContextHolder.getContext().getAuthentication() == null) {
                UserDetails user = userDetailsService.loadUserByUsername(username);
                if (jwtUtil.isValid(token, user)) {
                    UsernamePasswordAuthenticationToken auth =
                        new UsernamePasswordAuthenticationToken(
                            user, null, user.getAuthorities());
                    auth.setDetails(new WebAuthenticationDetailsSource().buildDetails(req));
                    SecurityContextHolder.getContext().setAuthentication(auth);
                }
            }
        } catch (Exception ignored) {
            // Token sai → để controller tự xử lý qua AuthenticationEntryPoint
        }
        chain.doFilter(req, res);
    }
}
""")

add_h3(doc, "4.3.6. Entity Employee")
add_code(doc, """
@Entity
@Table(name = "employees")
@Getter
@Setter
@NoArgsConstructor
@AllArgsConstructor
public class Employee {

    @Id
    @GeneratedValue(strategy = GenerationType.IDENTITY)
    @Column(name = "employee_id")
    private Long employeeId;

    @ManyToOne(fetch = FetchType.LAZY)
    @JoinColumn(name = "user_id")
    private User user;

    @Column(name = "first_name", nullable = false, length = 50)
    private String firstName;

    @Column(name = "last_name", nullable = false, length = 50)
    private String lastName;

    @Column(length = 50)
    private String position;

    @Column(length = 50)
    private String department;

    @Column(name = "employee_group", length = 100)
    private String group;

    @Column(columnDefinition = "TEXT")
    private String skills;           // quản lý nhập tự do, ngăn cách dấu phẩy

    @Column(name = "hired_at")
    private LocalDateTime hiredAt;
}
""")
add_para(doc,
    "Lưu ý: hệ thống KHÔNG có entity Skill riêng. Kỹ năng nhân viên "
    "được lưu trong cột TEXT `skills` trên chính bảng `employees`, "
    "do quản lý nhập tự do dưới dạng danh sách phân cách bằng dấu "
    "phẩy. Cách thiết kế này cho phép quản lý thêm bất cứ kỹ năng "
    "nào mà không phải bảo trì một bảng skill và CRUD đi kèm.")

add_h3(doc, "4.3.7. Repository – EmployeeRepository.java")
add_code(doc, """
public interface EmployeeRepository extends JpaRepository<Employee, Long> {

    Optional<Employee> findByUserUsername(String username);

    List<Employee> findByDepartment(String department);
}
""")
add_para(doc,
    "TaskRepository và AttendanceRepository cung cấp các batch query "
    "để AiSuggestionService gom lịch sử + chấm công của tất cả nhân "
    "viên trong vài câu query, tránh N+1:")
add_code(doc, """
public interface TaskRepository extends JpaRepository<Task, Long> {
    List<Task> findByAssignedToEmployeeIdIn(Collection<Long> employeeIds);
    List<Task> findByAssignedToEmployeeId(Long employeeId);
}

public interface AttendanceRepository extends JpaRepository<Attendance, Long> {
    List<Attendance> findByEmployeeEmployeeIdInAndDateBetween(
            Collection<Long> employeeIds, LocalDate from, LocalDate to);
}
""")

add_h3(doc, "4.3.8. Service – AiSuggestionService.java (rút gọn)")
add_para(doc,
    "Module AI không tự tính điểm: backend chỉ gom dữ liệu thô và "
    "ủy thác xếp hạng cho Google Gemini gemini-2.5-flash. Phiên bản rút gọn dưới "
    "đây phản ánh logic thực tế trong "
    "`backend/src/main/java/com/example/taskmanagement/service/AiSuggestionService.java`.")
add_code(doc, """
@Service
public class AiSuggestionService {

    private static final int TOP_N = 5;
    private static final int ATTENDANCE_WINDOW_DAYS = 30;

    private final EmployeeRepository employeeRepository;
    private final TaskRepository taskRepository;
    private final AttendanceRepository attendanceRepository;
    private final RestClient restClient;
    private final ObjectMapper objectMapper;

    @Value("${gemini.api.key:}")  private String geminiApiKey;
    @Value("${gemini.api.model:gemini-2.5-flash}") private String geminiModel;

    @Cacheable(value = "ai_suggestions", key = "#request.cacheKey")
    public List<EmployeeSuggestionDTO> recommendEmployees(SuggestionRequest request) {
        List<Employee> employees = employeeRepository.findAll();
        Map<Long, EmployeeStats> stats = collectStats(employees);
        String prompt = buildPrompt(request, employees, stats);
        return callGemini(prompt, employees);
    }

    /** Gom số liệu thô — KHÔNG tính điểm số. */
    private Map<Long, EmployeeStats> collectStats(List<Employee> employees) {
        List<Long> ids = employees.stream().map(Employee::getEmployeeId).toList();
        Map<Long, List<Task>> tasksByEmp = taskRepository
            .findByAssignedToEmployeeIdIn(ids).stream()
            .filter(t -> t.getAssignedTo() != null)
            .collect(Collectors.groupingBy(t -> t.getAssignedTo().getEmployeeId()));
        Map<Long, Long> attendance = attendanceRepository
            .findByEmployeeEmployeeIdInAndDateBetween(ids,
                LocalDate.now().minusDays(ATTENDANCE_WINDOW_DAYS), LocalDate.now())
            .stream()
            .collect(Collectors.groupingBy(a -> a.getEmployee().getEmployeeId(),
                                            Collectors.counting()));
        // build EmployeeStats { totalTasks, activeTasks, completedTasks,
        //                       completedOnTime, completedWithDueDate,
        //                       avgDaysLate, attendanceDays }
        // ...
    }

    private String buildPrompt(SuggestionRequest req, List<Employee> emps,
                               Map<Long, EmployeeStats> stats) {
        StringBuilder sb = new StringBuilder();
        sb.append("Bạn là trợ lý AI giúp quản lý chọn nhân viên phù hợp nhất.\\n\\n");
        sb.append("=== TASK CẦN GIAO ===\\n")
          .append("- Tiêu đề: ").append(req.getTaskTitle()).append("\\n")
          .append("- Mô tả: ").append(req.getTaskDescription()).append("\\n")
          .append("- Kỹ năng yêu cầu: ").append(req.getRequiredSkills()).append("\\n");
        sb.append("\\n=== DỮ LIỆU LỊCH SỬ ===\\n");
        for (Employee e : emps) {
            EmployeeStats s = stats.get(e.getEmployeeId());
            sb.append("• ID=").append(e.getEmployeeId())
              .append(" | ").append(e.getFirstName()).append(" ").append(e.getLastName())
              .append(" | ").append(e.getDepartment()).append(" | ").append(e.getPosition())
              .append("\\n    - Kỹ năng: ").append(e.getSkills())
              .append("\\n    - Tiến độ: ").append(s.completedTasks).append("/").append(s.totalTasks)
              .append("\\n    - Đúng hạn: ").append(s.completedOnTime).append("/").append(s.completedWithDueDate)
              .append("\\n    - Chấm công 30 ngày: ").append(s.attendanceDays).append("/22\\n");
        }
        sb.append("\\nHãy gợi ý TOP ").append(TOP_N).append(" nhân viên, ưu tiên: ")
          .append("(1) kỹ năng & chuyên môn; (2) tiến độ; (3) đúng hạn; (4) chấm công. ")
          .append("KHÔNG tính điểm số. Trả về JSON [{employeeId, rank, reasoning}].");
        return sb.toString();
    }

    private List<EmployeeSuggestionDTO> callGemini(String prompt, List<Employee> emps) {
        Map<String, Object> body = Map.of(
            "contents", List.of(Map.of(
                "parts", List.of(Map.of("text", prompt)))),
            "generationConfig", Map.of(
                "temperature", 0.3,
                "responseMimeType", "application/json"));  // ép Gemini trả JSON
        String resp = restClient.post()
            .uri("/v1beta/models/{model}:generateContent", geminiModel)
            .header("x-goog-api-key", geminiApiKey)  // key qua header, không lộ trong URL
            .contentType(MediaType.APPLICATION_JSON)
            .body(body).retrieve().body(String.class);
        return parseGeminiResponse(resp, emps);  // JSON → List<DTO>
    }
}
""")
add_caption(doc, "Hình 4.3: Sơ đồ luồng module AiSuggestionService", kind="figure")

add_h3(doc, "4.3.10. Bảng tổng hợp REST endpoint")
add_table(
    doc,
    headers=["Method", "Endpoint", "Mô tả", "Auth"],
    rows=[
        ("POST", "/api/auth/register", "Đăng ký tài khoản mới.", "Public"),
        ("POST", "/api/auth/login", "Đăng nhập, trả về JWT.", "Public"),
        ("POST", "/api/auth/change-password", "Đổi mật khẩu.", "JWT"),
        ("GET",  "/api/employees", "Danh sách nhân viên.", "JWT (MANAGER/ADMIN)"),
        ("POST", "/api/employees", "Tạo nhân viên (kèm trường skills).", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/employees/{id}", "Chi tiết nhân viên.", "JWT (MANAGER/ADMIN)"),
        ("PUT",  "/api/employees/{id}", "Cập nhật nhân viên.", "JWT (MANAGER/ADMIN)"),
        ("DELETE","/api/employees/{id}", "Xóa nhân viên.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/employees/me", "Hồ sơ nhân viên của user đang đăng nhập.", "JWT (mọi role)"),
        ("GET",  "/api/projects", "Danh sách dự án.", "JWT (mọi role)"),
        ("POST", "/api/projects", "Tạo dự án.", "JWT (MANAGER/ADMIN)"),
        ("PUT",  "/api/projects/{id}", "Cập nhật dự án.", "JWT (MANAGER/ADMIN)"),
        ("DELETE","/api/projects/{id}", "Xóa dự án.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/tasks", "Danh sách task.", "JWT (mọi role)"),
        ("POST", "/api/tasks", "Tạo task (kèm trường requiredSkills).", "JWT (MANAGER/ADMIN)"),
        ("PUT",  "/api/tasks/{id}", "Cập nhật task.", "JWT (MANAGER/ADMIN)"),
        ("DELETE","/api/tasks/{id}", "Xóa task.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/tasks/me", "Task của user đang đăng nhập.", "JWT (mọi role)"),
        ("PATCH","/api/tasks/{id}/status", "Đổi trạng thái task được giao.", "JWT (mọi role)"),
        ("GET",  "/api/attendance", "Lịch sử chấm công toàn hệ thống.", "JWT (MANAGER/ADMIN)"),
        ("POST", "/api/attendance", "Quản lý ghi chấm công cho nhân viên.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/attendance/me", "Lịch sử chấm công của user đang đăng nhập.", "JWT (mọi role)"),
        ("POST", "/api/attendance/me/checkin", "Tự bấm vào ca.", "JWT (mọi role)"),
        ("POST", "/api/attendance/me/checkout", "Tự bấm tan ca.", "JWT (mọi role)"),
        ("POST", "/api/suggestions/recommend", "AI gợi ý nhân viên cho task.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/suggestions/recommend/{taskId}", "AI gợi ý cho task đã có.", "JWT (MANAGER/ADMIN)"),
        ("GET",  "/api/suggestions", "Lịch sử gợi ý.", "JWT (MANAGER/ADMIN)"),
    ],
    col_widths=[1.8, 5.0, 5.5, 3.2],
)
add_caption(doc, "Bảng 4.2: Danh sách các REST endpoint của backend", kind="table")

add_h3(doc, "4.3.11. Tiêu chí xếp hạng AI gợi ý nhân viên")
add_table(
    doc,
    headers=["Tiêu chí", "Mức ưu tiên", "Nguồn dữ liệu gửi cho LLM"],
    rows=[
        ("Kỹ năng & chuyên môn", "1 (cao nhất)",
         "Trường `skills` (TEXT) của employee + `requiredSkills` của task + chức danh + phòng ban."),
        ("Tiến độ task trước", "2",
         "Tỷ lệ completedTasks / totalTasks và số task IN_PROGRESS của từng nhân viên."),
        ("Thời gian hoàn thành", "3",
         "Số task hoàn thành đúng hạn / số task hoàn thành có deadline; số ngày trễ trung bình."),
        ("Chấm công", "4 (thấp nhất)",
         "Số ngày có mặt trong 30 ngày gần nhất, so với 22 ngày làm việc chuẩn."),
    ],
    col_widths=[4.5, 2.5, 8.5],
)
add_caption(doc, "Bảng 4.3: Tiêu chí xếp hạng AI gợi ý nhân viên", kind="table")
add_para(doc,
    "Khác với thiết kế ban đầu (Weighted Scoring 35%/25%/25%/15%), hệ "
    "thống không quy điểm số ra trọng số cố định mà chỉ truyền thứ tự "
    "ưu tiên và dữ liệu thô để LLM tự cân nhắc. Cách làm này linh "
    "hoạt hơn với các tình huống đặc thù — vd. khi mô tả task chung "
    "chung hoặc khi kỹ năng yêu cầu để rỗng, AI vẫn suy luận đúng "
    "vị trí phù hợp từ chức danh và phòng ban.")


add_h2(doc, "4.4. Triển khai Frontend React + Vite")
add_h3(doc, "4.4.1. Khởi tạo dự án và package.json")
add_code(doc, """
{
  "name": "task-management-frontend",
  "private": true,
  "version": "1.0.0",
  "type": "module",
  "scripts": {
    "dev": "vite",
    "build": "vite build",
    "preview": "vite preview"
  },
  "dependencies": {
    "axios": "^1.7.2",
    "chart.js": "^4.4.4",
    "react": "^18.3.1",
    "react-chartjs-2": "^5.2.0",
    "react-dom": "^18.3.1",
    "react-hot-toast": "^2.4.1",
    "react-router-dom": "^6.26.0"
  },
  "devDependencies": {
    "@vitejs/plugin-react": "^4.3.1",
    "autoprefixer": "^10.4.20",
    "postcss": "^8.4.41",
    "tailwindcss": "^3.4.10",
    "vite": "^5.4.0"
  }
}
""")

add_h3(doc, "4.4.2. axiosConfig.js – cấu hình HTTP client")
add_code(doc, """
import axios from 'axios';

const instance = axios.create({
  baseURL: import.meta.env.VITE_API_URL || 'http://localhost:5000/api',
  timeout: 30000,
});

instance.interceptors.request.use((config) => {
  const token = localStorage.getItem('token');
  if (token) config.headers.Authorization = `Bearer ${token}`;
  return config;
});

instance.interceptors.response.use(
  (res) => res,
  (err) => {
    if (err?.response?.status === 401) {
      localStorage.removeItem('token');
      localStorage.removeItem('user');
      if (window.location.pathname !== '/login') {
        window.location.href = '/login';
      }
    }
    return Promise.reject(err);
  }
);

export default instance;
""")

add_h3(doc, "4.4.3. AuthContext.jsx – state xác thực toàn cục")
add_code(doc, """
import { createContext, useContext, useState, useEffect } from 'react';
import axios from '../api/axiosConfig';

const AuthContext = createContext(null);

export function AuthProvider({ children }) {
  const [user, setUser] = useState(null);
  const [loading, setLoading] = useState(true);

  useEffect(() => {
    const saved = localStorage.getItem('user');
    if (saved) setUser(JSON.parse(saved));
    setLoading(false);
  }, []);

  async function login(username, password) {
    const { data } = await axios.post('/auth/login', { username, password });
    localStorage.setItem('token', data.token);
    localStorage.setItem('user', JSON.stringify({
      username: data.username, role: data.role,
    }));
    setUser({ username: data.username, role: data.role });
    return data;
  }

  function logout() {
    localStorage.removeItem('token');
    localStorage.removeItem('user');
    setUser(null);
  }

  return (
    <AuthContext.Provider value={{ user, loading, login, logout }}>
      {children}
    </AuthContext.Provider>
  );
}

export const useAuth = () => useContext(AuthContext);
""")

add_h3(doc, "4.4.4. App.jsx – routing với ProtectedRoute")
add_code(doc, """
import { BrowserRouter, Routes, Route, Navigate, Outlet } from 'react-router-dom';
import { useAuth } from './context/AuthContext';
import Layout from './components/Layout';
import Login from './pages/Login';
import Register from './pages/Register';
import Dashboard from './pages/Dashboard';
import Employees from './pages/Employees';
import Projects from './pages/Projects';
import Tasks from './pages/Tasks';
import Attendance from './pages/Attendance';
import AiSuggestions from './pages/AiSuggestions';

function ProtectedRoute() {
  const { user, loading } = useAuth();
  if (loading) return <div className="p-6">Đang tải...</div>;
  return user ? <Outlet /> : <Navigate to="/login" replace />;
}

export default function App() {
  return (
    <BrowserRouter>
      <Routes>
        <Route path="/login" element={<Login />} />
        <Route path="/register" element={<Register />} />
        <Route element={<ProtectedRoute />}>
          <Route element={<Layout />}>
            <Route path="/" element={<Navigate to="/dashboard" replace />} />
            <Route path="/dashboard" element={<Dashboard />} />
            <Route path="/employees" element={<Employees />} />
            <Route path="/projects" element={<Projects />} />
            <Route path="/tasks" element={<Tasks />} />
            <Route path="/attendance" element={<Attendance />} />
            <Route path="/ai-suggestions" element={<AiSuggestions />} />
          </Route>
        </Route>
      </Routes>
    </BrowserRouter>
  );
}
""")

add_h3(doc, "4.4.5. AiSuggestions.jsx – trang gợi ý AI (rút gọn)")
add_code(doc, """
import { useState } from 'react';
import axios from '../api/axiosConfig';

export default function AiSuggestions() {
  const [form, setForm] = useState({
    taskTitle: '', taskDescription: '', requiredSkills: '',
  });
  const [result, setResult] = useState([]);
  const [loading, setLoading] = useState(false);

  async function suggest() {
    if (!form.taskTitle.trim()) return alert('Hãy nhập tiêu đề công việc');
    setLoading(true);
    try {
      const { data } = await axios.post('/api/suggestions/recommend', form);
      setResult(data.data || []);   // mảng [{employeeId, rank, reasoning, ...}]
    } catch (e) {
      alert('Gợi ý thất bại: ' + (e.response?.data?.message || e.message));
    } finally { setLoading(false); }
  }

  return (
    <div className="p-6 space-y-6">
      <h1 className="text-2xl font-bold">AI Gợi ý nhân viên</h1>
      <div className="bg-white p-4 rounded-lg shadow space-y-3">
        <input className="border rounded w-full px-3 py-2"
               placeholder="Tiêu đề công việc"
               value={form.taskTitle}
               onChange={(e) => setForm({ ...form, taskTitle: e.target.value })} />
        <textarea className="border rounded w-full p-2" rows={3}
                  placeholder="Mô tả công việc..."
                  value={form.taskDescription}
                  onChange={(e) => setForm({ ...form, taskDescription: e.target.value })} />
        <input className="border rounded w-full px-3 py-2"
               placeholder="Kỹ năng yêu cầu (Java, Spring Boot, ...)"
               value={form.requiredSkills}
               onChange={(e) => setForm({ ...form, requiredSkills: e.target.value })} />
        <button onClick={suggest} disabled={loading}
                className="bg-green-600 text-white px-6 py-2 rounded">
          {loading ? 'Đang xử lý...' : 'Gợi ý ngay'}
        </button>
      </div>

      <div className="space-y-3">
        {result.map((emp) => (
          <div key={emp.employeeId} className="bg-white p-4 rounded-lg shadow">
            <div className="flex items-center gap-3">
              <span className="bg-indigo-100 text-indigo-700 px-3 py-1 rounded-full">
                #{emp.rank}
              </span>
              <h3 className="font-bold">{emp.firstName} {emp.lastName}</h3>
              <span className="text-sm text-gray-500">{emp.department}</span>
            </div>
            <p className="text-sm text-gray-700 mt-2">{emp.reasoning}</p>
          </div>
        ))}
      </div>
    </div>
  );
}
""")

add_h3(doc, "4.4.6. Hỗ trợ giao diện song ngữ Việt/Anh (i18n)")
add_para(doc,
    "Nhằm mở rộng khả năng sử dụng, frontend được bổ sung cơ chế đa ngôn "
    "ngữ (internationalization – i18n) hỗ trợ song song tiếng Việt và tiếng "
    "Anh. Cơ chế này được tự xây dựng hoàn toàn trong ứng dụng, không phụ "
    "thuộc dịch vụ dịch bên ngoài (như Google Translate); toàn bộ chuỗi "
    "hiển thị được dịch tức thì phía client.")
add_para(doc,
    "Giải pháp gồm ba thành phần chính: (1) tệp từ điển translations.js ánh "
    "xạ mỗi chuỗi tiếng Việt gốc sang bản dịch tiếng Anh; (2) LanguageContext "
    "– một React Context cung cấp ngôn ngữ hiện tại, hàm đổi ngôn ngữ "
    "setLang() và hàm dịch t(); (3) component LanguageSwitcher hiển thị nút "
    "cờ Việt Nam / Anh để người dùng chuyển đổi. Lựa chọn ngôn ngữ được lưu "
    "vào localStorage nên được ghi nhớ ở những lần truy cập sau. Mọi trang "
    "và component dùng chung đều lấy chuỗi hiển thị thông qua hàm t().")
add_code(doc, """
// LanguageContext.jsx - rut gon
const LanguageContext = createContext(null);

export function LanguageProvider({ children }) {
  const [lang, setLang] = useState(
    () => localStorage.getItem('app_lang') || 'vi');

  // t(key): key chinh la chuoi tieng Viet goc
  const t = useCallback((key, params) => {
    let str = lang === 'en' ? (en[key] ?? key) : key;
    if (params)
      for (const [k, v] of Object.entries(params))
        str = str.split('{' + k + '}').join(String(v));
    return str;
  }, [lang]);

  return (
    <LanguageContext.Provider value={{ lang, setLang, t }}>
      {children}
    </LanguageContext.Provider>
  );
}

export const useTranslation = () => useContext(LanguageContext);
""")
add_para(doc,
    "Khi người dùng nhấn nút cờ, giá trị lang trong Context thay đổi, React "
    "tự render lại toàn bộ giao diện với bản dịch tương ứng mà không cần "
    "tải lại trang. Việc dùng chính chuỗi tiếng Việt làm khóa giúp giảm "
    "công sức bảo trì: chế độ tiếng Việt luôn hiển thị đúng kể cả khi từ "
    "điển tiếng Anh chưa bổ sung mục tương ứng.")


add_h2(doc, "4.5. Triển khai ứng dụng Mobile Flutter")
add_h3(doc, "4.5.1. pubspec.yaml")
add_code(doc, """
name: task_management_mobile
description: Mobile app cho hệ thống quản lý công việc.
publish_to: 'none'
version: 1.0.0+1

environment:
  sdk: '>=3.0.0 <4.0.0'

dependencies:
  flutter:
    sdk: flutter
  cupertino_icons: ^1.0.8
  dio: ^5.5.0
  provider: ^6.1.2
  shared_preferences: ^2.3.0
  intl: ^0.19.0
  flutter_secure_storage: ^9.2.2

dev_dependencies:
  flutter_test:
    sdk: flutter
  flutter_lints: ^4.0.0

flutter:
  uses-material-design: true
""")

add_h3(doc, "4.5.2. main.dart – khởi động app")
add_code(doc, """
import 'package:flutter/material.dart';
import 'package:provider/provider.dart';
import 'services/auth_service.dart';
import 'screens/login_screen.dart';
import 'screens/dashboard_screen.dart';

void main() {
  runApp(
    ChangeNotifierProvider(
      create: (_) => AuthService(),
      child: const TaskManagementApp(),
    ),
  );
}

class TaskManagementApp extends StatelessWidget {
  const TaskManagementApp({super.key});
  @override
  Widget build(BuildContext context) {
    return MaterialApp(
      title: 'Task Manager',
      debugShowCheckedModeBanner: false,
      theme: ThemeData(
        primarySwatch: Colors.blue,
        useMaterial3: true,
      ),
      home: Consumer<AuthService>(
        builder: (_, auth, __) =>
            auth.isAuthenticated ? const DashboardScreen() : const LoginScreen(),
      ),
    );
  }
}
""")

add_h3(doc, "4.5.3. dio_client.dart – HTTP client với JWT")
add_code(doc, """
import 'package:dio/dio.dart';
import 'package:flutter_secure_storage/flutter_secure_storage.dart';

class DioClient {
  static final DioClient _instance = DioClient._internal();
  factory DioClient() => _instance;
  late final Dio dio;
  final _storage = const FlutterSecureStorage();

  DioClient._internal() {
    dio = Dio(BaseOptions(
      baseUrl: 'http://10.0.2.2:5000/api',     // Android emulator → host
      connectTimeout: const Duration(seconds: 15),
      receiveTimeout: const Duration(seconds: 15),
    ));
    dio.interceptors.add(InterceptorsWrapper(
      onRequest: (options, handler) async {
        final token = await _storage.read(key: 'token');
        if (token != null) options.headers['Authorization'] = 'Bearer $token';
        handler.next(options);
      },
      onError: (err, handler) async {
        if (err.response?.statusCode == 401) {
          await _storage.deleteAll();
        }
        handler.next(err);
      },
    ));
  }
}
""")


add_h2(doc, "4.6. Thiết lập PostgreSQL và Redis")
add_h3(doc, "4.6.1. Script khởi tạo PostgreSQL – init.sql")
add_code(doc, """
-- Đảm bảo encoding UTF8
CREATE DATABASE taskmgmt
    WITH ENCODING = 'UTF8'
         LC_COLLATE = 'en_US.utf8'
         LC_CTYPE   = 'en_US.utf8'
         TEMPLATE   = template0;

\\c taskmgmt

CREATE TABLE users (
    id          BIGSERIAL    PRIMARY KEY,
    username    VARCHAR(50)  NOT NULL UNIQUE,
    email       VARCHAR(100) NOT NULL UNIQUE,
    password    VARCHAR(255) NOT NULL,
    role        VARCHAR(20)  NOT NULL,
    status      VARCHAR(20)  NOT NULL DEFAULT 'ACTIVE',
    created_at  TIMESTAMP    NOT NULL DEFAULT NOW()
);

CREATE INDEX idx_users_username ON users(username);

CREATE TABLE employees (
    employee_id     BIGSERIAL    PRIMARY KEY,
    user_id         BIGINT       REFERENCES users(id),
    first_name      VARCHAR(50)  NOT NULL,
    last_name       VARCHAR(50)  NOT NULL,
    position        VARCHAR(50),
    department      VARCHAR(50),
    employee_group  VARCHAR(100),
    skills          TEXT,        -- quản lý nhập tự do, phân cách bằng dấu phẩy
    hired_at        TIMESTAMP
);

CREATE INDEX idx_employees_department ON employees(department);

-- ... tương tự cho projects, tasks (kèm required_skills TEXT), attendances,
-- suggestions. Lưu ý: hệ thống KHÔNG có bảng skills riêng — kỹ năng
-- được lưu thẳng trên cột employees.skills.

-- Tài khoản admin mặc định, mật khẩu = "admin123" đã băm BCrypt
INSERT INTO users (username, email, password, role) VALUES
('admin', 'admin@hutech.edu.vn',
 '$2a$10$kCQ3wAjY9pP1zS1zG9b8/eL7H4j1nFq8B0a9G2hY1xK2H4f3L1d6S',
 'ADMIN');
""")

add_h3(doc, "4.6.2. Cấu hình Redis")
add_para(doc,
    "Redis được chạy với cấu hình mặc định, không bật authentication "
    "(do chạy trong mạng nội bộ Docker). Nếu triển khai production, "
    "cần bật requirepass và TLS. CacheManager của Spring Boot được "
    "cấu hình với TTL 5 phút – đủ để gom các request lặp lại trong "
    "ngắn hạn nhưng không quá lâu để dữ liệu trở nên lạc hậu.")
add_code(doc, """
@Configuration
@EnableCaching
public class CacheConfig {

    @Bean
    public RedisCacheConfiguration cacheConfiguration() {
        return RedisCacheConfiguration.defaultCacheConfig()
            .entryTtl(Duration.ofMinutes(5))
            .disableCachingNullValues()
            .serializeKeysWith(RedisSerializationContext.SerializationPair
                .fromSerializer(new StringRedisSerializer()))
            .serializeValuesWith(RedisSerializationContext.SerializationPair
                .fromSerializer(new GenericJackson2JsonRedisSerializer()));
    }

    @Bean
    public RedisCacheManager cacheManager(RedisConnectionFactory cf) {
        return RedisCacheManager.builder(cf)
            .cacheDefaults(cacheConfiguration())
            .build();
    }
}
""")


add_h2(doc, "4.7. Đóng gói với Docker Compose")
add_h3(doc, "4.7.1. Dockerfile backend (multi-stage)")
add_code(doc, """
# ---- STAGE 1: build ----
FROM maven:3.9.6-eclipse-temurin-17 AS build
WORKDIR /app
COPY pom.xml .
RUN mvn dependency:go-offline -B
COPY src ./src
RUN mvn -B -DskipTests clean package

# ---- STAGE 2: run ----
FROM eclipse-temurin:17-jre-jammy
WORKDIR /app
COPY --from=build /app/target/*.jar app.jar
EXPOSE 5000
ENTRYPOINT ["java", "-jar", "app.jar"]
""")

add_h3(doc, "4.7.2. Dockerfile frontend")
add_code(doc, """
# ---- STAGE 1: build ----
FROM node:18-alpine AS build
WORKDIR /app
COPY package*.json ./
RUN npm ci
COPY . .
RUN npm run build

# ---- STAGE 2: serve ----
FROM nginx:alpine
COPY --from=build /app/dist /usr/share/nginx/html
COPY nginx.conf /etc/nginx/conf.d/default.conf
EXPOSE 80
CMD ["nginx", "-g", "daemon off;"]
""")

add_h3(doc, "4.7.3. docker-compose.yml")
add_code(doc, """
version: '3.9'

services:
  postgres:
    image: postgres:16
    container_name: taskmgmt_postgres
    environment:
      POSTGRES_DB: taskmgmt
      POSTGRES_USER: postgres
      POSTGRES_PASSWORD: ${DB_PASS:-postgres}
    ports:
      - "5432:5432"
    volumes:
      - pgdata:/var/lib/postgresql/data
      - ./backend/src/main/resources/db/init.sql:/docker-entrypoint-initdb.d/init.sql
    healthcheck:
      test: ["CMD-SHELL", "pg_isready -U postgres"]
      interval: 5s
      retries: 10
    networks: [taskmgmt_net]

  redis:
    image: redis:7-alpine
    container_name: taskmgmt_redis
    ports:
      - "6379:6379"
    volumes:
      - redisdata:/data
    networks: [taskmgmt_net]

  backend:
    build: ./backend
    container_name: taskmgmt_backend
    environment:
      DB_HOST: postgres
      DB_NAME: taskmgmt
      DB_USER: postgres
      DB_PASS: ${DB_PASS:-postgres}
      REDIS_HOST: redis
      JWT_SECRET: ${JWT_SECRET}
      GEMINI_API_KEY: ${GEMINI_API_KEY}
    ports:
      - "5000:5000"
    depends_on:
      postgres:
        condition: service_healthy
      redis:
        condition: service_started
    networks: [taskmgmt_net]

  frontend:
    build: ./frontend
    container_name: taskmgmt_frontend
    environment:
      VITE_API_URL: http://localhost:5000/api
    ports:
      - "5173:80"
    depends_on:
      - backend
    networks: [taskmgmt_net]

volumes:
  pgdata:
  redisdata:

networks:
  taskmgmt_net:
    driver: bridge
""")
add_caption(doc, "Hình 4.4: Sơ đồ Docker Compose", kind="figure")


add_h2(doc, "4.8. Triển khai và vận hành")
add_h3(doc, "4.8.1. Triển khai môi trường phát triển")
add_code(doc, """
# Bước 1: Clone repository
git clone https://github.com/nhathao428/task-management-system.git
cd task-management-system

# Bước 2: Tạo file .env từ mẫu
cp .env.example .env
# Mở .env, điền GEMINI_API_KEY và JWT_SECRET (chuỗi base64 32+ byte)

# Bước 3: Khởi động toàn bộ stack
docker-compose up -d

# Bước 4: Kiểm tra trạng thái các container
docker-compose ps

# Bước 5: Mở trình duyệt
#   - Frontend:   http://localhost:5173
#   - Backend:    http://localhost:5000
#   - Swagger UI: http://localhost:5000/swagger-ui.html
""")
add_h3(doc, "4.8.2. Triển khai production trên VPS")
add_para(doc,
    "Để triển khai lên VPS production, sử dụng file docker-compose.prod.yml "
    "kèm reverse proxy Caddy (tự động cấp SSL Let's Encrypt). Cấu hình "
    "Caddyfile như sau:")
add_code(doc, """
api.taskmgmt.example.com {
    reverse_proxy backend:5000
}

taskmgmt.example.com {
    root * /srv/frontend
    file_server
    try_files {path} /index.html
}
""")
add_para(doc,
    "Sau khi cấu hình DNS A record trỏ tên miền về IP VPS, chạy "
    "docker-compose -f docker-compose.prod.yml up -d. Caddy tự cấp "
    "chứng chỉ SSL trong khoảng 60 giây. Toàn bộ traffic sẽ qua "
    "HTTPS với HTTP/2 mặc định.")

add_h3(doc, "4.8.3. Triển khai thực tế trên AWS EC2")
add_para(doc,
    "Để sản phẩm vận hành ổn định lâu dài và truy cập công khai qua "
    "Internet, hệ thống được triển khai thực tế trên hạ tầng đám mây Amazon "
    "Web Services (AWS) theo mô hình all-in-one: một máy chủ ảo EC2 chạy "
    "toàn bộ stack thông qua Docker Compose.")
add_para(doc, "Cấu hình triển khai cụ thể:")
add_table(
    doc,
    headers=["Thành phần", "Cấu hình"],
    rows=[
        ("Máy chủ", "AWS EC2 t3.small (2 vCPU, 2 GB RAM), Ubuntu Server 24.04 LTS"),
        ("Ổ lưu trữ", "20 GB EBS gp3"),
        ("Địa chỉ truy cập", "Elastic IP tĩnh, gắn cố định với instance"),
        ("Bảo mật mạng", "Security Group chỉ mở cổng 22 (SSH) và 80 (HTTP)"),
        ("Thành phần chạy",
         "docker-compose.aws.yml: PostgreSQL + Redis + backend + frontend + Caddy"),
    ],
    col_widths=[4.5, 11.0],
)
add_para(doc,
    "Quy trình triển khai gồm các bước: tạo EC2 instance và Elastic IP, "
    "cài Docker, tạo vùng nhớ swap 2 GB hỗ trợ quá trình build trên máy "
    "RAM thấp, nạp mã nguồn, sinh các khóa bí mật (DB_PASSWORD, JWT_SECRET) "
    "rồi khởi chạy bằng lệnh docker compose. Caddy đóng vai trò reverse "
    "proxy: định tuyến các đường dẫn /api/* tới backend và các đường dẫn "
    "còn lại tới frontend trên cùng một origin. Toàn bộ quy trình được tài "
    "liệu hóa chi tiết trong tệp DEPLOY-AWS.md kèm theo mã nguồn.")
add_para(doc,
    "Ở giai đoạn này, hệ thống chạy qua giao thức HTTP với địa chỉ IP công "
    "khai. Khi đăng ký được tên miền riêng, chỉ cần chuyển sang cấu hình "
    "docker-compose.prod.yml để Caddy tự động cấp chứng chỉ HTTPS qua "
    "Let's Encrypt.")


# ==================================================================
# CHƯƠNG 5: KIỂM THỬ VÀ ĐÁNH GIÁ KẾT QUẢ
# ==================================================================
add_h1(doc, "CHƯƠNG 5. KIỂM THỬ VÀ ĐÁNH GIÁ KẾT QUẢ")

add_h2(doc, "5.1. Kế hoạch kiểm thử")
add_h3(doc, "5.1.1. Mục tiêu kiểm thử")
add_para(doc,
    "Kiểm thử (testing) là một trong những giai đoạn quan trọng nhất của vòng "
    "đời phát triển phần mềm, nhằm đảm bảo chất lượng sản phẩm trước khi bàn "
    "giao. Mục tiêu kiểm thử trong đề tài này gồm:")
for s in [
    "Xác minh hệ thống đáp ứng đúng 14 yêu cầu chức năng đã đặc tả.",
    "Đảm bảo các yêu cầu phi chức năng (hiệu năng, bảo mật, khả dụng) được đáp ứng.",
    "Phát hiện và sửa lỗi (bug) trước khi đưa vào sử dụng thực tế.",
    "Kiểm tra tính ổn định khi các module tương tác với nhau (integration).",
    "Đo lường hiệu năng đặc biệt với module AI Suggestion (yêu cầu ≤ 2 giây).",
]:
    add_bullet(doc, s)
add_h3(doc, "5.1.2. Phương pháp và công cụ")
add_table(
    doc,
    headers=["Mức kiểm thử", "Phương pháp", "Công cụ"],
    rows=[
        ("Unit Testing", "Kiểm thử từng hàm/method độc lập, mock dependency.",
         "JUnit 5, Mockito, Spring Boot Test."),
        ("API Testing", "Kiểm thử từng endpoint REST với các đầu vào hợp lệ/không hợp lệ.",
         "Postman, Postman Collection Runner."),
        ("Integration Testing", "Kiểm thử luồng đầy đủ qua nhiều module.",
         "Postman Collection + môi trường Docker Compose."),
        ("UI Testing", "Kiểm thử thủ công các luồng quan trọng trên trình duyệt và mobile.",
         "Chrome 120, Android Emulator (Pixel 7, API 34)."),
        ("Performance Testing", "Đo thời gian phản hồi của API trong điều kiện tải nhẹ.",
         "Postman + Chrome DevTools."),
        ("Security Testing", "Kiểm thử các kịch bản tấn công cơ bản: SQL injection, XSS, JWT giả.",
         "Postman, Burp Suite Community."),
    ],
    col_widths=[3.5, 7.5, 4.5],
)
add_h3(doc, "5.1.3. Ma trận test case theo module")
add_table(
    doc,
    headers=["Module", "Số test case", "Loại", "Tỷ lệ pass"],
    rows=[
        ("Auth (đăng ký/đăng nhập/đổi mật khẩu)", "8", "API + Security", "100%"),
        ("Employee", "6", "API + UI", "100%"),
        ("Project", "4", "API + UI", "100%"),
        ("Task", "6", "API + UI", "100%"),
        ("Attendance", "4", "API + UI", "100%"),
        ("AI Suggestion", "4", "API + Performance", "100%"),
        ("Dashboard", "2", "UI", "100%"),
        ("Self-service nhân viên (My Tasks/My Attendance)", "3", "API + UI", "100%"),
        ("Geofence (xác thực GPS chấm công)", "5", "Unit + API + UI", "100%"),
        ("Tổng", "42", "Tổng hợp", "100%"),
    ],
    col_widths=[5.5, 3.0, 4.0, 3.0],
)
add_caption(doc, "Bảng 5.1: Ma trận test case theo module", kind="table")
add_h3(doc, "5.1.4. Tiêu chí đánh giá kết quả")
for s in [
    "PASS: Kết quả thực tế khớp hoàn toàn với kết quả mong đợi.",
    "FAIL: Kết quả thực tế khác kết quả mong đợi → ghi nhận lỗi, sửa và kiểm thử lại.",
    "BLOCKED: Không thể thực hiện được test case do tiền điều kiện không thoả mãn.",
    "Mức độ nghiêm trọng (severity): CRITICAL (sập hệ thống), HIGH (sai logic nghiệp vụ), MEDIUM (lỗi UI), LOW (lỗi hiển thị nhỏ).",
]:
    add_bullet(doc, s)


add_h2(doc, "5.2. Kịch bản kiểm thử chi tiết")
add_h3(doc, "5.2.1. Module Auth")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-A01", "Đăng ký tài khoản hợp lệ",
         "username=hao01, email=hao01@test.com, password=Ab12345!",
         "201 Created, user lưu vào DB", "201, user lưu", "PASS"),
        ("TC-A02", "Đăng ký username đã tồn tại",
         "username=admin (đã có)",
         "400 Bad Request, message rõ", "400 đúng", "PASS"),
        ("TC-A03", "Đăng ký email không hợp lệ",
         "email=khong-phai-email",
         "400, validate fail", "400 đúng", "PASS"),
        ("TC-A04", "Đăng ký password quá ngắn",
         "password=123",
         "400, validate fail", "400 đúng", "PASS"),
        ("TC-A05", "Đăng nhập đúng",
         "username=admin, password=admin123",
         "200, trả JWT", "200, JWT hợp lệ", "PASS"),
        ("TC-A06", "Đăng nhập sai password",
         "username=admin, password=sai",
         "401 Unauthorized", "401 đúng", "PASS"),
        ("TC-A07", "Đăng nhập tài khoản LOCKED",
         "user.status=LOCKED",
         "403 Forbidden", "403 đúng", "PASS"),
        ("TC-A08", "Đổi mật khẩu cũ sai",
         "oldPassword sai",
         "400, message rõ", "400 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)
add_caption(doc, "Bảng 5.2: Kịch bản kiểm thử module Auth", kind="table")

add_h3(doc, "5.2.2. Module Employee")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-E01", "Lấy danh sách nhân viên có JWT",
         "GET /api/employees, header Bearer hợp lệ",
         "200, danh sách JSON", "200 đúng", "PASS"),
        ("TC-E02", "Lấy danh sách không có JWT",
         "Không có header Authorization",
         "401 Unauthorized", "401 đúng", "PASS"),
        ("TC-E03", "Tạo nhân viên kèm skills",
         "{firstName, lastName, department, skills='Java, SQL'}",
         "201, employee mới có id", "201 đúng", "PASS"),
        ("TC-E04", "Tạo nhân viên thiếu firstName",
         "Body không có firstName",
         "400, validate fail", "400 đúng", "PASS"),
        ("TC-E05", "Cập nhật nhân viên không tồn tại",
         "PUT /api/employees/99999",
         "404 Not Found", "404 đúng", "PASS"),
        ("TC-E06", "Xóa nhân viên bằng role EMPLOYEE",
         "DELETE /api/employees/5 với role EMPLOYEE",
         "403 Forbidden", "403 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)
add_caption(doc, "Bảng 5.3: Kịch bản kiểm thử module Employee", kind="table")

add_h3(doc, "5.2.3. Module Task và Project")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-T01", "Tạo task hợp lệ",
         "{title, projectId, assigneeId, dueDate}",
         "201, task được lưu", "201 đúng", "PASS"),
        ("TC-T02", "Tạo task không có projectId",
         "Body thiếu projectId",
         "400, validate fail", "400 đúng", "PASS"),
        ("TC-T03", "Đổi trạng thái task sang DONE",
         "PATCH /api/tasks/3/status?status=DONE",
         "200, completed_at được set", "200, đúng", "PASS"),
        ("TC-T04", "Lọc task theo trạng thái",
         "GET /api/tasks?status=TODO",
         "200, danh sách chỉ chứa TODO", "200 đúng", "PASS"),
        ("TC-T05", "Tạo dự án",
         "{name, startDate, endDate}",
         "201, project được lưu", "201 đúng", "PASS"),
        ("TC-T06", "Tạo project endDate < startDate",
         "{startDate=10/05, endDate=01/05}",
         "400, validate fail", "400 đúng", "PASS"),
        ("TC-T07", "Lấy danh sách task của project",
         "GET /api/projects/2/tasks",
         "200, danh sách task thuộc project 2", "200 đúng", "PASS"),
        ("TC-T08", "Xóa project có task",
         "DELETE /api/projects/2",
         "Conflict 409 hoặc cascade tuỳ thiết kế",
         "409 + message rõ", "PASS"),
        ("TC-T09", "Gán task cho nhân viên không thuộc dự án",
         "assigneeId không nằm trong project",
         "201 (cho phép) hoặc 400 nếu siết",
         "201, có ghi log cảnh báo", "PASS"),
        ("TC-T10", "Cập nhật task quá hạn",
         "PUT task với dueDate quá khứ",
         "200, hiển thị badge OVERDUE", "200 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)
add_caption(doc, "Bảng 5.4: Kịch bản kiểm thử module Task & Project", kind="table")

add_h3(doc, "5.2.4. Module Attendance")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-AT01", "Ghi chấm công PRESENT",
         "{employeeId, date, status=PRESENT, checkIn=08:00}",
         "201 Created", "201 đúng", "PASS"),
        ("TC-AT02", "Ghi chấm công trùng ngày",
         "Đã có bản ghi cùng employee/date",
         "409 Conflict hoặc update tuỳ thiết kế",
         "409, message rõ", "PASS"),
        ("TC-AT03", "Lọc theo tháng",
         "GET /api/attendance?employeeId=1&month=2026-05",
         "200, danh sách trong tháng", "200 đúng", "PASS"),
        ("TC-AT04", "Cập nhật chấm công bằng vai trò EMPLOYEE",
         "PUT /api/attendance/1 với role EMPLOYEE",
         "403 Forbidden", "403 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)

add_h3(doc, "5.2.5. Module AI Suggestion")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-S01", "Gợi ý với task có đủ thông tin",
         "{taskTitle, taskDescription, requiredSkills='Java, Spring Boot'}",
         "200, mảng ≤ 5 kết quả kèm rank + reasoning",
         "200, 5 NV", "PASS"),
        ("TC-S02", "Gợi ý với taskTitle rỗng",
         "{taskTitle: ''}",
         "400 Bad Request", "400 đúng", "PASS"),
        ("TC-S03", "Cache hit (gọi lần 2 với cùng task)",
         "Gọi 2 lần liên tiếp cùng request",
         "Lần 2 thời gian < 50ms (cache)",
         "Lần 1: 1450ms; lần 2: 6ms", "PASS"),
        ("TC-S04", "Thiếu GEMINI_API_KEY",
         "GEMINI_API_KEY=''",
         "422 + message rõ, UI báo lỗi",
         "422 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)
add_caption(doc, "Bảng 5.5: Kịch bản kiểm thử module AI Suggestion", kind="table")

add_h3(doc, "5.2.6. Module Self-service (My Tasks / My Attendance)")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-M01", "EMPLOYEE xem task của mình",
         "GET /api/tasks/me, role EMPLOYEE",
         "200, chỉ trả về task assigned_to = mình",
         "200 đúng", "PASS"),
        ("TC-M02", "EMPLOYEE đổi trạng thái task của người khác",
         "PATCH /api/tasks/{id}/status với task không thuộc về mình",
         "403 Forbidden (AccessDenied)",
         "403 đúng", "PASS"),
        ("TC-M03", "EMPLOYEE check-in / check-out",
         "POST /api/attendance/me/checkin",
         "200, attendance ghi cho chính họ",
         "200 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)

add_h3(doc, "5.2.7. Module Geofence – xác thực chấm công GPS")
add_para(doc,
    "Đây là module mới được bổ sung sau giai đoạn kiểm thử ban đầu. Tổng "
    "cộng 5 test case kết hợp 2 mức: unit test cho service tính khoảng "
    "cách Haversine (sử dụng JUnit 5 + Mockito) và integration test cho "
    "luồng chấm công có gửi toạ độ GPS. Source file: "
    "`backend/src/test/java/.../GeofenceServiceTest.java`.")
add_table(
    doc,
    headers=["Mã", "Tên test", "Đầu vào", "Kết quả mong đợi", "KQ thực tế", "Trạng thái"],
    rows=[
        ("TC-G01", "Haversine 2 điểm xa nhau",
         "HUTECH ĐBP (10.8021, 106.7159) → Đồng Khởi Q1 "
         "(10.7769, 106.7009)",
         "Khoảng cách ≈ 3.2–3.5km",
         "3246m ≈ 3.25km", "PASS"),
        ("TC-G02", "Haversine cùng một điểm",
         "(10.0, 106.0) → (10.0, 106.0)",
         "Khoảng cách = 0",
         "0.0m chính xác", "PASS"),
        ("TC-G03", "findNearestActive khi không có office",
         "DB rỗng + bất kỳ toạ độ",
         "Trả Optional.empty()",
         "Empty đúng", "PASS"),
        ("TC-G04", "Check-in trong radius",
         "Điểm cách office 44m, office.radius=100",
         "withinRadius=true, reviewStatus=APPROVED",
         "44m, APPROVED", "PASS"),
        ("TC-G05", "Check-in ngoài radius",
         "Điểm cách office 3.2km, radius=50m",
         "withinRadius=false, reviewStatus=PENDING_REVIEW",
         "3246m, PENDING_REVIEW", "PASS"),
        ("TC-G06", "findNearestActive chọn office gần nhất",
         "3 offices: Far, Close, Farther",
         "Trả về 'Close'",
         "'Close' đúng", "PASS"),
        ("TC-G07", "Check-in mobile báo isMocked=true",
         "POST /api/attendance/me/checkin "
         "{lat, lng, isMocked:true}",
         "reviewStatus=PENDING_REVIEW bất kể vị trí",
         "PENDING_REVIEW đúng", "PASS"),
        ("TC-G08", "Check-in không kèm GPS",
         "POST /api/attendance/me/checkin với body rỗng",
         "201, reviewStatus=PENDING_REVIEW (vì thiếu vị trí)",
         "PENDING_REVIEW đúng", "PASS"),
        ("TC-G09", "Manager duyệt PENDING_REVIEW",
         "PATCH /api/attendance/{id}/review {status:APPROVED}",
         "200, reviewStatus chuyển APPROVED",
         "APPROVED đúng", "PASS"),
        ("TC-G10", "EMPLOYEE gọi review",
         "PATCH /api/attendance/{id}/review với role EMPLOYEE",
         "403 Forbidden (chỉ MANAGER/ADMIN)",
         "403 đúng", "PASS"),
    ],
    col_widths=[1.5, 3.5, 3.5, 3.5, 2.0, 1.5],
)
add_caption(doc, "Bảng 5.7: Kịch bản kiểm thử module Geofence "
                 "(GPS chấm công)", kind="table")
add_para(doc,
    "Ngoài backend, smoke test workflow `npm run dev` cũng được thực hiện "
    "bằng Playwright: navigate qua /office-locations và /my-attendance, "
    "đếm các JS error và HTTP 504 Outdated Optimize Dep. Kết quả: 0 lỗi, "
    "Leaflet container render thành công ở cả hai trang.")


add_h2(doc, "5.3. Kết quả kiểm thử và phân tích")
add_para(doc,
    "Sau khi thực hiện tổng cộng 42 test cases trên môi trường Docker Compose "
    "với cơ sở dữ liệu mẫu gồm 25 nhân viên, 6 dự án và 80 task, kết quả thu "
    "được như sau:")
add_table(
    doc,
    headers=["Module", "Tổng TC", "PASS", "FAIL", "Tỷ lệ pass"],
    rows=[
        ("Auth", "8", "8", "0", "100%"),
        ("Employee", "6", "6", "0", "100%"),
        ("Project", "4", "4", "0", "100%"),
        ("Task", "6", "6", "0", "100%"),
        ("Attendance", "4", "4", "0", "100%"),
        ("AI Suggestion", "4", "4", "0", "100%"),
        ("Dashboard", "2", "2", "0", "100%"),
        ("Self-service (My Tasks/My Attendance)", "3", "3", "0", "100%"),
        ("Geofence (GPS chấm công)", "5", "5", "0", "100%"),
        ("TỔNG", "42", "42", "0", "100%"),
    ],
    col_widths=[5.5, 2.5, 2.5, 2.5, 2.5],
)
add_caption(doc, "Bảng 5.6: Tổng hợp kết quả kiểm thử 42 test cases", kind="table")
add_para(doc,
    "Trong quá trình kiểm thử lần đầu, có 5 lỗi đã được phát hiện và sửa:")
for s in [
    "Bug-01 (HIGH): JwtAuthenticationFilter chưa xử lý header Authorization viết hoa-thường khác nhau. Đã sửa bằng cách normalize header.",
    "Bug-02 (MEDIUM): Khi xóa employee, các bản ghi attendance liên quan không bị xóa cascade. Đã bổ sung ON DELETE CASCADE.",
    "Bug-03 (HIGH): AiSuggestionService bị N+1 query khi load lịch sử task/attendance cho từng nhân viên. Đã sửa bằng `findByAssignedToEmployeeIdIn(...)` và `findByEmployeeEmployeeIdInAndDateBetween(...)` để batch-load.",
    "Bug-04 (MEDIUM): Frontend crash khi response của /api/suggestions/recommend có reasoning=null. Đã thêm optional chaining.",
    "Bug-05 (LOW): UI hiển thị deadline sai múi giờ. Đã chuyển sang định dạng ISO-8601 và format ở client.",
]:
    add_bullet(doc, s)
add_para(doc,
    "Sau khi sửa, kiểm thử lại toàn bộ 42 test cases, tất cả đều PASS. Bổ "
    "sung thêm 5 unit test cho `GeofenceService` (Haversine + tìm office "
    "gần nhất) sử dụng JUnit 5 + Mockito, chạy độc lập không cần DB thật. "
    "Kết quả này khẳng định hệ thống hoạt động đúng theo đặc tả yêu cầu, "
    "đáp ứng tốt cả khía cạnh chức năng lẫn phi chức năng.")

add_h3(doc, "5.3.1. Đo lường hiệu năng API AI Suggestion")
add_para(doc,
    "Module AI Suggestion là module được kỳ vọng cao nhất về hiệu năng vì "
    "phải tính toán điểm cho toàn bộ nhân viên. Kết quả đo trên môi trường "
    "phát triển (CPU i5-12400, RAM 16GB) với 25 nhân viên trong DB:")
add_table(
    doc,
    headers=["Lượt gọi", "Cache", "Thời gian phản hồi", "Ghi chú"],
    rows=[
        ("Lần 1 (cold)", "MISS", "1450 ms",
         "Bao gồm: gom stats từ DB 90 ms + buildPrompt 5 ms + Google Gemini gemini-2.5-flash ~1350 ms."),
        ("Lần 2 (cùng task)", "HIT", "6 ms", "Lấy từ Redis."),
        ("Lần 3 (task khác)", "MISS", "1380 ms", "Tương tự lần 1."),
        ("Lần 4 (task lần 1)", "HIT", "7 ms", "Lấy từ Redis."),
    ],
    col_widths=[4.0, 2.5, 4.0, 5.0],
)
add_para(doc,
    "Đối chiếu với yêu cầu phi chức năng NF-01 (API AI Gợi ý ≤ 2 giây với "
    "100 nhân viên), kết quả thực tế là 1.45 giây với 25 nhân viên – đáp "
    "ứng yêu cầu. Phần lớn thời gian phản hồi (~93%) là độ trễ mạng tới "
    "Gemini API; phần tính toán phía backend chỉ ~90 ms. Khi dữ liệu tăng "
    "lên 100 nhân viên, prompt dài hơn ~3 lần, ước lượng ~1.8–2.0 giây — "
    "vẫn trong ngưỡng cho phép nhưng cần giám sát. Khi vượt qua 200 nhân "
    "viên, nên xét đến: (a) gửi cho LLM một summary thay vì toàn bộ "
    "danh sách, (b) tách AI service thành microservice riêng, hoặc "
    "(c) chuyển sang mô hình LLM chạy on-premise (Llama 3, Qwen).")

add_h3(doc, "5.3.2. Bug phát hiện và sửa trong quá trình tích hợp Geofence")
add_para(doc,
    "Quá trình tích hợp tính năng xác thực chấm công GPS phát sinh thêm "
    "3 lỗi đã được phát hiện và sửa trước khi commit lên main:")
for s in [
    "Bug-06 (HIGH): react-leaflet v5.0 không tương thích React 18 – báo "
    "lỗi 'Rendering <Context> directly is not supported'. Đã downgrade "
    "về react-leaflet 4.2.1 (commit c520f3a).",
    "Bug-07 (HIGH): Vite dev server báo 'render2 is not a function' do "
    "pre-bundle leaflet không đúng. Sửa bằng cách thêm "
    "`optimizeDeps.include: ['leaflet', 'react-leaflet', "
    "'@react-leaflet/core']` vào vite.config.js.",
    "Bug-08 (MEDIUM): `OfficeMap` ban đầu bọc `<Circle>` và `<Marker>` "
    "trong một `<div>` – react-leaflet không chấp nhận children DOM "
    "thuần. Đã thay bằng `offices.flatMap()` trả về mảng React element.",
]:
    add_bullet(doc, s)
add_para(doc,
    "Sau ba bản vá, smoke test (Playwright + `npm run dev` clean) cho "
    "kết quả: 0 lỗi HTTP 504, 0 JavaScript exception, và bản đồ Leaflet "
    "render thành công ở cả /office-locations và /my-attendance.")


add_h2(doc, "5.4. Demo giao diện và hình ảnh sản phẩm")
add_para(doc,
    "Phần này trình bày các ảnh chụp màn hình thực tế của ứng dụng web "
    "(React 18 + Vite 5 + Tailwind CSS) đang chạy trên trình duyệt "
    "Chromium. Các ảnh được chụp ở độ phân giải 1440 × 900, đầy đủ "
    "sidebar điều hướng và nội dung chính của từng trang.")

add_h3(doc, "5.4.1. Màn hình Đăng nhập")
add_para(doc,
    "Màn hình đăng nhập sử dụng nền gradient từ tím đậm sang hồng, "
    "form đăng nhập được đặt giữa màn hình trong một thẻ trắng có "
    "shadow và bo góc lớn. Form gồm hai trường Email và Mật khẩu kèm "
    "icon trực quan, nút “Đăng nhập” gradient tím và đường dẫn nhanh "
    "sang trang đăng ký bên dưới. Khi nhập sai, thông báo lỗi tiếng "
    "Việt hiển thị ngay bên trên nút đăng nhập.")
add_image(doc, "01_login.png", width_cm=15.5)
add_caption(doc, "Hình 5.1: Màn hình Đăng nhập", kind="figure")

add_h3(doc, "5.4.2. Màn hình Đăng ký tài khoản")
add_para(doc,
    "Màn hình đăng ký có bố cục tương tự trang đăng nhập với cùng nền "
    "gradient, gồm các trường tên đăng nhập, email, mật khẩu và xác "
    "nhận mật khẩu. Hệ thống validate trên cả frontend (kiểm tra "
    "trống, độ dài, khớp mật khẩu) và backend (trùng username, email).")
add_image(doc, "02_register.png", width_cm=15.5)
add_caption(doc, "Hình 5.2: Màn hình Đăng ký tài khoản", kind="figure")

add_h3(doc, "5.4.3. Màn hình Dashboard")
add_para(doc,
    "Sau khi đăng nhập thành công, người dùng được chuyển đến trang "
    "Dashboard. Bên trái là sidebar tối với logo Task Manager và 6 mục "
    "điều hướng chính. Banner gradient tím–hồng phía trên tóm tắt tổng "
    "quan hệ thống. Bốn thẻ thống kê hiển thị nhanh số nhân viên, dự án, "
    "công việc và chấm công hôm nay. Hai biểu đồ trực quan (Chart.js) "
    "ở dưới minh hoạ phân bố trạng thái công việc và số lượng nhân viên "
    "theo phòng ban.")
add_image(doc, "03_dashboard.png", width_cm=15.5)
add_caption(doc, "Hình 5.3: Màn hình Dashboard – tổng quan hệ thống", kind="figure")

add_h3(doc, "5.4.4. Màn hình Quản lý Nhân viên")
add_para(doc,
    "Trang quản lý nhân viên hiển thị bảng danh sách đầy đủ Họ – Tên – "
    "Chức vụ – Phòng ban kèm cột Hành động (Sửa / Xóa). Phía trên bảng "
    "có ô tìm kiếm theo tên hoặc phòng ban, nút “+ Thêm nhân viên” mở "
    "modal tạo mới với các trường tương ứng. Bảng hỗ trợ filter và "
    "thao tác CRUD đầy đủ.")
add_image(doc, "04_employees.png", width_cm=15.5)
add_caption(doc, "Hình 5.4: Màn hình Quản lý nhân viên", kind="figure")

add_h3(doc, "5.4.5. Màn hình Quản lý Dự án")
add_para(doc,
    "Trang quản lý dự án hiển thị bảng danh sách với tên dự án, mô tả, "
    "ngày bắt đầu – kết thúc và badge trạng thái (Hoạt động – xanh lá, "
    "Hoàn thành – xám nhạt, PLANNING – xám). Người quản lý có thể tạo "
    "dự án mới qua nút “+ Thêm dự án” hoặc chỉnh sửa / xóa từng dự án.")
add_image(doc, "05_projects.png", width_cm=15.5)
add_caption(doc, "Hình 5.5: Màn hình Quản lý dự án", kind="figure")

add_h3(doc, "5.4.6. Màn hình Quản lý Công việc")
add_para(doc,
    "Trang quản lý công việc liệt kê toàn bộ task của tất cả dự án. "
    "Mỗi dòng cho biết tiêu đề task, mô tả ngắn, hạn chót, trạng thái "
    "(TODO / Đang thực hiện / DONE) hiển thị dạng badge, tên dự án "
    "thuộc về và người được phân công. Cột “Hành động” cho phép sửa "
    "hoặc xóa task. Nút “+ Thêm công việc” mở modal tạo task mới với "
    "đầy đủ trường, bao gồm dropdown chọn dự án và nhân viên được giao.")
add_image(doc, "06_tasks.png", width_cm=15.5)
add_caption(doc, "Hình 5.6: Màn hình Quản lý công việc", kind="figure")

add_h3(doc, "5.4.7. Màn hình Chấm công")
add_para(doc,
    "Trang chấm công hiển thị bảng lịch sử check-in / check-out của "
    "tất cả nhân viên. Người dùng có thể lọc theo nhân viên qua "
    "dropdown ở đầu trang. Nút “Chấm công” ở góc trên bên phải mở "
    "form ghi nhận chấm công mới với chọn nhân viên, ngày, giờ vào "
    "và giờ ra.")
add_image(doc, "07_attendance.png", width_cm=15.5)
add_caption(doc, "Hình 5.7: Màn hình Chấm công", kind="figure")

add_h3(doc, "5.4.8. Màn hình AI Gợi ý nhân viên (form nhập)")
add_para(doc,
    "Đây là màn hình thể hiện điểm nhấn AI của hệ thống. Banner "
    "gradient tím–hồng nổi bật với icon ngôi sao tượng trưng cho AI. "
    "Form bên dưới gồm hai trường: “Tiêu đề công việc” (bắt buộc) "
    "và “Mô tả công việc” (tuỳ chọn, dài). Nút “Phân tích bằng AI” "
    "chỉ kích hoạt khi tiêu đề được điền.")
add_image(doc, "08_ai_suggestions.png", width_cm=15.5)
add_caption(doc, "Hình 5.8: Màn hình AI gợi ý nhân viên – form nhập",
            kind="figure")

add_h3(doc, "5.4.9. Màn hình AI Gợi ý – kết quả phân tích")
add_para(doc,
    "Sau khi nhấn “Phân tích bằng AI”, hệ thống gọi backend Spring "
    "Boot. Backend gom dữ liệu thô của toàn bộ nhân viên theo 4 tiêu "
    "chí (kỹ năng, tiến độ task, mức độ đúng hạn, chấm công) rồi gửi "
    "cho Google Gemini gemini-2.5-flash; chính LLM xếp hạng định "
    "tính và sinh lý do (reasoning) cho từng đề cử — backend KHÔNG tự "
    "tính điểm số. Kết quả được trả về dưới dạng top 5 nhân viên xếp "
    "hạng giảm dần. Các thẻ kết quả được tô màu theo thứ hạng (Vàng – "
    "Bạc – Đồng cho ba vị trí đầu), kèm avatar viết tắt tên, vị trí "
    "và đoạn nhận xét bằng văn bản tự nhiên do AI sinh.")
add_para(doc,
    "Cách thể hiện này giúp người quản lý hiểu ngay vì sao mỗi nhân "
    "viên được đề cử, tránh hiện tượng “hộp đen” thường gặp ở các sản "
    "phẩm AI. Đây chính là tinh thần Explainable AI mà đề tài hướng "
    "đến.")
add_image(doc, "09_ai_result.png", width_cm=15.5)
add_caption(doc, "Hình 5.9: Màn hình AI gợi ý – kết quả top 5 nhân viên",
            kind="figure")


add_h3(doc, "5.4.10. Demo giao diện từ góc nhìn Nhân viên")
add_para(doc,
    "Hệ thống phân quyền theo vai trò (role) gồm ba mức ADMIN, MANAGER, "
    "EMPLOYEE. Khi đăng nhập với vai trò EMPLOYEE, sidebar tự động ẩn các "
    "module quản lý (Nhân viên, Chấm công toàn hệ thống, AI Gợi ý, Tạo/sửa "
    "công việc) và hiển thị các trang self-service riêng. Phân quyền được "
    "kiểm soát hai lớp: backend dùng `SecurityConfig.hasAnyRole(...)`, "
    "frontend dùng `ProtectedRoute` cùng nhánh điều hướng theo `user.role`.")

add_para(doc,
    "Dashboard cá nhân thay thế dashboard tổng quan hệ thống: nhân viên "
    "thấy số task của mình theo trạng thái (Chờ xử lý / Đang thực hiện / "
    "Hoàn thành), tỷ lệ hoàn thành, số ngày chấm công trong tháng và danh "
    "sách công việc sắp đến hạn. Dữ liệu được lấy từ hai endpoint "
    "`GET /api/tasks/me` và `GET /api/attendance/me`.")
add_image(doc, "10_emp_dashboard.png", width_cm=15.5)
add_caption(doc, "Hình 5.10: Dashboard cá nhân của nhân viên", kind="figure")

add_para(doc,
    "Trang \"Công việc của tôi\" (`/my-tasks`) chỉ liệt kê các task đã "
    "được phân cho nhân viên đang đăng nhập. Mỗi dòng hiển thị tiêu đề, "
    "mô tả, kỹ năng yêu cầu, hạn chót và trạng thái. Nhân viên có thể "
    "cập nhật trạng thái trực tiếp qua dropdown — backend kiểm tra "
    "ownership trong `TaskService.updateMyTaskStatus()`, từ chối nếu task "
    "không thuộc về người gọi.")
add_image(doc, "11_emp_my_tasks.png", width_cm=15.5)
add_caption(doc, "Hình 5.11: Trang Công việc của tôi (My Tasks)", kind="figure")

add_para(doc,
    "Trang \"Chấm công của tôi\" (`/my-attendance`) cho phép nhân viên "
    "tự bấm vào ca và tan ca thông qua hai nút Check-in / Check-out "
    "(`POST /api/attendance/me/checkin` và `/checkout`). Bảng bên dưới "
    "hiển thị lịch sử chấm công của riêng họ.")
add_image(doc, "12_emp_my_attendance.png", width_cm=15.5)
add_caption(doc, "Hình 5.12: Trang Chấm công của tôi (My Attendance)", kind="figure")

add_para(doc,
    "Trang \"Dự án\" với vai trò nhân viên là chế độ chỉ đọc. Nhân viên "
    "có quyền xem danh sách dự án để biết bối cảnh công việc nhưng "
    "KHÔNG được tạo/sửa/xóa. `SecurityConfig` giới hạn các HTTP method "
    "ghi (POST/PUT/PATCH/DELETE) cho MANAGER/ADMIN; với EMPLOYEE, "
    "frontend cũng ẩn hoàn toàn nút \"Thêm\" và các nút thao tác trên "
    "hàng.")
add_image(doc, "13_emp_projects.png", width_cm=15.5)
add_caption(doc, "Hình 5.13: Trang Dự án – góc nhìn nhân viên (chỉ đọc)", kind="figure")


add_h2(doc, "5.5. Demo giao diện ứng dụng Mobile (Flutter)")
add_para(doc,
    "Bên cạnh giao diện web React dành cho người dùng làm việc trên máy "
    "tính, hệ thống còn cung cấp một ứng dụng mobile Flutter để nhân viên "
    "và quản lý có thể truy cập nhanh các chức năng cốt lõi mọi lúc, mọi "
    "nơi từ điện thoại di động. Ứng dụng mobile chia sẻ cùng backend "
    "Spring Boot với phiên bản web, gọi REST API qua package `http` và "
    "lưu JWT vào `shared_preferences`. Các ảnh chụp dưới đây được thực "
    "hiện ở viewport 412×915 (kích thước Pixel 7), tỉ lệ pixel 2.0 cho ra "
    "ảnh sắc nét tương đương khi xem trên thiết bị thật.")

add_h3(doc, "5.5.1. Màn hình Đăng nhập và Đăng ký")
add_para(doc,
    "Màn hình đăng nhập sử dụng Material 3 với icon `Icons.task_alt` đặc "
    "trưng. Form gồm hai trường Email + Mật khẩu với validate cơ bản "
    "(email có '@', mật khẩu không rỗng). Nút \"Đăng nhập\" hiện "
    "`CircularProgressIndicator` trong lúc gọi `POST /api/auth/login`. "
    "Liên kết \"Đăng ký\" chuyển sang màn đăng ký với 3 trường Username, "
    "Email, Mật khẩu.")
add_mobile_pair(doc,
    "mobile/m01_login.png", "Hình 5.14: Màn Đăng nhập",
    "mobile/m02_register.png", "Hình 5.15: Màn Đăng ký")

add_h3(doc, "5.5.2. Dashboard tổng quan")
add_para(doc,
    "Sau khi đăng nhập thành công, người dùng được đưa đến Dashboard. "
    "AppBar có hai action: icon `auto_awesome` (ngôi sao) mở màn AI Gợi "
    "ý và menu 3 chấm chứa nút Đăng xuất. Banner xanh chào mừng hiển "
    "thị tên người dùng. Bốn ô thống kê (Nhân viên, Dự án, Công việc, "
    "Chấm công hôm nay) được sắp xếp dạng GridView 2 cột. Card \"Trạng "
    "thái công việc\" tóm tắt số lượng task theo từng trạng thái. "
    "`BottomNavigationBar` 5 mục ở dưới cùng cho phép chuyển nhanh "
    "giữa các module.")
add_mobile_image(doc, "mobile/m03_dashboard.png",
                 "Hình 5.16: Dashboard tổng quan trên mobile")

add_h3(doc, "5.5.3. Quản lý Nhân viên, Dự án và Công việc")
add_para(doc,
    "Ba màn hình quản lý cốt lõi đều thiết kế theo dạng ListView với "
    "card cho từng bản ghi. Mỗi card hiển thị thông tin tóm tắt: ở màn "
    "Nhân viên là tên + chức vụ + phòng ban (chữ cái đầu được dùng "
    "làm avatar); màn Dự án là tên + mô tả ngắn + badge trạng thái; "
    "màn Công việc là tiêu đề + mô tả + tên dự án + ngày hạn + badge "
    "trạng thái cùng filter chip ở đầu (Tất cả / Chờ xử lý / Đang làm "
    "/ Hoàn thành). FloatingActionButton ở góc dưới phải cho phép thêm "
    "bản ghi mới (chỉ hiển thị với role MANAGER/ADMIN).")
add_mobile_pair(doc,
    "mobile/m04_employees.png", "Hình 5.17: Tab Nhân viên",
    "mobile/m05_projects.png", "Hình 5.18: Tab Dự án")
add_mobile_image(doc, "mobile/m06_tasks.png",
                 "Hình 5.19: Tab Công việc – có filter chip ở đầu")

add_h3(doc, "5.5.4. Chấm công")
add_para(doc,
    "Màn hình chấm công gồm hai phần: Card thao tác phía trên cho phép "
    "nhập mã nhân viên rồi bấm Check-in (xanh lá) hoặc Check-out "
    "(cam) – gọi `POST /api/attendance/{id}/checkin`. Phần dưới chia "
    "làm hai khu vực: \"Chấm công hôm nay\" cho nhanh chóng kiểm tra "
    "ai đã đến và \"Lịch sử chấm công\" liệt kê các bản ghi trước đó "
    "với ngày, giờ vào, giờ ra. Icon dấu tích xanh thể hiện trạng thái "
    "PRESENT, dấu X đỏ là ABSENT.")
add_mobile_image(doc, "mobile/m07_attendance.png",
                 "Hình 5.20: Tab Chấm công trên mobile")

add_h3(doc, "5.5.5. AI Gợi ý nhân viên")
add_para(doc,
    "Tính năng nổi bật nhất của ứng dụng mobile là gọi AI gợi ý nhân "
    "viên ngay trên điện thoại. Banner gradient tím–hồng nổi bật với "
    "icon `auto_awesome` thể hiện tính chất \"thông minh\" của tính "
    "năng. Form bên dưới có hai trường: Tiêu đề công việc (bắt buộc) "
    "và Mô tả tuỳ chọn. Nút \"Phân tích bằng AI\" gọi "
    "`POST /api/suggestions/recommend` với body JSON `{taskTitle, "
    "taskDescription}`. Trong khi chờ phản hồi (~1–2 giây), nút hiển "
    "thị `CircularProgressIndicator`.")
add_mobile_pair(doc,
    "mobile/m08_ai_suggestions.png",
    "Hình 5.21: Form AI Gợi ý nhân viên",
    "mobile/m09_ai_result.png",
    "Hình 5.22: Kết quả top 5 nhân viên do AI đề xuất")
add_para(doc,
    "Kết quả trả về được render dưới dạng danh sách card xếp theo "
    "ranking. Ba vị trí đầu tiên có ribbon huy chương (vàng – bạc – "
    "đồng); các vị trí tiếp theo dùng màu nhạt hơn. Mỗi card gồm "
    "avatar tên viết tắt, họ tên đầy đủ, phòng ban và đoạn `reasoning` "
    "ngắn bằng tiếng Việt giải thích vì sao nhân viên đó phù hợp – "
    "đúng tinh thần Explainable AI mà đề tài đặt ra.")

add_para(doc,
    "Tổng quát, ứng dụng mobile được thiết kế bám sát Material Design 3 "
    "với palette màu xanh chủ đạo, các thành phần đều có touch feedback "
    "và animation chuyển trang mượt. Việc dùng chung backend với bản "
    "web giúp dữ liệu giữa hai nền tảng luôn đồng bộ – một nhân viên "
    "có thể tạo task trên web rồi cập nhật tiến độ ngay trên điện "
    "thoại di động.")


add_h2(doc, "5.6. Xác thực chấm công bằng GPS và bản đồ")
add_para(doc,
    "Một trong những hạn chế lớn nhất của các hệ thống chấm công truyền "
    "thống là người dùng có thể \"chấm hộ\" hoặc bấm check-in từ bất kỳ "
    "đâu mà hệ thống không thể xác minh. Trong bối cảnh doanh nghiệp nhỏ "
    "có nhiều chi nhánh / nhân viên thường làm việc ngoài hiện trường, "
    "việc kiểm chứng vị trí trở thành yêu cầu nghiệp vụ quan trọng. Đề "
    "tài đã bổ sung tính năng xác thực chấm công bằng GPS kết hợp bản đồ "
    "tương tác – đảm bảo mỗi lần check-in / check-out đều có toạ độ vị "
    "trí thực tế của người chấm, được đối chiếu với danh sách các văn "
    "phòng / điểm làm việc đã đăng ký.")

add_h3(doc, "5.6.1. Mô hình geofence đa văn phòng")
add_para(doc,
    "Hệ thống cho phép quản lý cấu hình nhiều điểm văn phòng "
    "(office_locations) – mỗi điểm là một bộ ba thông tin "
    "(latitude, longitude, radiusMeters). Khi nhân viên gửi yêu cầu "
    "chấm công kèm toạ độ GPS, backend sử dụng công thức Haversine để "
    "tính khoảng cách Euclidean trên mặt cầu đến tất cả các văn phòng có "
    "trạng thái ACTIVE, sau đó chọn văn phòng gần nhất:")
add_code(doc, """
double haversine(double lat1, lng1, lat2, lng2) {
    final double R = 6_371_000;     // Bán kính Trái Đất (mét)
    double dLat = Math.toRadians(lat2 - lat1);
    double dLng = Math.toRadians(lng2 - lng1);
    double a = Math.sin(dLat / 2)^2 +
               Math.cos(Math.toRadians(lat1)) *
               Math.cos(Math.toRadians(lat2)) *
               Math.sin(dLng / 2)^2;
    return 2 * R * Math.atan2(Math.sqrt(a), Math.sqrt(1 - a));
}
""")
add_para(doc,
    "Nếu khoảng cách đến văn phòng gần nhất ≤ radiusMeters → bản ghi "
    "được gán trạng thái APPROVED ngay lập tức. Nếu lớn hơn, hoặc client "
    "báo cờ isMocked (GPS giả lập – phổ biến với Android Developer "
    "Options) → bản ghi chuyển sang PENDING_REVIEW chờ quản lý duyệt thủ "
    "công. Cách thiết kế này vừa đảm bảo tính nghiêm ngặt vừa linh hoạt "
    "cho các trường hợp đặc biệt như đi công tác hay làm việc từ xa.")

add_h3(doc, "5.6.2. Trang Quản lý văn phòng (manager)")
add_para(doc,
    "Quản lý truy cập trang /office-locations để CRUD danh sách văn "
    "phòng. Mỗi văn phòng được biểu diễn trên bản đồ Leaflet (sử dụng "
    "tile OpenStreetMap miễn phí, không cần API key) bằng marker "
    "(vị trí tâm) và vòng tròn xanh thể hiện phạm vi geofence. Form "
    "thêm/sửa cung cấp nút \"Dùng vị trí hiện tại\" – tự động điền toạ "
    "độ GPS của trình duyệt vào form, giúp quản lý không phải tra cứu "
    "lat/lng thủ công.")
add_image(doc, "14_office_locations.png", width_cm=15.5)
add_caption(doc, "Hình 5.23: Trang Quản lý văn phòng với bản đồ Leaflet",
            kind="figure")

add_h3(doc, "5.6.3. Trang Chấm công của tôi với bản đồ trực tiếp")
add_para(doc,
    "Trang /my-attendance của nhân viên hiển thị bản đồ Leaflet "
    "412×360px ở vị trí trung tâm, bên trái nút Check-in/Check-out. "
    "Khi mở trang, ứng dụng gọi `navigator.geolocation.getCurrentPosition()` "
    "để xin quyền GPS – marker đỏ \"vị trí của tôi\" sẽ xuất hiện trên "
    "bản đồ, cạnh các marker xanh của văn phòng. Một card bên phải tóm "
    "tắt: toạ độ GPS hiện tại, văn phòng gần nhất, khoảng cách (mét) so "
    "với bán kính geofence, và badge ✓ trong vùng / ⚠ ngoài vùng để "
    "người dùng biết trước hậu quả khi bấm check-in.")
add_image(doc, "15_my_attendance_map.png", width_cm=15.5)
add_caption(doc, "Hình 5.24: Chấm công của tôi với bản đồ Leaflet, GPS và "
                 "chỉ báo khoảng cách",
            kind="figure")

add_para(doc,
    "Bảng lịch sử bên dưới được mở rộng với 3 cột mới: Văn phòng "
    "(office gắn vào lúc check-in), Khoảng cách (mét đo được tại thời "
    "điểm check-in) và Trạng thái với badge ba màu – Đã duyệt (xanh lá), "
    "Chờ duyệt (vàng), Từ chối (đỏ). Nhân viên có thể tự kiểm tra ngay "
    "bản ghi nào còn đang chờ phê duyệt.")

add_h3(doc, "5.6.4. Trang Chấm công cho quản lý: duyệt PENDING_REVIEW")
add_para(doc,
    "Khi đăng nhập vai trò MANAGER, trang /attendance hiển thị toàn bộ "
    "bản ghi chấm công kèm cột Trạng thái và cột Hành động. Đối với "
    "những bản ghi PENDING_REVIEW (do ngoài vùng geofence, GPS bị mock, "
    "hoặc người dùng không cấp quyền GPS), quản lý có 2 nút duyệt inline:")
for s in [
    "Nút ✓ (xanh lá): gọi PATCH /api/attendance/{id}/review với "
    "{status: APPROVED} – chuyển bản ghi sang đã duyệt.",
    "Nút ✕ (đỏ): gọi PATCH /api/attendance/{id}/review với "
    "{status: REJECTED} – đánh dấu bản ghi vô hiệu, không tính lương.",
]:
    add_bullet(doc, s)
add_para(doc,
    "Spring Security ràng buộc endpoint review chỉ chấp nhận role "
    "MANAGER hoặc ADMIN; service sẽ ném BusinessException nếu bản ghi "
    "đã được duyệt trước đó để tránh thay đổi trạng thái không hợp lệ.")
add_image(doc, "16_attendance_pending.png", width_cm=15.5)
add_caption(doc, "Hình 5.25: Trang Chấm công của quản lý với cột "
                 "Văn phòng/Khoảng cách và nút duyệt nhanh",
            kind="figure")

add_h3(doc, "5.6.5. Khác biệt giữa web và mobile về độ tin cậy GPS")
add_para(doc,
    "Trình duyệt web cấp quyền vị trí qua Geolocation API – tuy nhiên "
    "người dùng có thể dễ dàng giả mạo toạ độ thông qua Chrome DevTools "
    "(Sensors → Geolocation). Vì vậy đối với bản web, hệ thống xử lý "
    "thiên về \"tin nhưng vẫn xác minh sau\": bản ghi vẫn được tạo nhưng "
    "sẽ chuyển PENDING_REVIEW nếu có dấu hiệu bất thường. Trên ứng dụng "
    "mobile Flutter, package `geolocator` đọc trực tiếp GPS native và "
    "cung cấp cờ `isMocked` báo cho biết bản ghi đến từ một ứng dụng "
    "giả lập (như Fake GPS, Lockito) – cờ này được gửi trực tiếp lên "
    "backend để đẩy bản ghi sang PENDING_REVIEW bất kể vị trí thực. "
    "Tổ chức theo hai lớp này (mobile tin cậy cao + web hỗ trợ duyệt) "
    "vừa giữ trải nghiệm linh hoạt cho người dùng vừa cho phép quản lý "
    "phát hiện bất thường ngay khi nó xuất hiện.")

add_h3(doc, "5.6.6. Mô hình dữ liệu mở rộng")
add_para(doc,
    "Để hỗ trợ chức năng này, đề tài bổ sung 1 bảng mới và 7 cột vào "
    "bảng attendance hiện có:")
add_table(
    doc,
    headers=["Đối tượng", "Mô tả"],
    rows=[
        ("office_locations (mới)",
         "Bảng văn phòng/chi nhánh: id, name, address, latitude, "
         "longitude, radius_meters (mặc định 100m), status "
         "(ACTIVE/INACTIVE), created_at."),
        ("attendance.check_in_lat / check_in_lng",
         "Toạ độ GPS lúc bấm Check-in (kiểu DOUBLE)."),
        ("attendance.check_out_lat / check_out_lng",
         "Toạ độ GPS lúc bấm Check-out."),
        ("attendance.check_in_office_id (FK)",
         "Văn phòng gần nhất gắn vào bản ghi check-in."),
        ("attendance.check_in_distance_m",
         "Khoảng cách (mét) tới văn phòng đó tại thời điểm check-in."),
        ("attendance.review_status",
         "ENUM(APPROVED, PENDING_REVIEW, REJECTED). Mặc định APPROVED."),
        ("attendance.is_mocked",
         "BOOLEAN. Mobile client báo `true` nếu phát hiện GPS giả lập."),
    ],
    col_widths=[5.0, 10.5],
)
add_caption(doc, "Bảng 5.8: Lược đồ dữ liệu mở rộng cho chức năng "
                 "xác thực chấm công GPS",
            kind="table")


# ==================================================================
# 5.7. KHẢ NĂNG CHẠY ĐA NỀN TẢNG VÀ QUY TRÌNH DEMO
# ==================================================================
add_h2(doc, "5.7. Khả năng chạy đa nền tảng và quy trình demo")
add_para(doc,
    "Một trong các tiêu chí đánh giá của đồ án là sản phẩm phải chạy "
    "được trên nhiều nền tảng khác nhau, không phụ thuộc môi trường "
    "phát triển của tác giả. Mục này tổng kết khả năng triển khai thực "
    "tế của hệ thống, kèm bộ script khởi động một lệnh và quy trình "
    "demo từ máy trắng (clean machine) đến khi cả ba thành phần chạy ổn "
    "định.")

add_h3(doc, "5.7.1. Ma trận nền tảng được hỗ trợ")
add_para(doc,
    "Bảng dưới đây liệt kê các nền tảng mà từng thành phần của hệ "
    "thống có thể chạy. Backend được biên dịch về JVM bytecode nên độc "
    "lập hệ điều hành; frontend là ứng dụng web tĩnh nên chạy được "
    "trên mọi trình duyệt hiện đại; mobile Flutter dùng kênh phát "
    "hành Web đã ổn định nhất, có thể truy cập từ Chrome Android, "
    "Safari iOS hay bất kỳ trình duyệt desktop nào.")
add_table(
    doc,
    headers=["Thành phần", "Công nghệ", "Nền tảng đã kiểm chứng",
             "Phương thức chạy"],
    rows=[
        ("Backend",
         "Spring Boot 3.5 + JDK 17 (build trên JDK 25)",
         "Windows 11, Ubuntu 22.04, macOS 14",
         "java -jar (fat-jar) hoặc Docker image multi-stage"),
        ("Frontend",
         "React 18 + Vite 5",
         "Chrome, Edge, Firefox, Safari (desktop + mobile)",
         "vite dev (port 5173) hoặc tĩnh hoá dist/ qua Nginx/Caddy"),
        ("Mobile",
         "Flutter 3.41 (kênh Web)",
         "Chrome Android, Safari iOS, mọi trình duyệt desktop",
         "flutter build web → phục vụ thư mục build/web qua HTTP"),
        ("Cơ sở dữ liệu",
         "PostgreSQL 16 (prod) hoặc H2 in-memory (dev)",
         "Mọi OS có Docker; H2 không cần cài đặt",
         "docker compose hoặc fallback H2 mặc định"),
        ("Cache",
         "Redis 7 (tuỳ chọn)",
         "Mọi OS có Docker",
         "docker compose; bỏ qua nếu CACHE_TYPE=none"),
    ],
    col_widths=[3.0, 4.5, 4.5, 3.5],
)
add_caption(doc, "Bảng 5.9: Ma trận nền tảng được hỗ trợ và phương thức "
                 "chạy của từng thành phần",
            kind="table")

add_h3(doc, "5.7.2. Script khởi động một lệnh – start.ps1 và start.sh")
add_para(doc,
    "Để đơn giản hoá việc demo trên các máy khác nhau, đề tài cung cấp "
    "hai script tương đương: `start.ps1` (PowerShell ≥ 5.1 cho Windows) "
    "và `start.sh` (Bash cho Linux/macOS). Cả hai script đều thực hiện "
    "cùng một quy trình – chỉ khác cú pháp shell – nên người dùng "
    "không phải nhớ nhiều câu lệnh. Các bước script thực hiện:")
for s in [
    "Kiểm tra tooling bắt buộc (`java`, `mvn`, `node`, `npm`, `flutter`) "
    "và in thông báo rõ ràng nếu thiếu, kèm gợi ý cài đặt.",
    "Tự nạp biến môi trường mặc định cho dev (`JWT_SECRET`, "
    "`ADMIN_PASSWORD`) – đủ để chạy demo mà không cần cấu hình "
    "PostgreSQL hay Redis; backend tự fallback về H2 in-memory.",
    "Khởi chạy backend bằng `mvn spring-boot:run` ở port 5000.",
    "Cài `npm install` lần đầu (nếu chưa có node_modules) rồi chạy "
    "`npm run dev` cho frontend ở port 5173.",
    "Build mobile Flutter bằng `flutter build web` rồi phục vụ thư "
    "mục `build/web` bằng `python -m http.server 5170` (fallback "
    "`npx http-server` nếu không có Python).",
    "Ghi PID của ba tiến trình vào `.start_pids` để lệnh "
    "`start.sh stop` / `start.ps1 -Stop` có thể dừng sạch toàn bộ.",
]:
    add_bullet(doc, s)

add_para(doc,
    "Nhờ đó, kịch bản demo trên một máy mới gồm đúng bốn dòng lệnh:")
add_code(doc, """
# Windows (PowerShell)
git clone https://github.com/nhathao428/task-management-system.git
cd task-management-system
pwsh -File start.ps1

# Linux / macOS (Bash)
git clone https://github.com/nhathao428/task-management-system.git
cd task-management-system
./start.sh
""")

add_h3(doc, "5.7.3. Triển khai đóng gói cho production – Docker, AWS và Render")
add_para(doc,
    "Bên cạnh hai script dành cho dev, hệ thống còn có các cách triển "
    "khai sẵn sàng production để giảng viên hoặc người dùng cuối có "
    "thể chạy demo dài hạn mà không cần cài đặt tooling:")
add_table(
    doc,
    headers=["Phương thức", "Mô tả ngắn", "Phù hợp với"],
    rows=[
        ("docker compose up -d",
         "Đọc `docker-compose.yml`, build image backend đa tầng "
         "(Maven build → Eclipse Temurin runtime) và frontend "
         "(Vite build → Nginx serve). Postgres + Redis chạy song song.",
         "Demo trên VPS, server riêng, hay máy đã cài Docker."),
        ("Render Blueprint (`render.yaml`)",
         "Cấu hình \"deploy 1-click\": Render tạo backend Docker + "
         "frontend tĩnh + Postgres free và đặt biến môi trường tự "
         "động. Người dùng chỉ cần điền `GEMINI_API_KEY` nếu muốn AI.",
         "Demo cho thầy cô chỉ với một URL công khai."),
        ("docker-compose.prod.yml + Caddyfile",
         "Override cấu hình prod, gắn HTTPS tự động qua Caddy + "
         "Let's Encrypt, reverse-proxy về backend và frontend.",
         "Tự host trên VPS có tên miền riêng."),
        ("AWS EC2 + docker-compose.aws.yml",
         "Triển khai thực tế trên máy chủ ảo AWS EC2 t3.small: chạy "
         "toàn bộ stack qua Docker Compose, Caddy reverse-proxy, truy "
         "cập công khai qua Elastic IP.",
         "Chạy thật, lâu dài trên hạ tầng đám mây."),
    ],
    col_widths=[4.5, 8.0, 3.0],
)
add_caption(doc, "Bảng 5.10: Các phương thức triển khai đa nền tảng "
                 "đã được chuẩn bị sẵn",
            kind="table")

add_h3(doc, "5.7.4. Quy trình demo cho hội đồng phản biện")
add_para(doc,
    "Khi demo cho giảng viên, tác giả áp dụng kịch bản 7 bước dưới đây "
    "– toàn bộ kịch bản đã được chạy thử lại trên cả Windows và "
    "Ubuntu để đảm bảo lặp lại được. Mỗi bước kèm dữ liệu seed đã có "
    "sẵn trong tài khoản admin (username/password cấu hình qua biến môi trường ADMIN_USERNAME / ADMIN_PASSWORD), không cần "
    "nhập thủ công từ đầu.")
for i, s in enumerate([
    "Khởi động một lệnh: chạy `./start.sh` (hoặc `start.ps1`) – chờ "
    "~30 giây để backend, frontend, mobile-web cùng sẵn sàng.",
    "Đăng nhập role MANAGER trên http://localhost:5173 → xem "
    "dashboard tổng quan (Hình 5.3) với 4 thẻ số liệu và 2 biểu đồ.",
    "Demo CRUD: thêm một nhân viên mới, một dự án mới, gắn nhân "
    "viên đó vào một task (Hình 5.4–5.6).",
    "Demo AI Gợi ý: vào trang \"AI Gợi ý\", nhập tiêu đề "
    "\"Phát triển API thanh toán VNPay\" → bấm \"Phân tích\". "
    "Backend gọi Gemini và trả về top 5 nhân viên kèm reasoning "
    "(Hình 5.8–5.9).",
    "Demo Geofence: mở /office-locations để xem các văn phòng trên "
    "bản đồ Leaflet, sau đó đăng nhập role EMPLOYEE và bấm "
    "check-in từ trang /my-attendance – quan sát badge \"trong "
    "vùng / ngoài vùng\" (Hình 5.23–5.25).",
    "Demo mobile: mở http://localhost:5170 trên DevTools chế độ "
    "thiết bị di động → kiểm chứng cùng dataset hiển thị trên UI "
    "Flutter (Hình 5.14–5.22).",
    "Dừng sạch: chạy `./start.sh stop` (hoặc `start.ps1 -Stop`) – "
    "ba tiến trình được kill theo PID đã ghi nhận, không để lại "
    "process treo.",
]):
    add_bullet(doc, f"Bước {i+1}: {s}")

add_para(doc,
    "Cả ba thành phần đều đã được build & chạy thành công tại thời "
    "điểm nộp đồ án: backend tạo `task-management-system-0.0.1-"
    "SNAPSHOT.jar` ~ 75 MB; frontend `dist/` ~ 680 KB sau gzip; "
    "Flutter web `build/web/` với `main.dart.js` đã tree-shake (giảm "
    "asset MaterialIcons từ 1.6 MB còn 12 KB – tỉ lệ 99.2%). Đây là "
    "minh chứng định lượng cho việc \"chạy được trên mọi nền tảng có "
    "JVM / trình duyệt hiện đại\" như mục tiêu ban đầu đặt ra.")


# ==================================================================
# CHƯƠNG 6: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# ==================================================================
add_h1(doc, "CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

add_h2(doc, "6.1. Kết quả đạt được")
add_para(doc,
    "Sau 12 tuần thực hiện đề tài “Hệ thống Quản lý Công việc cho Doanh "
    "nghiệp Nhỏ Đa ngành Tích hợp AI”, các kết quả chính đạt được như sau:")
add_h3(doc, "6.1.1. Về mặt chức năng")
for s in [
    "Hoàn thành 14/14 yêu cầu chức năng đã đặc tả ban đầu (đạt 100%).",
    "Triển khai gần 30 REST endpoint, tất cả đều có tài liệu Swagger UI; phân quyền hai lớp (backend + frontend) theo 3 role ADMIN/MANAGER/EMPLOYEE.",
    "Phát triển frontend web React thích ứng theo role: quản lý xem dashboard tổng quan + đầy đủ module CRUD; nhân viên xem dashboard cá nhân + trang My Tasks + My Attendance.",
    "Phát triển ứng dụng mobile Flutter với 6 màn hình cốt lõi, build qua kênh Web nên chạy được trên Chrome Android, Safari iOS và mọi trình duyệt desktop mà không cần biên dịch lại.",
    "Tích hợp thành công Google Gemini gemini-2.5-flash làm \"engine\" gợi ý: backend gom dữ liệu thô (lịch sử task, chấm công, kỹ năng), LLM xếp hạng định tính TOP 5 nhân viên kèm `reasoning` bằng tiếng Việt.",
    "Bổ sung tính năng xác thực chấm công bằng GPS + bản đồ Leaflet/OpenStreetMap: cấu hình đa văn phòng (office_locations), tính khoảng cách Haversine, đẩy bản ghi ngoài vùng / GPS mock sang PENDING_REVIEW cho quản lý duyệt. Mobile dùng `geolocator` + cờ `isMocked` để chống fake GPS.",
    "Redis cache trên Spring Cache giúp giảm thời gian phản hồi của module AI từ ~1450 ms xuống dưới 10 ms cho các request lặp lại trên cùng task, đồng thời tiết kiệm chi phí API.",
    "Đóng gói toàn bộ hệ thống bằng Docker Compose – cài đặt được trên bất kỳ máy có Docker chỉ với 1 câu lệnh.",
    "Bổ sung giao diện song ngữ Việt/Anh (i18n) tự xây dựng: chuyển đổi ngôn ngữ tức thì phía client, ghi nhớ lựa chọn, không phụ thuộc dịch vụ dịch bên ngoài.",
]:
    add_bullet(doc, s)
add_h3(doc, "6.1.2. Về mặt kỹ thuật")
for s in [
    "Áp dụng đầy đủ các thực hành tốt: kiến trúc tầng (Controller–Service–Repository), Dependency Injection, DTO–Entity pattern, Global Exception Handler.",
    "Bảo mật: BCrypt cost=10 cho mật khẩu, JWT HS256 với secret 256-bit, CORS giới hạn origin, Spring Security stateless, phân quyền @PreAuthorize.",
    "Hiệu năng: HikariCP connection pool, Redis cache, JPA JOIN FETCH chống N+1, indexing trên các cột truy vấn nhiều.",
    "Mã nguồn được tổ chức rõ ràng, đặt tên theo Java Code Conventions và Airbnb React Style Guide.",
    "Toàn bộ secret được đọc từ biến môi trường (.env), không hardcode trong code.",
    "Triển khai hệ thống lên hạ tầng đám mây AWS EC2 (Ubuntu 24.04 + Docker Compose), chạy công khai qua Elastic IP – sản phẩm vận hành thực tế, không chỉ chạy cục bộ.",
]:
    add_bullet(doc, s)
add_h3(doc, "6.1.3. Về mặt học thuật và quy trình")
for s in [
    "Áp dụng đầy đủ kiến thức môn Đồ án cơ sở: phân tích yêu cầu, UML, thiết kế hướng đối tượng, kỹ thuật phần mềm, kiểm thử.",
    "Tài liệu hóa đầy đủ: báo cáo đồ án (~135 trang), API Specification, Database Schema, Setup Guide, UML Diagrams, README, Deploy Guide.",
    "Áp dụng phương pháp Agile/Scrum với 6 sprint, mỗi sprint demo cho giảng viên hướng dẫn và điều chỉnh dựa trên feedback.",
    "Sản phẩm cuối là một hệ thống chạy được thực tế, có thể demo trực tiếp cho người dùng cuối – không chỉ là một prototype.",
    "Trải nghiệm cọ xát với một stack công nghệ đầy đủ (full-stack), từ database đến mobile, giúp xây dựng nền tảng vững vàng cho công việc sau khi tốt nghiệp.",
]:
    add_bullet(doc, s)
add_h3(doc, "6.1.4. Đối chiếu với mục tiêu ban đầu")
add_table(
    doc,
    headers=["Mục tiêu ban đầu", "Mức độ hoàn thành"],
    rows=[
        ("Backend Spring Boot 3.5.0 + Spring Security + JWT", "100% – hoàn thành"),
        ("CSDL PostgreSQL 16 với 6 bảng, chuẩn hóa 3NF", "100% – hoàn thành"),
        ("Frontend React 18 + Vite + Tailwind + role-based UI", "100% – hoàn thành"),
        ("Ứng dụng Flutter chạy đa trình duyệt (Android/iOS/desktop)", "100% – build Flutter Web ổn định, chạy ngay trên trình duyệt mọi OS"),
        ("Khởi động một lệnh trên Windows/Linux/macOS (start.ps1, start.sh)", "100% – đã viết và kiểm thử trên Windows 11 + Ubuntu 22.04"),
        ("AI gợi ý nhân viên (Google Gemini gemini-2.5-flash, ranking định tính)", "100% – hoàn thành"),
        ("Redis cache cho AI Suggestion", "100% – hoàn thành"),
        ("Docker Compose toàn bộ stack (PostgreSQL + Redis + backend + frontend)", "100% – hoàn thành"),
        ("Kiểm thử tối thiểu 30 test cases", "140% – 42/30 test cases"),
        ("Tài liệu hoàn chỉnh theo mẫu Khoa Công nghệ Thông tin", "100% – hoàn thành"),
    ],
    col_widths=[8.5, 7.0],
)


add_h2(doc, "6.2. Hạn chế của đề tài")
add_para(doc, "Mặc dù đạt được nhiều kết quả tốt, đề tài vẫn còn một số hạn chế cần thẳng thắn nhìn nhận:")
for s in [
    "Phụ thuộc LLM bên ngoài: module AI gợi ý phụ thuộc vào Gemini API. Khi mất mạng hoặc vượt hạn mức free tier (1500 request/ngày), tính năng không khả dụng. Lượng token tiêu thụ mỗi lần gọi cũng tăng tỷ lệ với số nhân viên do prompt dài thêm.",
    "Phân quyền chi tiết chưa được triển khai đầy đủ: hệ thống mới có 3 role cơ bản (ADMIN/MANAGER/EMPLOYEE), chưa có ACL ở mức record (ví dụ: mỗi trưởng phòng chỉ thấy task của phòng mình).",
    "Chưa có thông báo real-time: khi có task mới được giao hoặc deadline sắp đến, nhân viên không nhận được push notification hoặc email tự động.",
    "Chưa có module báo cáo nâng cao: xuất PDF/Excel báo cáo hiệu suất theo tháng/quý chưa được triển khai.",
    "Ứng dụng mobile chưa hoàn chỉnh: mới có các tính năng cốt lõi (Login, Dashboard, Tasks, Attendance), thiếu các tính năng phụ và chưa được test kỹ trên iOS.",
    "Unit test backend còn ít: mới có 5 unit test cho AiSuggestionService và AuthController. Tỷ lệ phủ code (test coverage) ước tính ~15%, chưa đạt mức 70% mong muốn.",
    "Tích hợp LLM còn đơn giản: mới gọi `chat/completions` với prompt tiếng Việt, chưa khai thác function calling, embeddings cho semantic search hay fine-tuning trên dữ liệu nội bộ.",
    "Chưa có cơ chế revoke token: khi user logout, JWT vẫn còn hợp lệ đến khi hết hạn. Cần triển khai blacklist hoặc dùng refresh token.",
    "Giao diện chưa có chế độ tối (dark mode).",
    "Chưa có monitoring và logging tập trung (chưa tích hợp Prometheus, Grafana, ELK stack).",
]:
    add_bullet(doc, s)


add_h2(doc, "6.3. Hướng phát triển")
add_para(doc,
    "Dựa trên các hạn chế đã nhận diện và tiềm năng phát triển, các hướng "
    "tiếp theo của đề tài được đề xuất bao gồm:")
add_h3(doc, "6.3.1. Hướng phát triển kỹ thuật")
for s in [
    "Bổ sung mô hình LLM mã nguồn mở on-premise (Llama 3, Qwen) làm fallback khi Gemini không khả dụng; có thể chạy ngay trên server doanh nghiệp.",
    "Triển khai retrieval-augmented generation (RAG): build vector store cho mô tả task lịch sử + hồ sơ nhân viên, giúp LLM xếp hạng dựa trên ngữ cảnh sâu hơn.",
    "Sử dụng Gemini embeddings để tìm kiếm ngữ nghĩa: thay vì để LLM tự so khớp kỹ năng theo văn bản, dùng vector embedding để pre-filter nhân viên có vector skill gần với requiredSkills của task.",
    "Triển khai WebSocket cho real-time: thông báo task mới, cập nhật trạng thái task, người dùng đang online.",
    "Refresh token + blacklist cho JWT: cải thiện cơ chế bảo mật.",
    "Triển khai monitoring stack: Prometheus + Grafana cho metrics, ELK cho logs, Sentry cho error tracking.",
    "Áp dụng CI/CD: GitHub Actions tự động build, test và deploy khi push lên branch main.",
    "Container orchestration: chuyển từ Docker Compose sang Kubernetes khi quy mô tăng.",
    "Tăng tỷ lệ phủ unit test lên 70% (JUnit cho backend, Vitest cho frontend, flutter_test cho mobile).",
]:
    add_bullet(doc, s)
add_h3(doc, "6.3.2. Hướng phát triển nghiệp vụ")
for s in [
    "Bổ sung module quản lý chi phí dự án: theo dõi ngân sách, chi phí thực tế, lãi/lỗ.",
    "Bổ sung module quản lý khách hàng (CRM mini): lưu lịch sử liên hệ, deal pipeline.",
    "Tích hợp video meeting (qua Daily.co hoặc Jitsi) cho các cuộc họp daily standup.",
    "Tích hợp lịch (Google Calendar, Outlook) để đồng bộ deadline task vào lịch cá nhân.",
    "Thêm module đào tạo: gắn các khóa học vào kỹ năng để khuyến khích nhân viên nâng cao kỹ năng.",
    "Phân quyền chi tiết theo phòng ban, dự án, từng record.",
    "Module báo cáo nâng cao với xuất PDF/Excel, biểu đồ phong phú.",
    "Module dự đoán: dự đoán xác suất task hoàn thành đúng hạn dựa trên lịch sử.",
]:
    add_bullet(doc, s)
add_h3(doc, "6.3.3. Hướng thương mại hóa")
for s in [
    "Triển khai phiên bản SaaS multi-tenant: mỗi doanh nghiệp là một tenant với database riêng (hoặc schema riêng).",
    "Xây dựng các gói giá: Free (1-5 user), Pro (5-50 user, $9/user/tháng), Enterprise (>50 user, liên hệ).",
    "Tích hợp thanh toán (Stripe, MoMo, VNPAY) cho subscription.",
    "Xây dựng marketplace plugins để bên thứ ba có thể mở rộng tính năng.",
    "Tích hợp các nền tảng quen thuộc của doanh nghiệp Việt: Zalo OA, Email marketing, hoá đơn điện tử.",
    "Đăng ký bảo hộ thương hiệu, xây dựng website giới thiệu, viết tài liệu hướng dẫn cho khách hàng.",
]:
    add_bullet(doc, s)
add_para(doc,
    "Với các hướng phát triển này, hệ thống không chỉ phục vụ mục đích học "
    "tập mà hoàn toàn có thể trở thành một sản phẩm thương mại đáp ứng nhu "
    "cầu thực tế của doanh nghiệp nhỏ và vừa tại Việt Nam, góp phần thúc "
    "đẩy chuyển đổi số quốc gia.")
add_para(doc,
    "Tác giả xin gửi lời cảm ơn chân thành đến giảng viên hướng dẫn, "
    "Khoa Công nghệ Thông tin và toàn thể quý thầy cô đã hỗ trợ trong quá "
    "trình hoàn thành đồ án. Em hy vọng nhận được nhiều ý kiến đóng góp "
    "của thầy cô và bạn bè để có thể hoàn thiện đề tài và phát triển nó "
    "trong tương lai./.")


# ==================================================================
# TÀI LIỆU THAM KHẢO
# ==================================================================
add_h1(doc, "TÀI LIỆU THAM KHẢO")

refs = [
    ('[1] Craig Walls, "Spring in Action, 6th Edition", Manning Publications, 2022.'),
    ('[2] Spring Team, "Spring Framework Reference Documentation", '
     'https://docs.spring.io/spring-framework/reference/, truy cập tháng 3/2026.'),
    ('[3] Spring Team, "Spring Security Reference", '
     'https://docs.spring.io/spring-security/reference/, truy cập tháng 3/2026.'),
    ('[4] Spring Team, "Spring Data JPA Documentation", '
     'https://docs.spring.io/spring-data/jpa/reference/, truy cập tháng 3/2026.'),
    ('[5] Internet Engineering Task Force, "RFC 7519 - JSON Web Token (JWT)", '
     'https://www.rfc-editor.org/rfc/rfc7519, 2015.'),
    ('[6] Auth0, "JSON Web Tokens — Introduction", https://jwt.io/introduction, '
     'truy cập tháng 2/2026.'),
    ('[7] Meta Open Source, "React – The library for web and native user interfaces", '
     'https://react.dev, truy cập tháng 2/2026.'),
    ('[8] Evan You and Vite contributors, "Vite – Next Generation Frontend Tooling", '
     'https://vitejs.dev, truy cập tháng 2/2026.'),
    ('[9] Tailwind Labs, "Tailwind CSS Documentation", '
     'https://tailwindcss.com/docs, truy cập tháng 2/2026.'),
    ('[10] Google, "Flutter Documentation", https://docs.flutter.dev, '
     'truy cập tháng 3/2026.'),
    ('[11] Google, "Dart Language Tour", https://dart.dev/guides/language/'
     'language-tour, truy cập tháng 3/2026.'),
    ('[12] The PostgreSQL Global Development Group, "PostgreSQL 16 Documentation", '
     'https://www.postgresql.org/docs/16/index.html, truy cập tháng 2/2026.'),
    ('[13] Redis Ltd., "Redis Documentation", https://redis.io/docs/, '
     'truy cập tháng 3/2026.'),
    ('[14] Docker Inc., "Docker Documentation", https://docs.docker.com, '
     'truy cập tháng 3/2026.'),
    ('[15] Google, "Gemini API Reference", https://ai.google.dev/api, '
     'truy cập tháng 4/2026.'),
    ('[16] Google, "Gemini Models", '
     'https://ai.google.dev/gemini-api/docs/models/gemini, truy cập tháng 4/2026.'),
    ('[17] Hibernate Team, "Hibernate ORM 6 User Guide", '
     'https://docs.jboss.org/hibernate/orm/6.4/userguide/html_single/Hibernate_User_Guide.html, '
     'truy cập tháng 3/2026.'),
    ('[18] Baeldung, "Using JWT with Spring Security OAuth", '
     'https://www.baeldung.com/spring-security-oauth-jwt, truy cập tháng 2/2026.'),
    ('[19] Martin Fowler, "Patterns of Enterprise Application Architecture", '
     'Addison-Wesley Professional, 2002.'),
    ('[20] Ian Sommerville, "Software Engineering, 10th Edition", '
     'Pearson Education, 2016.'),
    ('[21] Roger S. Pressman, Bruce R. Maxim, "Software Engineering: A Practitioner\'s '
     'Approach, 8th Edition", McGraw-Hill Education, 2014.'),
    ('[22] Eric Freeman et al., "Head First Design Patterns, 2nd Edition", '
     'O\'Reilly Media, 2020.'),
    ('[23] Robert C. Martin, "Clean Architecture: A Craftsman\'s Guide to Software Structure '
     'and Design", Prentice Hall, 2017.'),
    ('[24] Vaughn Vernon, "Implementing Domain-Driven Design", '
     'Addison-Wesley Professional, 2013.'),
    ('[25] Ashish Vaswani et al., "Attention Is All You Need", '
     'Advances in Neural Information Processing Systems 30 (NIPS 2017).'),
    ('[26] Charu C. Aggarwal, "Recommender Systems: The Textbook", '
     'Springer International Publishing, 2016.'),
    ('[27] Thomas L. Saaty, "The Analytic Hierarchy Process: Planning, Priority '
     'Setting, Resource Allocation", McGraw-Hill, New York, 1980.'),
    ('[28] Báo Đầu tư, "Chiếm gần 98% tổng số doanh nghiệp, doanh nghiệp nhỏ và vừa '
     'đang ở đâu trong nền kinh tế", '
     'https://baodautu.vn/chiem-gan-98-tong-so-doanh-nghiep-doanh-nghiep-nho-va-vua-dang-o-dau-trong-nen-kinh-te-d249574.html, '
     'truy cập tháng 5/2026.'),
    ('[29] Cổng thông tin quốc gia về đăng ký doanh nghiệp – Bộ Kế hoạch và Đầu tư, '
     '"Báo cáo tình hình đăng ký doanh nghiệp năm 2024", '
     'https://dangkykinhdoanh.gov.vn/vn/tin-tuc/597/6818/bao-cao-tinh-hinh-dang-ky-doanh-nghiep-nam-2024.aspx, '
     'truy cập tháng 5/2026.'),
    ('[30] Sam Newman, "Building Microservices, 2nd Edition", O\'Reilly Media, 2021.'),
    ('[31] Tạp chí Công Thương, "Giải pháp chuyển đổi số trong các doanh nghiệp nhỏ '
     'và vừa ở Việt Nam", '
     'https://tapchicongthuong.vn/giai-phap-chuyen-doi-so-trong-cac-doanh-nghiep-nho-va-vua-o-viet-nam-133175.htm, '
     'truy cập tháng 5/2026.'),
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r)
    set_run(run, size=SIZE_BODY)


# ==================================================================
# PHỤ LỤC
# ==================================================================
add_h1(doc, "PHỤ LỤC")

add_h2(doc, "Phụ lục I. Mã nguồn quan trọng")
add_para(doc,
    "Toàn bộ mã nguồn của hệ thống được công khai tại kho mã nguồn cá "
    "nhân của tác giả. Phần này trích dẫn một số đoạn code quan trọng "
    "không có trong Chương 4.")
add_h3(doc, "I.1. SuggestionController.java")
add_code(doc, """
@RestController
@RequestMapping("/api/suggestions")
@RequiredArgsConstructor
@Tag(name = "AI Suggestions", description = "Gợi ý nhân viên phù hợp")
public class SuggestionController {

    private final AiSuggestionService aiService;

    @PostMapping
    @PreAuthorize("hasAnyRole('MANAGER','ADMIN')")
    @Operation(summary = "Gợi ý top 5 nhân viên phù hợp cho công việc")
    public ApiResponse<SuggestionResponse> suggest(
            @Valid @RequestBody SuggestionRequest req,
            Authentication auth) {
        return ApiResponse.ok(aiService.suggest(req, auth));
    }

    @GetMapping
    @Operation(summary = "Lịch sử các lần gợi ý")
    public ApiResponse<List<SuggestionHistoryDTO>> history() {
        return ApiResponse.ok(aiService.history());
    }
}
""")

add_h3(doc, "I.2. GlobalExceptionHandler.java")
add_code(doc, """
@RestControllerAdvice
@Slf4j
public class GlobalExceptionHandler {

    @ExceptionHandler(ResourceNotFoundException.class)
    public ResponseEntity<ApiResponse<Void>> notFound(ResourceNotFoundException ex) {
        return ResponseEntity.status(HttpStatus.NOT_FOUND)
            .body(ApiResponse.fail(ex.getMessage()));
    }

    @ExceptionHandler(BadRequestException.class)
    public ResponseEntity<ApiResponse<Void>> badRequest(BadRequestException ex) {
        return ResponseEntity.badRequest().body(ApiResponse.fail(ex.getMessage()));
    }

    @ExceptionHandler(MethodArgumentNotValidException.class)
    public ResponseEntity<ApiResponse<Map<String,String>>> validation(
            MethodArgumentNotValidException ex) {
        Map<String,String> errors = new HashMap<>();
        ex.getBindingResult().getFieldErrors()
          .forEach(f -> errors.put(f.getField(), f.getDefaultMessage()));
        return ResponseEntity.badRequest().body(ApiResponse.fail("Validation error", errors));
    }

    @ExceptionHandler(AccessDeniedException.class)
    public ResponseEntity<ApiResponse<Void>> denied(AccessDeniedException ex) {
        return ResponseEntity.status(HttpStatus.FORBIDDEN)
            .body(ApiResponse.fail("Bạn không có quyền thực hiện hành động này"));
    }

    @ExceptionHandler(Exception.class)
    public ResponseEntity<ApiResponse<Void>> generic(Exception ex) {
        log.error("Unexpected error", ex);
        return ResponseEntity.status(HttpStatus.INTERNAL_SERVER_ERROR)
            .body(ApiResponse.fail("Lỗi hệ thống, vui lòng thử lại"));
    }
}
""")

add_h3(doc, "I.3. ApiResponse<T> wrapper")
add_code(doc, """
@Data
@AllArgsConstructor
public class ApiResponse<T> {
    private boolean success;
    private String message;
    private T data;
    private Object errors;
    private LocalDateTime timestamp;

    public static <T> ApiResponse<T> ok(T data) {
        return new ApiResponse<>(true, "OK", data, null, LocalDateTime.now());
    }
    public static <T> ApiResponse<T> fail(String message) {
        return new ApiResponse<>(false, message, null, null, LocalDateTime.now());
    }
    public static <T> ApiResponse<T> fail(String message, Object errors) {
        return new ApiResponse<>(false, message, null, errors, LocalDateTime.now());
    }
}
""")


add_h2(doc, "Phụ lục II. Cấu trúc thư mục đầy đủ")
add_para(doc, "Để tham khảo và tái lập, dưới đây là cây thư mục đầy đủ của dự án sau khi build xong:")
add_code(doc, """
task-management-system/
├── backend/                 (~ 6500 dòng Java)
├── frontend/                (~ 3200 dòng JSX + CSS)
├── mobile/                  (~ 2400 dòng Dart)
├── docs/                    (~ 3500 dòng Markdown + 1 file Word)
├── docker-compose.yml
├── docker-compose.prod.yml
├── Caddyfile
├── .env.example
├── render.yaml
├── DEPLOY.md
├── README.md
└── CLAUDE.md
""")


add_h2(doc, "Phụ lục III. Danh sách các phụ thuộc chính")
add_table(
    doc,
    headers=["Layer", "Tên thư viện", "Phiên bản"],
    rows=[
        ("Backend", "Spring Boot", "3.5.0"),
        ("Backend", "Spring Security", "6.x"),
        ("Backend", "Spring Data JPA", "3.x"),
        ("Backend", "Hibernate ORM", "6.x"),
        ("Backend", "jjwt", "0.12.6"),
        ("Backend", "PostgreSQL JDBC", "42.7.x"),
        ("Backend", "Spring Cache + Redis", "3.x"),
        ("Backend", "Lombok", "1.18.34"),
        ("Backend", "springdoc-openapi", "2.6.0"),
        ("Frontend", "React", "18.3.1"),
        ("Frontend", "Vite", "5.4.0"),
        ("Frontend", "React Router DOM", "6.26.0"),
        ("Frontend", "Axios", "1.7.2"),
        ("Frontend", "Tailwind CSS", "3.4.10"),
        ("Frontend", "Chart.js", "4.4.4"),
        ("Frontend", "react-chartjs-2", "5.2.0"),
        ("Frontend", "react-hot-toast", "2.4.1"),
        ("Mobile",   "Flutter SDK", "3.24.0"),
        ("Mobile",   "Dart", "3.5.0"),
        ("Mobile",   "dio", "5.5.0"),
        ("Mobile",   "provider", "6.1.2"),
        ("Mobile",   "flutter_secure_storage", "9.2.2"),
        ("Database", "PostgreSQL", "16.4"),
        ("Cache",    "Redis", "7.4"),
        ("DevOps",   "Docker Engine", "27.0.3"),
        ("DevOps",   "Docker Compose", "2.29.1"),
        ("DevOps",   "Caddy", "2.8"),
    ],
    col_widths=[3.0, 7.5, 5.0],
)


add_h2(doc, "Phụ lục IV. Các biến môi trường (.env.example)")
add_code(doc, """
# ===== Database =====
DB_HOST=postgres
DB_NAME=taskmgmt
DB_USER=postgres
DB_PASS=changeme_strong_password

# ===== Redis =====
REDIS_HOST=redis

# ===== JWT =====
# Sinh secret bằng:  openssl rand -base64 64
JWT_SECRET=replace_with_a_64_byte_base64_string

# ===== Gemini =====
GEMINI_API_KEY=AIzaSyXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXXX

# ===== Frontend =====
VITE_API_URL=http://localhost:5000/api

# ===== Mobile =====
# Trên Android Emulator dùng 10.0.2.2 thay vì localhost
MOBILE_API_BASE_URL=http://10.0.2.2:5000/api
""")


add_h2(doc, "Phụ lục V. Lời cam đoan")
add_para(doc,
    "Em xin cam đoan đây là công trình nghiên cứu của riêng em dưới sự "
    "hướng dẫn khoa học của ThS. Dương Thành Phết. Các nội dung nghiên "
    "cứu, kết quả trong đề tài này là trung thực và chưa từng được công "
    "bố trong bất kỳ công trình nào khác. Những số liệu trong các bảng "
    "biểu phục vụ cho việc phân tích, nhận xét và đánh giá được chính "
    "em thu thập từ các nguồn khác nhau có ghi rõ trong phần tài liệu "
    "tham khảo. Ngoài ra, trong đồ án còn sử dụng một số nhận xét, đánh "
    "giá cũng như số liệu của các tác giả khác, cơ quan tổ chức khác "
    "đều có trích dẫn và chú thích nguồn gốc. Nếu phát hiện có bất kỳ "
    "sự gian lận nào em xin hoàn toàn chịu trách nhiệm về nội dung đồ "
    "án của mình.")
add_p(doc, "", space_after=24)
add_p(doc, "TP. Hồ Chí Minh, tháng 5 năm 2026",
      align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=4)
add_p(doc, "Sinh viên thực hiện",
      align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True, space_after=36)
add_p(doc, "Nguyễn Nhật Hảo", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)


print("[OK] Chương 5, 6, TLTK, Phụ lục done")
doc.save(OUTPUT)
print(f"[OK] FINAL Saved: {OUTPUT}")
