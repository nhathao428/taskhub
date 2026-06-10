"""
Restructure BAO_CAO_DO_AN_CO_SO.docx theo Hướng dẫn DACS:
  - Page A4 (21 × 29.7 cm)
  - Margins T=2 B=2.5 L=3 R=2 cm
  - Font Times New Roman 13pt (Normal)
  - Heading 1 = 16pt bold (chương)
  - Heading 2 = 14pt bold (mục)
  - Heading 3 = 14pt italic, NOT bold (tiểu mục)
  - Page numbers center bottom
  - Lời cam đoan đặt sau LỜI CẢM ƠN (chuyển từ Phụ lục V)

Giữ nguyên nội dung. Chỉ thay đổi cấu trúc + format.
"""
from copy import deepcopy
import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'
DST = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'  # overwrite in place

d = docx.Document(SRC)

# ============================================================
# 1. Page size A4 + margins
# ============================================================
for s in d.sections:
    s.page_width = Cm(21.0)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(2.0)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0)
    s.right_margin = Cm(2.0)
print('section: A4 + margins T=2 B=2.5 L=3 R=2 OK')

# ============================================================
# 2. Heading sizes
# ============================================================
def set_style(name, size_pt, bold=None, italic=None):
    st = d.styles[name]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(size_pt)
    if bold is not None:
        st.font.bold = bold
    if italic is not None:
        st.font.italic = italic

set_style('Heading 1', 16, bold=True, italic=False)
set_style('Heading 2', 14, bold=True, italic=False)
set_style('Heading 3', 14, bold=False, italic=True)
set_style('Heading 4', 13, bold=False, italic=True)
print('headings: H1=16 bold, H2=14 bold, H3=14 italic, H4=13 italic')

# ============================================================
# 3. Page numbers center bottom (footer field)
# ============================================================
for s in d.sections:
    footer = s.footer
    # Clear existing then add page number field
    for p in list(footer.paragraphs):
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.text = ''
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # PAGE field
    fld_begin = OxmlElement('w:fldChar'); fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar'); fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
print('footer: page number center bottom')

# ============================================================
# 4. Move "Lời cam đoan" from Phụ lục V to front (after LỜI CẢM ƠN)
# ============================================================
# Find Phụ lục V heading + content range
paragraphs = d.paragraphs
phu_luc_v_idx = None
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 2' and 'cam đoan' in p.text.lower() and 'V.' in p.text:
        phu_luc_v_idx = i
        break

# Collect content from Phụ lục V until next Heading or end
moved_paras = []
if phu_luc_v_idx is not None:
    j = phu_luc_v_idx + 1
    while j < len(paragraphs):
        np = paragraphs[j]
        if np.style.name in ('Heading 1', 'Heading 2'):
            break
        if np.text.strip():
            moved_paras.append((np.style.name, np.text))
        j += 1
    # Remove from Phụ lục V (including the heading)
    for k in range(j - 1, phu_luc_v_idx - 1, -1):
        el = paragraphs[k]._element
        el.getparent().remove(el)
    print(f'removed {j - phu_luc_v_idx} paras from Phụ lục V (Lời cam đoan)')

# Re-find paragraphs after removal
paragraphs = d.paragraphs

# Insert Lời cam đoan AFTER LỜI CẢM ƠN section, BEFORE LỜI MỞ ĐẦU/MỤC LỤC
loi_cam_on_end = None
for i, p in enumerate(paragraphs):
    if p.style.name == 'Heading 1' and 'CẢM ƠN' in p.text.upper():
        # Find next heading 1 → that's where we insert before
        for k in range(i + 1, len(paragraphs)):
            if paragraphs[k].style.name == 'Heading 1':
                loi_cam_on_end = k
                break
        break

if loi_cam_on_end is not None and moved_paras:
    anchor = paragraphs[loi_cam_on_end]._element
    body = anchor.getparent()

    # Create heading "LỜI CAM ĐOAN" first
    h_p = d.styles['Heading 1']
    new_h = OxmlElement('w:p')
    new_pPr = OxmlElement('w:pPr')
    new_pStyle = OxmlElement('w:pStyle'); new_pStyle.set(qn('w:val'), 'Heading1')
    new_pPr.append(new_pStyle)
    new_h.append(new_pPr)
    new_r = OxmlElement('w:r'); new_t = OxmlElement('w:t'); new_t.text = 'LỜI CAM ĐOAN'
    new_r.append(new_t); new_h.append(new_r)
    anchor.addprevious(new_h)

    # Then add content paragraphs
    for style_name, text in moved_paras:
        new_p = OxmlElement('w:p')
        if style_name and style_name != 'Normal':
            pPr = OxmlElement('w:pPr')
            pStyle = OxmlElement('w:pStyle')
            pStyle.set(qn('w:val'), style_name.replace(' ', ''))
            pPr.append(pStyle); new_p.append(pPr)
        r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = text
        # Preserve space in text
        t.set(qn('xml:space'), 'preserve')
        r.append(t); new_p.append(r)
        anchor.addprevious(new_p)
    print(f'inserted LỜI CAM ĐOAN with {len(moved_paras)} content paras after LỜI CẢM ƠN')

d.save(DST)
print(f'\nSaved → {DST}')
