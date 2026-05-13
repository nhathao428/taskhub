"""
Sinh 4 file Word phiếu hành chính của Đồ án cơ sở thành các file độc lập:
  1. PHIEU_GIAO_DE_TAI.docx
  2. PHIEU_THEO_DOI_TIEN_DO.docx
  3. PHIEU_CHAM_DIEM_GVHD.docx
  4. PHIEU_CHAM_DIEM_GVPB.docx

Trước đây 4 phiếu được gộp chung trong BAO_CAO_DO_AN_CO_SO.docx; tách ra để
nộp riêng theo yêu cầu của Khoa Công nghệ Thông tin.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(r"C:\Users\Admin\task-management-system\docs")

FONT = "Times New Roman"
SIZE_BODY = 13


def set_run(run, size=SIZE_BODY, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def add_p(doc, text="", size=SIZE_BODY, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=0.0,
          space_before=0, space_after=0, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(indent_first)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, italic=italic)
    return p


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, size=SIZE_BODY, bold=True)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_run(run, size=SIZE_BODY)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    add_p(doc, "", space_after=4)
    return tbl


def init_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE_BODY)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    return doc


def add_header_vien(doc, title, sub=None, title_size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("KHOA CÔNG NGHỆ THÔNG TIN")
    set_run(run, size=13, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6 if sub else 18)
    run = p.add_run(title)
    set_run(run, size=title_size, bold=True)

    if sub:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(sub)
        set_run(run, size=14, bold=True)


# ==================================================================
# 1) PHIẾU GIAO ĐỀ TÀI
# ==================================================================
def build_phieu_giao_de_tai():
    doc = init_doc()
    add_header_vien(doc, "PHIẾU GIAO ĐỀ TÀI")

    add_p(doc, "TÊN MÔN HỌC: ĐỒ ÁN CƠ SỞ", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, "NGÀNH: CÔNG NGHỆ THÔNG TIN", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

    add_p(doc, "Họ và tên sinh viên được giao đề tài (sĩ số trong nhóm: 01):",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "• NGUYỄN NHẬT HẢO    MSSV: 2380612688    Lớp: 23DTHC1",
          align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=0, space_after=10)

    add_p(doc, "Tên đề tài: Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ "
               "Đa ngành Tích hợp AI",
          align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, space_after=8)

    add_p(doc, "Các dữ liệu ban đầu:", align=WD_ALIGN_PARAGRAPH.LEFT,
          bold=True, space_after=2)
    for s in [
        "- Khảo sát thực tế quy trình quản lý công việc tại doanh nghiệp nhỏ đa ngành.",
        "- Tài liệu về Spring Boot 3.x, React 18, Flutter 3.x, PostgreSQL 16, Docker.",
        "- Tài liệu kỹ thuật về OpenAI API và phương pháp tích hợp LLM vào ứng dụng nội bộ.",
        "- Các nghiên cứu về hệ thống gợi ý đa tiêu chí và xu hướng ứng dụng "
        "LLM vào tự động hoá phân công nhân sự.",
    ]:
        add_p(doc, s, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=0, space_after=2)

    add_p(doc, "Nội dung nhiệm vụ:", align=WD_ALIGN_PARAGRAPH.LEFT,
          bold=True, space_before=6, space_after=2)
    for s in [
        "- Phân tích yêu cầu, thiết kế kiến trúc hệ thống và cơ sở dữ liệu.",
        "- Triển khai backend RESTful API bằng Spring Boot, bảo mật bằng "
        "Spring Security + JWT.",
        "- Triển khai frontend web bằng React 18 + Vite + Tailwind CSS.",
        "- Triển khai ứng dụng mobile bằng Flutter cho iOS và Android.",
        "- Tích hợp module AI gợi ý nhân viên phù hợp dùng OpenAI GPT-4o-mini "
        "xếp hạng định tính dựa trên kỹ năng, tiến độ, đúng hạn và chấm công.",
        "- Đóng gói hệ thống bằng Docker Compose với PostgreSQL 16 và Redis 7.",
        "- Kiểm thử và viết tài liệu hướng dẫn cài đặt, vận hành.",
    ]:
        add_p(doc, s, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=0, space_after=2)

    add_p(doc, "Kết quả tối thiểu phải có:", align=WD_ALIGN_PARAGRAPH.LEFT,
          bold=True, space_before=6, space_after=2)
    for s in [
        "1) Mã nguồn hoàn chỉnh backend Spring Boot, frontend React và ứng dụng "
        "Flutter chạy được.",
        "2) Cơ sở dữ liệu PostgreSQL có script khởi tạo schema và dữ liệu mẫu.",
        "3) Tài liệu API (Swagger / OpenAPI) và tài liệu hướng dẫn cài đặt.",
        "4) Báo cáo đồ án đầy đủ theo mẫu của Khoa Công nghệ Thông tin.",
    ]:
        add_p(doc, s, align=WD_ALIGN_PARAGRAPH.LEFT, indent_first=0, space_after=2)

    add_p(doc, "Ngày giao đề tài: 03/02/2026    Ngày nộp báo cáo: 11/05/2026",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=24)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = True
    c1 = tbl.rows[0].cells[0]
    c2 = tbl.rows[0].cells[1]
    for cell, txt in [(c1, "Sinh viên thực hiện"),
                      (c2, "TP. HCM, ngày … tháng … năm 2026")]:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        set_run(r, bold=True)
    for cell, sub in zip([c1, c2],
                         ["(Ký và ghi rõ họ tên các thành viên)",
                          "Giảng viên hướng dẫn"]):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sub)
        set_run(r, italic=True)
    for cell, sig in zip([c1, c2],
                         ["Nguyễn Nhật Hảo", "ThS. Dương Thành Phết"]):
        for _ in range(4):
            cell.add_paragraph("")
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(sig)
        set_run(r, bold=True)

    out = OUT_DIR / "PHIEU_GIAO_DE_TAI.docx"
    doc.save(out)
    return out


# ==================================================================
# 2) PHIẾU THEO DÕI TIẾN ĐỘ
# ==================================================================
def build_phieu_theo_doi_tien_do():
    doc = init_doc()
    add_header_vien(
        doc, "PHIẾU THEO DÕI TIẾN ĐỘ",
        sub="THỰC HIỆN ĐỒ ÁN MÔN HỌC & ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN",
        title_size=15,
    )

    add_p(doc, "TÊN MÔN HỌC: ĐỒ ÁN CƠ SỞ", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT)
    add_p(doc, "NGÀNH: CÔNG NGHỆ THÔNG TIN", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)
    add_p(doc, "Tên đề tài: Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ "
               "Đa ngành Tích hợp AI",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Giảng viên hướng dẫn: ThS. Dương Thành Phết",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Sinh viên thực hiện: NGUYỄN NHẬT HẢO – MSSV: 2380612688 – "
               "Lớp: 23DTHC1",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    progress_rows = [
        ("1", "03/02/2026", "Giao đề tài",
         "Sinh viên nhận đề tài, xác định mục tiêu, phạm vi và kết quả tối thiểu."),
        ("2", "10/02/2026", "Khảo sát hiện trạng, thu thập yêu cầu",
         "Hoàn thành khảo sát quy trình quản lý công việc, tổng hợp yêu cầu "
         "chức năng và phi chức năng."),
        ("3", "17/02/2026", "Phân tích yêu cầu, vẽ Use Case",
         "Hoàn thành sơ đồ Use Case và đặc tả use case chi tiết cho 8 nhóm chức năng."),
        ("4", "24/02/2026", "Thiết kế CSDL, ERD",
         "Hoàn thành ERD với 6 bảng và mô tả chi tiết các trường."),
        ("5", "03/03/2026", "Thiết kế kiến trúc, Class & Sequence Diagram",
         "Hoàn thành sơ đồ lớp, sơ đồ tuần tự cho các luồng nghiệp vụ chính."),
        ("6", "10/03/2026", "Triển khai backend – Auth, JWT, User",
         "Cài đặt Spring Security + JWT, module User/Auth chạy được, test bằng Postman."),
        ("7", "17/03/2026", "Triển khai backend – Employee, Project, Task",
         "Hoàn thành CRUD đầy đủ cho Employee, Project, Task; cache Redis hoạt động."),
        ("8", "24/03/2026", "Triển khai backend – Attendance & AiSuggestionService",
         "Module chấm công và AI gợi ý hoàn thiện, tích hợp OpenAI GPT-4o-mini."),
        ("9", "31/03/2026", "Triển khai frontend React",
         "Hoàn thành các trang Login/Register/Dashboard/Employees/Projects/"
         "Tasks/Attendance/AI."),
        ("10", "14/04/2026", "Triển khai ứng dụng mobile Flutter",
         "Hoàn thành các màn hình chính cho Android, gọi backend qua dio."),
        ("11", "28/04/2026", "Kiểm thử tổng thể, viết tài liệu",
         "Kiểm thử 35 test cases, viết API Specification, Setup Guide, UML Diagrams."),
        ("12", "11/05/2026", "Hoàn thành và bảo vệ đồ án",
         "Đóng quyển báo cáo, chuẩn bị slide trình bày, bảo vệ đồ án trước hội đồng."),
    ]
    add_table(
        doc,
        headers=["Tuần", "Ngày", "Nội dung thực hiện",
                 "Kết quả thực hiện của sinh viên"],
        rows=progress_rows,
        col_widths=[1.4, 2.4, 5.2, 7.0],
    )

    add_p(doc, "Đánh giá kết quả báo cáo:", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4)
    add_p(doc,
          "Hình thức trình bày đúng quy định; nội dung chi tiết, đầy đủ các "
          "chương theo mẫu của Khoa Công nghệ Thông tin. Sản phẩm thực hiện được – "
          "backend, frontend, mobile, AI gợi ý chạy ổn định. Sinh viên có thái "
          "độ học tập tốt, chủ động trao đổi tiến độ hàng tuần với giảng viên "
          "hướng dẫn.",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)

    add_p(doc, "Cách tính điểm:", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_p(doc,
          "Điểm đánh giá quá trình thực hiện đồ án = 50% × Tính chủ động, tích "
          "cực, sáng tạo + 50% × Đáp ứng nội dung nhiệm vụ.",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=0, space_after=2)
    add_p(doc,
          "Tổng điểm kết thúc học phần = Điểm đánh giá quá trình × 40% + "
          "Điểm chấm báo cáo GVHD × 30% + Điểm chấm báo cáo GVPB × 30%.",
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=0, space_after=12)

    out = OUT_DIR / "PHIEU_THEO_DOI_TIEN_DO.docx"
    doc.save(out)
    return out


# ==================================================================
# 3) PHIẾU CHẤM ĐIỂM GVHD
# ==================================================================
def build_phieu_cham_gvhd():
    doc = init_doc()
    add_header_vien(
        doc, "PHIẾU CHẤM ĐIỂM ĐỒ ÁN MÔN HỌC – GIẢNG VIÊN HƯỚNG DẪN",
        title_size=14,
    )

    add_p(doc, "Họ và tên sinh viên: NGUYỄN NHẬT HẢO – MSSV: 2380612688 – Lớp: 23DTHC1",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Tên đề tài: Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ "
               "Đa ngành Tích hợp AI",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Giảng viên hướng dẫn: ThS. Dương Thành Phết",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    add_table(
        doc,
        headers=["STT", "Tiêu chí đánh giá", "Điểm tối đa", "Điểm chấm"],
        rows=[
            ("1", "Hình thức trình bày báo cáo (đúng mẫu, bố cục rõ ràng, "
                  "ít sai chính tả)", "1.0", ""),
            ("2", "Mức độ hoàn thành nội dung báo cáo so với nhiệm vụ được giao",
             "2.0", ""),
            ("3", "Kết quả thực hiện đề tài (sản phẩm chạy được, có demo)",
             "3.0", ""),
            ("4", "Tính sáng tạo, ứng dụng công nghệ mới (AI, microservice...)",
             "1.5", ""),
            ("5", "Mức độ phức tạp của giải pháp kỹ thuật", "1.0", ""),
            ("6", "Khả năng trình bày, trả lời câu hỏi", "1.5", ""),
            ("TỔNG", "", "10.0", ""),
        ],
        col_widths=[1.3, 9.5, 2.5, 2.5],
    )

    add_p(doc, "Nhận xét của giảng viên hướng dẫn:", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
    for _ in range(8):
        add_p(doc, "." * 90, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_p(doc, "Điểm số (bằng chữ): ......................................................",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=18)
    add_p(doc, "TP. HCM, ngày … tháng … năm 2026",
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    add_p(doc, "Giảng viên hướng dẫn", bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=36)
    add_p(doc, "ThS. Dương Thành Phết", bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    out = OUT_DIR / "PHIEU_CHAM_DIEM_GVHD.docx"
    doc.save(out)
    return out


# ==================================================================
# 4) PHIẾU CHẤM ĐIỂM GVPB
# ==================================================================
def build_phieu_cham_gvpb():
    doc = init_doc()
    add_header_vien(
        doc, "PHIẾU CHẤM ĐIỂM ĐỒ ÁN MÔN HỌC – GIẢNG VIÊN PHẢN BIỆN",
        title_size=14,
    )

    add_p(doc, "Họ và tên sinh viên: NGUYỄN NHẬT HẢO – MSSV: 2380612688 – Lớp: 23DTHC1",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Tên đề tài: Hệ thống Quản lý Công việc cho Doanh nghiệp Nhỏ "
               "Đa ngành Tích hợp AI",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    add_p(doc, "Giảng viên phản biện: ………………………………………………………",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

    add_table(
        doc,
        headers=["STT", "Tiêu chí đánh giá", "Điểm tối đa", "Điểm chấm"],
        rows=[
            ("1", "Hình thức trình bày báo cáo", "1.0", ""),
            ("2", "Tính hợp lý của giải pháp – kiến trúc – thuật toán", "2.5", ""),
            ("3", "Kết quả thực hiện đề tài và demo sản phẩm", "3.0", ""),
            ("4", "Tính sáng tạo, ứng dụng công nghệ mới", "1.5", ""),
            ("5", "Khả năng trình bày, trả lời câu hỏi phản biện", "2.0", ""),
            ("TỔNG", "", "10.0", ""),
        ],
        col_widths=[1.3, 9.5, 2.5, 2.5],
    )

    add_p(doc, "Nhận xét của giảng viên phản biện:", bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4)
    for _ in range(8):
        add_p(doc, "." * 90, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

    add_p(doc, "Điểm số (bằng chữ): ......................................................",
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=18)
    add_p(doc, "TP. HCM, ngày … tháng … năm 2026",
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    add_p(doc, "Giảng viên phản biện", bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=36)
    add_p(doc, "(Ký và ghi rõ họ tên)", italic=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    out = OUT_DIR / "PHIEU_CHAM_DIEM_GVPB.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    files = [
        build_phieu_giao_de_tai(),
        build_phieu_theo_doi_tien_do(),
        build_phieu_cham_gvhd(),
        build_phieu_cham_gvpb(),
    ]
    for f in files:
        print(f"[OK] {f.name}")
