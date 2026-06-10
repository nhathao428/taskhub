"""
Chèn ảnh demo vào Chương 5.
Mapping:
  5.4.1 ↔ 01_login.png
  5.4.2 ↔ 02_register.png
  5.4.3 ↔ 03_dashboard.png
  5.4.4 ↔ 04_employees.png
  5.4.5 ↔ 05_projects.png
  5.4.6 ↔ 06_tasks.png
  5.4.7 ↔ 07_attendance.png
  5.4.8 ↔ 08_ai_suggestions.png
  5.4.9 ↔ 09_ai_result.png
  5.4.10 ↔ 10_emp_dashboard.png + 11_emp_my_tasks.png + 12_emp_my_attendance.png

  5.5 (mobile demo) ↔ m01_login.png + m03_dashboard.png + m08_ai_suggestions.png
"""
import re
import os
import docx
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

PATH = r'C:\Users\Admin\taskhub\docs\BAO_CAO_DO_AN_CO_SO.docx'
SCREENSHOT_DIR = r'C:\Users\Admin\taskhub\docs\screenshots'
MOBILE_DIR = os.path.join(SCREENSHOT_DIR, 'mobile')

d = docx.Document(PATH)

# Mapping: heading prefix → list of (image_path, caption)
MAPPING = {
    '5.4.1.': [('01_login.png', 'Hình 5.1: Màn hình Đăng nhập (Web)')],
    '5.4.2.': [('02_register.png', 'Hình 5.2: Màn hình Đăng ký tài khoản (Web)')],
    '5.4.3.': [('03_dashboard.png', 'Hình 5.3: Dashboard tổng quan (Quản lý)')],
    '5.4.4.': [('04_employees.png', 'Hình 5.4: Quản lý Nhân viên')],
    '5.4.5.': [('05_projects.png', 'Hình 5.5: Quản lý Dự án')],
    '5.4.6.': [('06_tasks.png', 'Hình 5.6: Quản lý Công việc')],
    '5.4.7.': [('07_attendance.png', 'Hình 5.7: Trang Chấm công')],
    '5.4.8.': [('08_ai_suggestions.png', 'Hình 5.8: AI Gợi ý nhân viên — form nhập')],
    '5.4.9.': [('09_ai_result.png', 'Hình 5.9: AI Gợi ý nhân viên — kết quả phân tích')],
    '5.4.10.': [
        ('10_emp_dashboard.png', 'Hình 5.10: Góc nhìn nhân viên — Dashboard'),
        ('11_emp_my_tasks.png', 'Hình 5.11: Góc nhìn nhân viên — Công việc của tôi'),
        ('12_emp_my_attendance.png', 'Hình 5.12: Góc nhìn nhân viên — Chấm công của tôi'),
    ],
    '5.5.1.': [
        ('m01_login.png', 'Hình 5.13: Mobile — Đăng nhập', True),
        ('m02_register.png', 'Hình 5.14: Mobile — Đăng ký', True),
    ],
}


def img_path(name, mobile=False):
    return os.path.join(MOBILE_DIR if mobile else SCREENSHOT_DIR, name)


# Find Heading 3 in Ch5 and insert image + caption after it
body = d.element.body
children = list(body)

# Track inserts: position, image_path, caption
to_insert = []
for i, child in enumerate(children):
    if not child.tag.endswith('}p'): continue
    p = Paragraph(child, d)
    if p.style.name != 'Heading 3': continue
    text = p.text.strip()
    for prefix, items in MAPPING.items():
        if text.startswith(prefix):
            to_insert.append((child, items))
            break


def make_caption_para(text):
    """Tạo paragraph caption — italic, center."""
    p_el = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle'); pStyle.set(qn('w:val'), 'Caption')
    pPr.append(pStyle)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p_el.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_el = OxmlElement('w:i'); rPr.append(i_el)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '24'); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t'); t.text = text; r.append(t)
    p_el.append(r)
    return p_el


def make_img_para(image_path, width_cm=14):
    """Tạo paragraph chứa ảnh + center alignment."""
    # Use python-docx to insert picture: add a new paragraph to doc, get its element, then move it
    para = d.add_paragraph()
    para.alignment = 1  # center
    run = para.add_run()
    try:
        run.add_picture(image_path, width=Cm(width_cm))
    except Exception as e:
        print(f'  ERROR adding {image_path}: {e}')
        # Remove the failed paragraph
        para._element.getparent().remove(para._element)
        return None
    # Detach from body so we can insert wherever
    p_el = para._element
    p_el.getparent().remove(p_el)
    return p_el


inserted_count = 0
for anchor_h3, items in to_insert:
    cursor = anchor_h3
    for item in items:
        if len(item) == 3:
            fname, caption, mobile = item
        else:
            fname, caption = item; mobile = False
        path = img_path(fname, mobile=mobile)
        if not os.path.exists(path):
            print(f'  MISSING: {path}')
            continue
        # Image (smaller for mobile portrait)
        width = 8 if mobile else 14
        img_p = make_img_para(path, width_cm=width)
        if img_p is None: continue
        cap_p = make_caption_para(caption)
        cursor.addnext(img_p)
        img_p.addnext(cap_p)
        cursor = cap_p
        inserted_count += 1

print(f'Inserted {inserted_count} screenshots')

d.save(PATH)
print(f'Saved → {PATH}')
